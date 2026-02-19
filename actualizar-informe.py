from docx import Document

doc = Document('/home/tunek/Universidad/MATERIAS/bd-clinica/Informe_Hospital_NE.docx')

# 1. Actualizar fecha
for p in doc.paragraphs:
    if '18/02/2025' in p.text:
        for run in p.runs:
            if '18/02/2025' in run.text:
                run.text = run.text.replace('18/02/2025', '18/02/2026')
        print('Fecha actualizada: 2025 -> 2026')

# 2. Corregir tablas del diccionario de datos
for ti, table in enumerate(doc.tables):
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]

        # Tabla PERSONA: Nombre VARCHAR(200) -> VARCHAR(150)
        if cells[0] == 'Nombre' and cells[1] == 'VARCHAR(200)':
            for run in row.cells[1].paragraphs[0].runs:
                if 'VARCHAR(200)' in run.text:
                    run.text = run.text.replace('VARCHAR(200)', 'VARCHAR(150)')
            print(f'Tabla {ti}: Nombre VARCHAR(200) -> VARCHAR(150)')

        # Tabla PERSONA: Direccion TEXT -> VARCHAR(255), quitar "opcional"
        if cells[0] == 'Direccion' and 'TEXT' in cells[1]:
            for run in row.cells[1].paragraphs[0].runs:
                if 'TEXT' in run.text:
                    run.text = run.text.replace('TEXT', 'VARCHAR(255)')
            for run in row.cells[2].paragraphs[0].runs:
                run.text = run.text.replace('(opcional)', '').replace('opcional', '').strip()
                if not run.text.endswith('.'):
                    run.text = run.text.rstrip('.') + '. No nulo.'
            print(f'Tabla {ti}: Direccion TEXT -> VARCHAR(255), NOT NULL')

        # Tabla PERSONA: Telefono quitar "opcional"
        if cells[0] == 'Telefono' and 'opcional' in cells[2].lower():
            for run in row.cells[2].paragraphs[0].runs:
                run.text = run.text.replace('(opcional)', '').replace('opcional', '').strip()
                if not run.text.endswith('.'):
                    run.text = run.text.rstrip('.') + '. No nulo.'
            print(f'Tabla {ti}: Telefono marcado como NOT NULL')

        # Tabla TIPO_DIAGNOSTICO: Categoria VARCHAR(50) -> VARCHAR(100)
        if cells[0] == 'Categoria' and cells[1] == 'VARCHAR(50)':
            for run in row.cells[1].paragraphs[0].runs:
                if 'VARCHAR(50)' in run.text:
                    run.text = run.text.replace('VARCHAR(50)', 'VARCHAR(100)')
            print(f'Tabla {ti}: Categoria VARCHAR(50) -> VARCHAR(100)')

        # Tabla CITA_MEDICA: Estado VARCHAR(20) -> VARCHAR(50)
        if cells[0] == 'Estado' and cells[1] == 'VARCHAR(20)':
            for run in row.cells[1].paragraphs[0].runs:
                if 'VARCHAR(20)' in run.text:
                    run.text = run.text.replace('VARCHAR(20)', 'VARCHAR(50)')
            print(f'Tabla {ti}: Estado VARCHAR(20) -> VARCHAR(50)')

doc.save('/home/tunek/Universidad/MATERIAS/bd-clinica/Informe_Hospital_NE.docx')
print('\nDocumento actualizado y guardado.')
