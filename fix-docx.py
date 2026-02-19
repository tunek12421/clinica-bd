from docx import Document

doc = Document('/home/tunek/Universidad/MATERIAS/bd-clinica/Informe_Hospital_NE.docx')

fixes = {
    ('Nombre', 'VARCHAR(200)'): 'VARCHAR(150)',
    ('Categoria', 'VARCHAR(50)'): 'VARCHAR(100)',
    ('Estado', 'VARCHAR(20)'): 'VARCHAR(50)',
}

for table in doc.tables:
    for row in table.rows:
        attr = row.cells[0].text.strip()
        tipo = row.cells[1].text.strip()
        key = (attr, tipo)
        if key in fixes:
            new_val = fixes[key]
            # Reemplazar en todos los runs de la celda de tipo
            for p in row.cells[1].paragraphs:
                for run in p.runs:
                    if tipo in run.text:
                        run.text = run.text.replace(tipo, new_val)
                # Si no habia runs con el texto, forzar
                if p.text.strip() == tipo:
                    p.runs[0].text = new_val
                    for run in p.runs[1:]:
                        run.text = ''
            print(f'{attr}: {tipo} -> {new_val}')

doc.save('/home/tunek/Universidad/MATERIAS/bd-clinica/Informe_Hospital_NE.docx')
print('Corregido.')
