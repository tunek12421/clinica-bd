#!/usr/bin/env python3
"""
Genera 4 documentos:
1. Informe_ETL_Grupo6.docx  (G3 + G6)
2. graficas_etl_grupo6.xlsx (G3 + G6)
3. Informe_ETL_Completo.docx (G3 + G1 + G6)
4. graficas_etl_completo.xlsx (G3 + G1 + G6)
"""

import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from openpyxl import Workbook
from openpyxl.drawing.image import Image as XlImage
from openpyxl.styles import Font as XlFont, PatternFill, Alignment, Border, Side
import os

BASE = '/home/tunek/Universidad/MATERIAS/bd-clinica'
IMG = '/tmp/etl_docs_charts'
os.makedirs(IMG, exist_ok=True)

# =============================================================================
# DATOS COMUNES (ANTES = Solo Grupo 3)
# =============================================================================
esp_antes = [
    ('Medicina General', 3542), ('Cardiología', 3097), ('Pediatría', 2170),
    ('Dermatología', 1905), ('Neurología', 1764), ('Ginecología', 1660),
    ('Oftalmología', 1220), ('Traumatología', 1001), ('Otorrinolaringología', 760),
    ('Urología', 741), ('Gastroenterología', 706), ('Neumología', 483),
    ('Psiquiatría', 426), ('Endocrinología', 331), ('Oncología', 194),
]
meses_n = ['Ene','Feb','Mar','Abr','May','Jun','Jul','Ago','Sep','Oct','Nov','Dic']
mes_antes = [1432, 1432, 1707, 1682, 1780, 2406, 2573, 2660, 1230, 1178, 1198, 722]
diag_antes = [
    ('Clínico', 2142), ('Por Imagen', 1735), ('Laboratorio', 1389),
    ('Diferencial', 1228), ('Presuntivo', 1084), ('Definitivo', 1027),
    ('Prenatal', 890), ('Molecular', 750), ('Patológico', 737), ('Funcional', 620),
]

# --- G3 + G6 ---
esp_g6 = [
    ('Medicina General', 65652), ('Traumatología', 24485), ('Urología', 14810),
    ('Cardiología', 11566), ('Neurología', 10712), ('Oftalmología', 10706),
    ('Pediatría', 10301), ('Oncología', 10229), ('Dermatología', 9542),
    ('Ginecología', 8950), ('Otorrinolaringología', 8756), ('Endocrinología', 6514),
    ('Psiquiatría', 6236), ('Gastroenterología', 706), ('Neumología', 483),
]
mes_g6 = [18347, 17043, 18595, 18179, 18649, 19070, 19412, 19454, 17652, 18079, 17796, 17724]
diag_g6 = [
    ('Definitivo', 43780), ('Presuntivo', 43721), ('Clínico', 2142),
    ('Por Imagen', 1735), ('Laboratorio', 1389), ('Diferencial', 1228),
    ('Prenatal', 890), ('Molecular', 750), ('Patológico', 737), ('Funcional', 620),
]

# --- G3 + G1 + G6 ---
esp_all = [
    ('Medicina General', 71813), ('Traumatología', 30307), ('Urología', 20429),
    ('Cardiología', 18921), ('Pediatría', 17088), ('Neurología', 16830),
    ('Dermatología', 16157), ('Ginecología', 14473), ('Oftalmología', 10706),
    ('Oncología', 10229), ('Otorrinolaringología', 8756), ('Endocrinología', 6514),
    ('Psiquiatría', 6236), ('Gastroenterología', 706), ('Neumología', 483),
]
mes_all = [23268, 21466, 23532, 22924, 22545, 22796, 23362, 23357, 21417, 22013, 21623, 21697]
diag_all = [
    ('Definitivo', 43780), ('Presuntivo', 43721), ('Clínico', 33940),
    ('Laboratorio', 15025), ('Ambulatorio', 8954), ('Por Imagen', 6307),
    ('Nutricional', 4802), ('Endoscópico', 4628), ('Diferencial', 1228), ('Prenatal', 890),
]

# =============================================================================
# FUNCIONES GRÁFICOS
# =============================================================================
plt.rcParams['font.size'] = 10; plt.rcParams['figure.dpi'] = 150
def clist(n, cm='tab20'): c = plt.get_cmap(cm); return [c(i/n) for i in range(n)]

def gen_esp(antes, despues, tag, title_d):
    fig, ax = plt.subplots(figsize=(10,5))
    n=[e[0] for e in antes]; v=[e[1] for e in antes]
    ax.barh(n[::-1],v[::-1],color=clist(len(n))[::-1])
    ax.set_title('ANTES (Solo Grupo 3)',fontweight='bold',fontsize=12)
    for i,val in enumerate(v[::-1]): ax.text(val+50,i,f'{val:,}',va='center',fontsize=8)
    plt.tight_layout(); plt.savefig(f'{IMG}/{tag}_esp_a.png'); plt.close()

    fig, ax = plt.subplots(figsize=(10,5))
    n=[e[0] for e in despues]; v=[e[1] for e in despues]
    ax.barh(n[::-1],v[::-1],color=clist(len(n))[::-1])
    ax.set_title(f'DESPUÉS ({title_d})',fontweight='bold',fontsize=12)
    for i,val in enumerate(v[::-1]): ax.text(val+max(v)*0.01,i,f'{val:,}',va='center',fontsize=8)
    plt.tight_layout(); plt.savefig(f'{IMG}/{tag}_esp_d.png'); plt.close()

def gen_mes(antes, despues, tag, title_d):
    fig, ax = plt.subplots(figsize=(10,4.5)); x=np.arange(12)
    ax.bar(x,antes,color='#4472C4',width=0.6); ax.set_xticks(x); ax.set_xticklabels(meses_n)
    ax.set_title('ANTES (Solo Grupo 3)',fontweight='bold',fontsize=12)
    for i,v in enumerate(antes): ax.text(i,v+max(antes)*0.02,f'{v:,}',ha='center',fontsize=8)
    plt.tight_layout(); plt.savefig(f'{IMG}/{tag}_mes_a.png'); plt.close()

    fig, ax = plt.subplots(figsize=(10,4.5))
    ax.bar(x,despues,color='#ED7D31',width=0.6); ax.set_xticks(x); ax.set_xticklabels(meses_n)
    ax.set_title(f'DESPUÉS ({title_d})',fontweight='bold',fontsize=12)
    for i,v in enumerate(despues): ax.text(i,v+max(despues)*0.02,f'{v:,}',ha='center',fontsize=8)
    plt.tight_layout(); plt.savefig(f'{IMG}/{tag}_mes_d.png'); plt.close()

def gen_diag(antes, despues, tag, title_d):
    fig, ax = plt.subplots(figsize=(10,5))
    n=[d[0] for d in antes]; v=[d[1] for d in antes]
    ax.barh(n[::-1],v[::-1],color=clist(len(n),'Set3')[::-1])
    ax.set_title('ANTES (Solo Grupo 3)',fontweight='bold',fontsize=12)
    for i,val in enumerate(v[::-1]): ax.text(val+max(v)*0.01,i,f'{val:,}',va='center',fontsize=8)
    plt.tight_layout(); plt.savefig(f'{IMG}/{tag}_diag_a.png'); plt.close()

    fig, ax = plt.subplots(figsize=(10,5))
    n=[d[0] for d in despues]; v=[d[1] for d in despues]
    ax.barh(n[::-1],v[::-1],color=clist(len(n),'Set3')[::-1])
    ax.set_title(f'DESPUÉS ({title_d})',fontweight='bold',fontsize=12)
    for i,val in enumerate(v[::-1]): ax.text(val+max(v)*0.01,i,f'{val:,}',va='center',fontsize=8)
    plt.tight_layout(); plt.savefig(f'{IMG}/{tag}_diag_d.png'); plt.close()

# =============================================================================
# FUNCIONES WORD
# =============================================================================
def make_word(filename, grupo_label, subtitle, esp_d, mes_d, diag_d,
              tabla_antes, tabla_despues, esquema_fuente, analisis_text,
              transform_text, obs_esp, obs_mes, obs_diag, tag, total_citas, total_diag):
    doc = Document()
    st = doc.styles['Normal']; st.font.name='Calibri'; st.font.size=Pt(11)
    st.paragraph_format.space_after = Pt(6)

    def ah(text, lv=1):
        h=doc.add_heading(text,level=lv)
        for r in h.runs: r.font.color.rgb=RGBColor(0x1F,0x3A,0x5F)

    def at(headers, rows):
        t=doc.add_table(rows=1+len(rows),cols=len(headers),style='Light Grid Accent 1')
        t.alignment=WD_TABLE_ALIGNMENT.CENTER
        for i,h in enumerate(headers):
            c=t.rows[0].cells[i]; c.text=h
            for p in c.paragraphs:
                p.alignment=WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs: r.bold=True; r.font.size=Pt(9)
        for ri,row in enumerate(rows):
            for ci,v in enumerate(row):
                c=t.rows[ri+1].cells[ci]; c.text=str(v)
                for p in c.paragraphs:
                    for r in p.runs: r.font.size=Pt(9)

    # Portada
    doc.add_paragraph(); doc.add_paragraph()
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run('Informe de Reconocimiento y Migración ETL'); r.bold=True; r.font.size=Pt(24); r.font.color.rgb=RGBColor(0x1F,0x3A,0x5F)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(subtitle); r.font.size=Pt(16); r.font.color.rgb=RGBColor(0x4F,0x81,0xBD)
    doc.add_paragraph()
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('Materia: Base de Datos\n').font.size=Pt(12)
    p.add_run('Universidad Privada Domingo Savio\n').font.size=Pt(12)
    p.add_run('Febrero 2026').font.size=Pt(12)
    doc.add_page_break()

    # Sec 1: Análisis
    ah('1. Análisis Comparativo')
    doc.add_paragraph(analisis_text)

    ah('1.1 Esquema del Grupo 3 (Destino)', 2)
    doc.add_paragraph('Base de datos PostgreSQL 17 con 8 tablas normalizadas (3FN).')
    at(['Tabla','Registros'], [
        ('ESPECIALIDAD','15'), ('TIPO_DIAGNOSTICO','25'), ('ZONA','20'),
        ('PERSONA','5,000'), ('HORARIO_MEDICO','2,500'), ('CITA_MEDICA','20,000'),
        ('DIAGNOSTICO','15,000'), ('RECETA','7,440'),
    ])

    ah(f'1.2 Esquema del {grupo_label} (Fuente)', 2)
    for row in esquema_fuente:
        doc.add_paragraph(row)

    ah('1.3 Diferencias Identificadas', 2)
    at(*tabla_antes)

    ah('1.4 Resultado de la Carga', 2)
    at(*tabla_despues)

    doc.add_page_break()

    # Sec 2: Transformaciones
    ah('2. Información a Transformar')
    for t in transform_text:
        if t.startswith('##'):
            ah(t.replace('## ',''), 2)
        elif t.startswith('### '):
            ah(t.replace('### ',''), 3)
        elif t.startswith('|'):
            # Parse table
            lines = t.strip().split('\n')
            headers = [c.strip() for c in lines[0].split('|')[1:-1]]
            rows = []
            for l in lines[2:]:
                rows.append([c.strip() for c in l.split('|')[1:-1]])
            at(headers, rows)
        else:
            doc.add_paragraph(t)

    doc.add_page_break()

    # Sec 3: Gráficos
    ah('3. Gráficos Estadísticos: Antes y Después')
    doc.add_paragraph(f'Comparación de datos originales (Grupo 3) vs consolidados ({grupo_label}).')

    ah('3.1 Citas por Especialidad', 2)
    ah(f'Antes (Solo Grupo 3 — 20,000 citas)', 3)
    doc.add_picture(f'{IMG}/{tag}_esp_a.png', width=Inches(6))
    doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
    ah(f'Después ({grupo_label} — {total_citas:,} citas)', 3)
    doc.add_picture(f'{IMG}/{tag}_esp_d.png', width=Inches(6))
    doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(obs_esp)

    ah('3.2 Citas por Mes', 2)
    ah(f'Antes (Solo Grupo 3 — 20,000 citas)', 3)
    doc.add_picture(f'{IMG}/{tag}_mes_a.png', width=Inches(6))
    doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
    ah(f'Después ({grupo_label} — {total_citas:,} citas)', 3)
    doc.add_picture(f'{IMG}/{tag}_mes_d.png', width=Inches(6))
    doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(obs_mes)

    ah('3.3 Diagnósticos por Tipo', 2)
    ah(f'Antes (Solo Grupo 3 — 15,000 diagnósticos)', 3)
    doc.add_picture(f'{IMG}/{tag}_diag_a.png', width=Inches(6))
    doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
    ah(f'Después ({grupo_label} — {total_diag:,} diagnósticos)', 3)
    doc.add_picture(f'{IMG}/{tag}_diag_d.png', width=Inches(6))
    doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(obs_diag)

    doc.save(f'{BASE}/{filename}')
    print(f"Word: {filename}")

# =============================================================================
# FUNCIONES EXCEL
# =============================================================================
hf=XlFont(bold=True,color='FFFFFF',size=11)
fa=PatternFill(start_color='4472C4',end_color='4472C4',fill_type='solid')
fd=PatternFill(start_color='ED7D31',end_color='ED7D31',fill_type='solid')
tf=XlFont(bold=True,size=13,color='1F3A5F')
bdr=Border(left=Side('thin'),right=Side('thin'),top=Side('thin'),bottom=Side('thin'))

def sh(ws,r,c1,c2,f):
    for c in range(c1,c2+1):
        cl=ws.cell(r,c); cl.font=hf; cl.fill=f; cl.alignment=Alignment(horizontal='center'); cl.border=bdr

def sd(ws,r,c1,c2):
    for c in range(c1,c2+1):
        cl=ws.cell(r,c); cl.border=bdr
        if c>c1: cl.number_format='#,##0'; cl.alignment=Alignment(horizontal='right')

def make_excel(filename, esp_d, mes_d, diag_d, label, tag):
    wb=Workbook()
    # Hoja 1
    ws=wb.active; ws.title='Citas por Especialidad'
    ws.cell(1,1,f'Citas por Especialidad — Antes y Después ({label})').font=tf; ws.merge_cells('A1:F1')
    ws.cell(3,1,'Especialidad'); ws.cell(3,2,'Antes (G3)'); sh(ws,3,1,2,fa)
    for i,(n,v) in enumerate(esp_antes):
        ws.cell(4+i,1,n); ws.cell(4+i,2,v); sd(ws,4+i,1,2)
    r=4+len(esp_antes)
    ws.cell(r,1,'TOTAL').font=XlFont(bold=True); ws.cell(r,2,sum(v for _,v in esp_antes)); ws.cell(r,2).font=XlFont(bold=True); ws.cell(r,2).number_format='#,##0'; sd(ws,r,1,2)

    ws.cell(3,4,'Especialidad'); ws.cell(3,5,f'Después ({label})'); ws.cell(3,6,'Diferencia'); sh(ws,3,4,6,fd)
    ad=dict(esp_antes)
    for i,(n,v) in enumerate(esp_d):
        ws.cell(4+i,4,n); ws.cell(4+i,5,v); d=v-ad.get(n,0); ws.cell(4+i,6,f'+{d:,}' if d>0 else str(d)); sd(ws,4+i,4,6)
    r=4+len(esp_d)
    ws.cell(r,4,'TOTAL').font=XlFont(bold=True); ws.cell(r,5,sum(v for _,v in esp_d)); ws.cell(r,5).font=XlFont(bold=True); ws.cell(r,5).number_format='#,##0'; sd(ws,r,4,6)
    for c,w in [('A',22),('B',16),('D',22),('E',20),('F',14)]: ws.column_dimensions[c].width=w
    im=XlImage(f'{IMG}/{tag}_esp_comp.png'); im.width=900; im.height=400; ws.add_image(im,'A22')

    # Hoja 2
    ws2=wb.create_sheet('Citas por Mes')
    ws2.cell(1,1,f'Citas por Mes — Antes y Después ({label})').font=tf; ws2.merge_cells('A1:E1')
    ws2.cell(3,1,'Mes'); ws2.cell(3,2,'Antes (G3)'); sh(ws2,3,1,2,fa)
    ws2.cell(3,3,f'Después ({label})'); ws2.cell(3,4,'Diferencia'); ws2.cell(3,5,'%'); sh(ws2,3,3,5,fd)
    for i,(m,a,d) in enumerate(zip(meses_n,mes_antes,mes_d)):
        ws2.cell(4+i,1,m); ws2.cell(4+i,2,a); ws2.cell(4+i,3,d)
        ws2.cell(4+i,4,f'+{d-a:,}'); ws2.cell(4+i,5,f'{((d-a)/a*100):.0f}%'); sd(ws2,4+i,1,5)
    r=4+len(meses_n)
    ws2.cell(r,1,'TOTAL').font=XlFont(bold=True)
    ws2.cell(r,2,sum(mes_antes)); ws2.cell(r,2).font=XlFont(bold=True); ws2.cell(r,2).number_format='#,##0'
    ws2.cell(r,3,sum(mes_d)); ws2.cell(r,3).font=XlFont(bold=True); ws2.cell(r,3).number_format='#,##0'
    sa=sum(mes_antes); sd2=sum(mes_d)
    ws2.cell(r,4,f'+{sd2-sa:,}').font=XlFont(bold=True); ws2.cell(r,5,f'{((sd2-sa)/sa*100):.0f}%').font=XlFont(bold=True); sd(ws2,r,1,5)
    for c,w in [('A',10),('B',16),('C',20),('D',14),('E',12)]: ws2.column_dimensions[c].width=w
    im2=XlImage(f'{IMG}/{tag}_mes_comp.png'); im2.width=850; im2.height=380; ws2.add_image(im2,'A19')

    # Hoja 3
    ws3=wb.create_sheet('Diagnósticos por Tipo')
    ws3.cell(1,1,f'Diagnósticos por Tipo (Top 10) — Antes y Después ({label})').font=tf; ws3.merge_cells('A1:E1')
    ws3.cell(3,1,'Tipo'); ws3.cell(3,2,'Antes (G3)'); sh(ws3,3,1,2,fa)
    for i,(n,v) in enumerate(diag_antes):
        ws3.cell(4+i,1,n); ws3.cell(4+i,2,v); sd(ws3,4+i,1,2)
    r=4+len(diag_antes)
    ws3.cell(r,1,'TOTAL').font=XlFont(bold=True); ws3.cell(r,2,sum(v for _,v in diag_antes)); ws3.cell(r,2).font=XlFont(bold=True); ws3.cell(r,2).number_format='#,##0'; sd(ws3,r,1,2)

    ws3.cell(3,4,'Tipo'); ws3.cell(3,5,f'Después ({label})'); sh(ws3,3,4,5,fd)
    for i,(n,v) in enumerate(diag_d):
        ws3.cell(4+i,4,n); ws3.cell(4+i,5,v); sd(ws3,4+i,4,5)
    r=4+len(diag_d)
    ws3.cell(r,4,'TOTAL').font=XlFont(bold=True); ws3.cell(r,5,sum(v for _,v in diag_d)); ws3.cell(r,5).font=XlFont(bold=True); ws3.cell(r,5).number_format='#,##0'; sd(ws3,r,4,5)
    for c,w in [('A',20),('B',16),('D',20),('E',20)]: ws3.column_dimensions[c].width=w
    im3=XlImage(f'{IMG}/{tag}_diag_comp.png'); im3.width=900; im3.height=380; ws3.add_image(im3,'A17')

    wb.save(f'{BASE}/{filename}')
    print(f"Excel: {filename}")

def gen_comp_charts(antes_e, desp_e, antes_m, desp_m, antes_d, desp_d, tag, lbl):
    """Generate comparison charts for Excel"""
    fig,(a1,a2)=plt.subplots(1,2,figsize=(16,7))
    n=[e[0] for e in antes_e]; v=[e[1] for e in antes_e]
    a1.barh(n[::-1],v[::-1],color=clist(len(n))[::-1])
    a1.set_title('ANTES (Solo Grupo 3)',fontweight='bold',color='#4472C4')
    for i,val in enumerate(v[::-1]): a1.text(val+50,i,f'{val:,}',va='center',fontsize=8)
    n=[e[0] for e in desp_e]; v=[e[1] for e in desp_e]
    a2.barh(n[::-1],v[::-1],color=clist(len(n))[::-1])
    a2.set_title(f'DESPUÉS ({lbl})',fontweight='bold',color='#ED7D31')
    for i,val in enumerate(v[::-1]): a2.text(val+max(v)*0.01,i,f'{val:,}',va='center',fontsize=8)
    fig.suptitle('Citas por Especialidad',fontsize=14,fontweight='bold')
    plt.tight_layout(); plt.savefig(f'{IMG}/{tag}_esp_comp.png',bbox_inches='tight'); plt.close()

    fig,ax=plt.subplots(figsize=(14,6)); x=np.arange(12); w=0.35
    ax.bar(x-w/2,antes_m,w,label='Antes (G3)',color='#4472C4')
    ax.bar(x+w/2,desp_m,w,label=f'Después ({lbl})',color='#ED7D31')
    ax.set_xticks(x); ax.set_xticklabels(meses_n); ax.legend()
    ax.set_title('Citas por Mes',fontweight='bold',fontsize=14)
    for b in ax.patches[:12]: ax.text(b.get_x()+b.get_width()/2,b.get_height()+max(desp_m)*0.01,f'{int(b.get_height()):,}',ha='center',fontsize=7,color='#4472C4')
    for b in ax.patches[12:]: ax.text(b.get_x()+b.get_width()/2,b.get_height()+max(desp_m)*0.01,f'{int(b.get_height()):,}',ha='center',fontsize=7,color='#ED7D31')
    plt.tight_layout(); plt.savefig(f'{IMG}/{tag}_mes_comp.png',bbox_inches='tight'); plt.close()

    fig,(a1,a2)=plt.subplots(1,2,figsize=(16,6))
    n=[d[0] for d in antes_d]; v=[d[1] for d in antes_d]
    a1.barh(n[::-1],v[::-1],color=clist(len(n),'Set3')[::-1])
    a1.set_title('ANTES (Solo Grupo 3)',fontweight='bold',color='#4472C4')
    for i,val in enumerate(v[::-1]): a1.text(val+max(v)*0.01,i,f'{val:,}',va='center',fontsize=8)
    n=[d[0] for d in desp_d]; v=[d[1] for d in desp_d]
    a2.barh(n[::-1],v[::-1],color=clist(len(n),'Set3')[::-1])
    a2.set_title(f'DESPUÉS ({lbl})',fontweight='bold',color='#ED7D31')
    for i,val in enumerate(v[::-1]): a2.text(val+max(v)*0.01,i,f'{val:,}',va='center',fontsize=8)
    fig.suptitle('Diagnósticos por Tipo (Top 10)',fontsize=14,fontweight='bold')
    plt.tight_layout(); plt.savefig(f'{IMG}/{tag}_diag_comp.png',bbox_inches='tight'); plt.close()

# =============================================================================
# GENERAR GRÁFICOS PARA WORD Y EXCEL
# =============================================================================
print("Generando gráficos G3+G6...")
gen_esp(esp_antes, esp_g6, 'g6', 'G3 + G6')
gen_mes(mes_antes, mes_g6, 'g6', 'G3 + G6')
gen_diag(diag_antes, diag_g6, 'g6', 'G3 + G6')
gen_comp_charts(esp_antes, esp_g6, mes_antes, mes_g6, diag_antes, diag_g6, 'g6', 'G3+G6')

print("Generando gráficos G3+G1+G6...")
gen_esp(esp_antes, esp_all, 'all', 'G3 + G1 + G6')
gen_mes(mes_antes, mes_all, 'all', 'G3 + G1 + G6')
gen_diag(diag_antes, diag_all, 'all', 'G3 + G1 + G6')
gen_comp_charts(esp_antes, esp_all, mes_antes, mes_all, diag_antes, diag_all, 'all', 'G3+G1+G6')

# =============================================================================
# DOCUMENTO 1: G3 + G6
# =============================================================================
make_word(
    'Informe_ETL_Grupo6.docx',
    'G3 + G6', 'Grupo 6 → Grupo 3',
    esp_g6, mes_g6, diag_g6,
    (['Aspecto', 'Grupo 3', 'Grupo 6'], [
        ('Motor', 'PostgreSQL 17 (Docker)', 'SQLite'),
        ('Tablas', '8', '9'),
        ('Registros', '50,000', '959,866'),
        ('Normalización', 'Alta (3FN)', 'Media'),
        ('CI / Documento', 'Sí', 'No'),
        ('Dirección / Teléfono', 'Sí', 'Sí (pacientes)'),
        ('Zonas', 'Sí (20)', 'No'),
        ('Recetas', 'Sí', 'Sí (detallado)'),
        ('Historia Clínica', 'No', 'Sí'),
        ('CIE-10', 'No', 'Sí'),
    ]),
    (['Tabla', 'Grupo 3', '+ Grupo 6', 'Total'], [
        ('PERSONA', '5,000', '+303,000', '308,000'),
        ('CITA_MEDICA', '20,000', '+200,000', '220,000'),
        ('DIAGNOSTICO', '15,000', '+85,390', '100,390'),
        ('RECETA', '7,440', '+230,696', '238,136'),
    ]),
    [
        'Base de datos SQLite (hospital.db, 121 MB) con 9 tablas.',
        '300,000 pacientes, 3,000 personal, 200,000 citas, 85,390 diagnósticos, 230,696 recetas.',
        '24 especialidades (vs 15 nuestras). Incluye Historia Clínica y Signos Vitales.',
        'Tipos de diagnóstico: Presuntivo y Confirmado (vs nuestros 25 tipos por método).',
    ],
    'Se compara la BD del Grupo 6 (SQLite, 959,866 registros) con la del Grupo 3 (PostgreSQL, 50,000 registros).',
    [
        '## 2.1 Estrategia de IDs',
        'Pacientes: +300,000 | Personal: +600,000 | Citas: +300,000 | Diagnósticos: +300,000 | Recetas: +300,000',
        '## 2.2 Campos Transformados',
        'PERSONA.Nombre: nombres + apellidos concatenados. PERSONA.CI: marcador G6-PAC/G6-PER. '
        'PERSONA.Sexo: "X" (no disponible). PERSONA.ID_Zona: zona especial ID=98.',
        '## 2.3 Mapeo de Especialidades (24 → 15)',
        '13 coincidencias directas. 11 mapeadas: Medicina Interna, Anestesiología, Radiología, '
        'Hematología, Infectología, Med. Familiar, Med. Preventiva, Med. Emergencias, Patología → Medicina General. '
        'Cirugía General, Reumatología → Traumatología. Nefrología → Urología.',
        '## 2.4 Mapeo de Diagnósticos',
        'Grupo 6 usa tipos Presuntivo y Confirmado. Mapeados a: Presuntivo → Diagnóstico Presuntivo (ID 5), '
        'Confirmado → Diagnóstico Definitivo (ID 6).',
        '## 2.5 Recetas',
        'Grupo 6 tiene recetas detalladas por medicamento. Transformación: medicamento + dosis → Medicamentos, '
        'frecuencia + duración + indicaciones → Indicaciones.',
    ],
    'Medicina General domina con 65,652 citas por el mapeo de múltiples especialidades del G6. '
    'Traumatología sube al 2do lugar (24,485) incluyendo Cirugía General y Reumatología.',
    'El patrón estacional del G3 (pico invernal) se suaviza con los datos uniformes del G6. '
    'Distribución final más equilibrada (~18,000-19,000/mes).',
    'Diagnóstico Definitivo (43,780) y Presuntivo (43,721) del G6 dominan masivamente. '
    'Los tipos por método del G3 mantienen sus valores originales.',
    'g6', sum(mes_g6), sum(v for _,v in diag_g6)
)

# =============================================================================
# DOCUMENTO 2: G3 + G1 + G6
# =============================================================================
make_word(
    'Informe_ETL_Completo.docx',
    'G3 + G1 + G6', 'Grupo 1 y Grupo 6 → Grupo 3',
    esp_all, mes_all, diag_all,
    (['Aspecto', 'Grupo 3', 'Grupo 1', 'Grupo 6'], [
        ('Motor', 'PostgreSQL 17', 'PostgreSQL (Supabase)', 'SQLite'),
        ('Tablas', '8', '4', '9'),
        ('Registros', '50,000', '218,123', '959,866'),
        ('Normalización', 'Alta (3FN)', 'Baja', 'Media'),
        ('CI / Documento', 'Sí', 'No', 'No'),
        ('Dirección', 'Sí', 'No', 'Sí (pacientes)'),
        ('Zonas', 'Sí (20)', 'No', 'No'),
        ('Recetas', 'Sí', 'No', 'Sí'),
        ('CIE-10', 'No', 'Sí', 'Sí'),
    ]),
    (['Tabla', 'Grupo 3', '+ Grupo 1', '+ Grupo 6', 'Total'], [
        ('PERSONA', '5,000', '+100,000', '+303,000', '408,000'),
        ('CITA_MEDICA', '20,000', '+50,000', '+200,000', '270,000'),
        ('DIAGNOSTICO', '15,000', '+68,123', '+85,390', '168,513'),
        ('RECETA', '7,440', '—', '+230,696', '238,136'),
        ('TOTAL', '47,440', '+218,123', '+819,086', '1,084,649'),
    ]),
    [
        'Grupo 1: PostgreSQL en Supabase. 4 tablas, 218,123 registros. Sin normalización de catálogos.',
        'Grupo 6: SQLite (hospital.db). 9 tablas, 959,866 registros. Incluye Hist. Clínica, Recetas.',
    ],
    'Se comparan las bases de datos de los Grupos 1 y 6 con la del Grupo 3, '
    'identificando diferencias estructurales, volumétricas y de calidad para la migración ETL.',
    [
        '## 2.1 Transformaciones Grupo 1',
        'IDs offset +100,000 (pacientes) y +200,000 (personal). '
        '8 especialidades mapeadas directamente. 15 códigos CIE-10 mapeados a tipos diagnósticos. '
        'Campos faltantes (CI, dirección, teléfono, zona) marcados con identificadores de origen.',
        '## 2.2 Transformaciones Grupo 6',
        'IDs offset +300,000 (pacientes/citas/diag/recetas) y +600,000 (personal). '
        '24 especialidades mapeadas a 15 (13 directas + 11 aproximadas). '
        'Nombres concatenados (nombres + apellidos). Recetas detalladas fusionadas.',
        '## 2.3 Resultado Consolidado',
        'Base de datos final: 1,084,649 registros totales. '
        '408,000 personas, 270,000 citas, 168,513 diagnósticos, 238,136 recetas. '
        'Integridad referencial verificada: 0 errores.',
    ],
    'Medicina General lidera con 71,813 citas (35.3x más que el original). '
    'El volumen total pasa de 20,000 a 270,000 citas (incremento de 13.5x). '
    'Las especialidades exclusivas del G3 (Gastroenterología, Neumología) mantienen sus valores.',
    'La distribución mensual se uniformiza (~22,000/mes). El pico invernal del G3 (35%) '
    'se diluye con los datos uniformes de G1 y G6. Variación final: solo 1.1x entre máximo y mínimo.',
    'Definitivo (43,780) y Presuntivo (43,721) del G6 dominan. Clínico (33,940) crece por G1. '
    'Aparecen categorías nuevas significativas: Ambulatorio (8,954) y Endoscópico (4,628).',
    'all', sum(mes_all), sum(v for _,v in diag_all)
)

# =============================================================================
# EXCEL 1: G3 + G6
# =============================================================================
make_excel('graficas_etl_grupo6.xlsx', esp_g6, mes_g6, diag_g6, 'G3+G6', 'g6')

# =============================================================================
# EXCEL 2: G3 + G1 + G6
# =============================================================================
make_excel('graficas_etl_completo.xlsx', esp_all, mes_all, diag_all, 'G3+G1+G6', 'all')

print("\n¡4 documentos generados!")
