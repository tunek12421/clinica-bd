"""
Genera el informe de la Actividad 6:
Preparacion de la herramienta de analisis de datos para graficos de toma de decisiones
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ============================================================
# ESTILOS
# ============================================================
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    h = doc.styles[f'Heading {level}']
    h.font.color.rgb = RGBColor(0, 0, 0)

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            if cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].font.size = Pt(10)
    return table

def add_code(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(18)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)

def add_screenshot_placeholder(label):
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.rows[0].cells[0]
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'[ CAPTURA: {label} ]')
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(128, 128, 128)
    trPr = table.rows[0]._tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), '1800')
    trPr.append(trHeight)
    doc.add_paragraph()

# ============================================================
# PORTADA
# ============================================================
for _ in range(3):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('UNIVERSIDAD PRIVADA DOMINGO SAVIO')
run.bold = True; run.font.size = Pt(14)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('FACULTAD DE INGENIERIA\nCARRERA DE INGENIERIA DE SISTEMAS')
run.font.size = Pt(12)

for _ in range(3):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('PREPARACION DE LA HERRAMIENTA DE ANALISIS\nDE DATOS PARA GRAFICOS DE TOMA DE DECISIONES')
run.bold = True; run.font.size = Pt(16)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Aplicada al Data Warehouse de la Clinica')
run.font.size = Pt(12)

for _ in range(2):
    doc.add_paragraph()

for nombre in [
    'Argote Gonzales Marco Ariel',
    'Benavides Arancibia Jorge Enrique',
    'Lujan Arispe Enrique',
    'Lujan Arispe William',
    'Mejillones Iraizos Rebeca',
    'Valencia Vidaurre Marcos Daniel',
    'Villarroel Espinoza Gustavo Ernesto',
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(nombre)
    run.font.size = Pt(11)

for _ in range(3):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Cochabamba - Bolivia\n2026')
run.bold = True; run.font.size = Pt(12)

doc.add_page_break()

# ============================================================
# 1. INTRODUCCION
# ============================================================
doc.add_heading('1. INTRODUCCION', level=1)
doc.add_paragraph(
    'El presente informe documenta la preparacion de la herramienta de analisis de datos '
    'seleccionada para la visualizacion y toma de decisiones sobre el Data Warehouse (DW) '
    'de la Clinica. El DW integra datos de cuatro fuentes heterogeneas (Grupo 1, Grupo 3, '
    'Grupo 4 y Grupo 6) bajo un esquema estrella con snowflake parcial, consolidando '
    '424,369 atenciones medicas, 426,487 pacientes y 314,668 medicos.'
)
doc.add_paragraph(
    'La herramienta seleccionada permite conectarse directamente al DW PostgreSQL, '
    'construir consultas analiticas comparativas entre sucursales y generar visualizaciones '
    'interactivas que apoyan la toma de decisiones en el ambito clinico y administrativo.'
)

# ============================================================
# 2. JUSTIFICACION DE LA HERRAMIENTA
# ============================================================
doc.add_heading('2. JUSTIFICACION DE LA HERRAMIENTA: METABASE', level=1)

doc.add_heading('2.1. Que es Metabase?', level=2)
doc.add_paragraph(
    'Metabase es una herramienta de inteligencia de negocios (BI) de codigo abierto que '
    'permite conectarse a bases de datos relacionales, ejecutar consultas analiticas y '
    'generar visualizaciones interactivas sin necesidad de conocimientos avanzados de programacion. '
    'Se despliega como contenedor Docker y expone una interfaz web accesible desde el navegador.'
)

doc.add_heading('2.2. Comparacion con alternativas', level=2)
add_table(
    ['Criterio', 'Metabase', 'Power BI', 'Tableau'],
    [
        ['Sistema operativo', 'Linux, Windows, Mac', 'Solo Windows (nativo)', 'Windows, Mac'],
        ['Costo', 'Gratuito (Open Source)', 'Requiere licencia', 'Requiere licencia'],
        ['Despliegue', 'Docker (local)', 'Instalacion nativa', 'Instalacion nativa'],
        ['Conexion PostgreSQL', 'Nativa y directa', 'Requiere conector', 'Requiere conector'],
        ['Curva de aprendizaje', 'Baja', 'Media', 'Alta'],
        ['Dashboards interactivos', 'Si', 'Si', 'Si'],
        ['Consultas SQL directas', 'Si', 'Limitado', 'Si'],
    ]
)

doc.add_paragraph()
doc.add_heading('2.3. Razones de seleccion', level=2)
razones = [
    'Compatibilidad con Linux: el entorno de desarrollo del equipo utiliza Linux, '
    'donde Power BI no esta disponible de forma nativa.',
    'Integracion directa con PostgreSQL: se conecta al contenedor DW (puerto 5434) '
    'sin configuraciones adicionales.',
    'Despliegue en Docker: se levanta con un unico comando y se integra a la misma '
    'red Docker del proyecto (bd-clinica_default).',
    'Gratuito y open source: no requiere licencias pagas.',
    'Soporte para SQL nativo: permite escribir consultas SQL directamente sobre '
    'las tablas del DW (fact_atenciones, dim_medico, dim_paciente, dim_sucursal).',
    'Dashboards interactivos: agrupa multiples graficos en un panel de control unificado.',
]
for r in razones:
    p = doc.add_paragraph(r, style='List Bullet')
    p.runs[0].font.size = Pt(11)

doc.add_paragraph()
add_screenshot_placeholder('Metabase - Vista general de la base de datos DATA WAREHOUSE')

# ============================================================
# 3. MODELADO DE LOS DIAGRAMAS
# ============================================================
doc.add_heading('3. MODELADO DE LOS DIAGRAMAS', level=1)

doc.add_heading('3.1. Esquema del Data Warehouse', level=2)
doc.add_paragraph(
    'El DW utiliza un modelo estrella con snowflake parcial. '
    'Las consultas para los graficos se construyen sobre este esquema:'
)
add_table(
    ['Tabla', 'Tipo', 'Registros', 'Rol en los graficos'],
    [
        ['dim_sucursal', 'Subdimension', '4', 'Segmentacion por grupo (G1, G3, G4, G6)'],
        ['dim_paciente', 'Dimension', '426,487', 'Datos demograficos de pacientes'],
        ['dim_medico', 'Dimension', '314,668', 'Especialidades y distribucion de medicos'],
        ['fact_atenciones', 'Hechos', '424,369', 'Base de todos los indicadores analiticos'],
    ]
)

doc.add_paragraph()
doc.add_heading('3.2. Diagrama del modelo estrella en Metabase', level=2)
doc.add_paragraph(
    'Metabase detecta automaticamente las relaciones entre tablas (foreign keys) '
    'del esquema estrella, permitiendo navegar entre dimensiones y hechos de forma visual:'
)
add_screenshot_placeholder('Diagrama del modelo en Metabase - relaciones entre tablas')

doc.add_heading('3.3. Estructura del Dashboard', level=2)
doc.add_paragraph(
    'El panel de control agrupa las 6 visualizaciones en un dashboard denominado '
    '"Dashboard Clinica - Toma de Decisiones". Cada grafico compara las 4 sucursales '
    'y responde a una pregunta de negocio orientada a la gestion clinica y administrativa:'
)
add_table(
    ['#', 'Nombre del grafico', 'Tipo', 'Tablas involucradas'],
    [
        ['1', 'Carga laboral por sucursal', 'Barras agrupadas', 'fact_atenciones + dim_paciente + dim_sucursal'],
        ['2', 'Atenciones anuales por sucursal', 'Barras agrupadas', 'fact_atenciones + dim_paciente + dim_sucursal'],
        ['3', 'Top especialidades por sucursal', 'Barras apiladas 100%', 'fact_atenciones + dim_medico + dim_paciente + dim_sucursal'],
        ['4', 'Top diagnosticos por sucursal', 'Barras apiladas 100%', 'fact_atenciones + dim_paciente + dim_sucursal'],
        ['5', 'Atenciones por dia de la semana', 'Barras apiladas 100%', 'fact_atenciones + dim_paciente + dim_sucursal'],
        ['6', 'Estado de atenciones por sucursal', 'Barras apiladas 100%', 'fact_atenciones + dim_paciente + dim_sucursal'],
    ]
)
doc.add_paragraph()
add_screenshot_placeholder('Dashboard completo en Metabase - Panel Analitica Clinica')

doc.add_page_break()

# ============================================================
# 4. LAS 6 CONSULTAS
# ============================================================
doc.add_heading('4. CONSULTAS PARA GRAFICOS DE TOMA DE DECISIONES', level=1)

consultas = [
    {
        'num': '4.1',
        'titulo': 'Carga Laboral por Sucursal — Ultima Gestion Completa',
        'tipo': 'Barras agrupadas',
        'dimension': 'dim_sucursal (nombre)',
        'sql': (
            "SELECT ds.nombre AS sucursal,\n"
            "       COUNT(*) AS total_atenciones,\n"
            "       COUNT(DISTINCT fa.medico_key)\n"
            "           AS medicos_activos,\n"
            "       ROUND(COUNT(*)::numeric /\n"
            "             NULLIF(COUNT(DISTINCT fa.medico_key), 0), 1)\n"
            "             AS carga_por_medico\n"
            "FROM fact_atenciones fa\n"
            "JOIN dim_paciente dp ON dp.paciente_key = fa.paciente_key\n"
            "JOIN dim_sucursal ds ON ds.sucursal_key = dp.sucursal_key\n"
            "WHERE fa.anio = (SELECT MAX(anio) - 1\n"
            "                 FROM fact_atenciones)\n"
            "GROUP BY ds.nombre\n"
            "ORDER BY carga_por_medico DESC;"
        ),
        'decision': (
            'Compara la carga de trabajo entre sucursales durante la ultima gestion completa, '
            'midiendo atenciones totales, medicos activos y el ratio atenciones/medico. '
            'Una sucursal con ratio alto indica sobrecarga del personal medico, lo que '
            'justifica contrataciones o redistribucion de recursos humanos. Permite a la '
            'gerencia priorizar inversiones en las sucursales con mayor presion operativa.'
        ),
    },
    {
        'num': '4.2',
        'titulo': 'Atenciones Anuales por Sucursal — Gestiones Completas',
        'tipo': 'Barras agrupadas',
        'dimension': 'fact_atenciones (anio) + dim_sucursal',
        'sql': (
            "SELECT fa.anio::text AS año,\n"
            "       ds.nombre AS sucursal,\n"
            "       COUNT(*) AS total_atenciones\n"
            "FROM fact_atenciones fa\n"
            "JOIN dim_paciente dp ON dp.paciente_key = fa.paciente_key\n"
            "JOIN dim_sucursal ds ON ds.sucursal_key = dp.sucursal_key\n"
            "WHERE fa.anio < (SELECT MAX(anio)\n"
            "                 FROM fact_atenciones)\n"
            "GROUP BY fa.anio, ds.nombre\n"
            "ORDER BY fa.anio, ds.nombre;"
        ),
        'decision': (
            'Revela la evolucion historica del volumen de atenciones por sucursal, '
            'excluyendo la gestion en curso (datos parciales) para una comparacion justa. '
            'Permite identificar cuando cada grupo inicio operaciones, detectar tendencias '
            'de crecimiento o decrecimiento, y comparar la escala operativa entre sedes a '
            'lo largo del tiempo.'
        ),
    },
    {
        'num': '4.3',
        'titulo': 'Top Especialidades por Sucursal — Ultima Gestion Completa',
        'tipo': 'Barras apiladas 100%',
        'dimension': 'dim_medico (especialidad) + dim_sucursal',
        'sql': (
            "SELECT dm.especialidad,\n"
            "       ds.nombre AS sucursal,\n"
            "       COUNT(*) AS total_atenciones\n"
            "FROM fact_atenciones fa\n"
            "JOIN dim_medico dm ON dm.medico_key = fa.medico_key\n"
            "JOIN dim_paciente dp ON dp.paciente_key = fa.paciente_key\n"
            "JOIN dim_sucursal ds ON ds.sucursal_key = dp.sucursal_key\n"
            "WHERE fa.anio = (SELECT MAX(anio) - 1\n"
            "                 FROM fact_atenciones)\n"
            "GROUP BY dm.especialidad, ds.nombre\n"
            "ORDER BY total_atenciones DESC\n"
            "LIMIT 20;"
        ),
        'decision': (
            'Muestra la distribucion porcentual de las especialidades mas demandadas por '
            'sucursal durante la ultima gestion completa. El apilado al 100% normaliza las '
            'diferencias de volumen, permitiendo comparar que especialidades concentran mayor '
            'carga asistencial en cada grupo. Facilita la planificacion de contrataciones y '
            'la redistribucion de especialistas entre sedes.'
        ),
    },
    {
        'num': '4.4',
        'titulo': 'Top Diagnosticos por Sucursal — Ultima Gestion Completa',
        'tipo': 'Barras apiladas 100%',
        'dimension': 'fact_atenciones (tipo_diagnostico) + dim_sucursal',
        'sql': (
            "SELECT fa.tipo_diagnostico,\n"
            "       ds.nombre AS sucursal,\n"
            "       COUNT(*) AS total\n"
            "FROM fact_atenciones fa\n"
            "JOIN dim_paciente dp ON dp.paciente_key = fa.paciente_key\n"
            "JOIN dim_sucursal ds ON ds.sucursal_key = dp.sucursal_key\n"
            "WHERE fa.anio = (SELECT MAX(anio) - 1\n"
            "                 FROM fact_atenciones)\n"
            "GROUP BY fa.tipo_diagnostico, ds.nombre\n"
            "ORDER BY total DESC\n"
            "LIMIT 20;"
        ),
        'decision': (
            'Compara la distribucion porcentual de tipos de diagnostico entre sucursales '
            'durante la ultima gestion completa. El apilado al 100% normaliza las diferencias '
            'de volumen entre grupos, permitiendo comparar el perfil diagnostico de cada sede. '
            'Si una sucursal presenta alta concentracion de un tipo especifico, la administracion '
            'puede asignar recursos especializados de forma dirigida.'
        ),
    },
    {
        'num': '4.5',
        'titulo': 'Atenciones por Dia de la Semana por Sucursal — Ultima Gestion Completa',
        'tipo': 'Barras apiladas 100%',
        'dimension': 'fact_atenciones (dia_semana) + dim_sucursal',
        'sql': (
            "SELECT CASE fa.dia_semana\n"
            "           WHEN 'Monday' THEN 'Lunes'\n"
            "           WHEN 'Tuesday' THEN 'Martes'\n"
            "           WHEN 'Wednesday' THEN 'Miercoles'\n"
            "           WHEN 'Thursday' THEN 'Jueves'\n"
            "           WHEN 'Friday' THEN 'Viernes'\n"
            "           WHEN 'Saturday' THEN 'Sabado'\n"
            "           WHEN 'Sunday' THEN 'Domingo'\n"
            "       END AS dia_semana,\n"
            "       ds.nombre AS sucursal,\n"
            "       COUNT(*) AS total_atenciones\n"
            "FROM fact_atenciones fa\n"
            "JOIN dim_paciente dp ON dp.paciente_key = fa.paciente_key\n"
            "JOIN dim_sucursal ds ON ds.sucursal_key = dp.sucursal_key\n"
            "WHERE fa.anio = (SELECT MAX(anio) - 1\n"
            "                 FROM fact_atenciones)\n"
            "GROUP BY fa.dia_semana, ds.nombre\n"
            "ORDER BY CASE fa.dia_semana\n"
            "    WHEN 'Monday' THEN 1 WHEN 'Tuesday' THEN 2\n"
            "    WHEN 'Wednesday' THEN 3 WHEN 'Thursday' THEN 4\n"
            "    WHEN 'Friday' THEN 5 WHEN 'Saturday' THEN 6\n"
            "    WHEN 'Sunday' THEN 7\n"
            "END;"
        ),
        'decision': (
            'Muestra la distribucion porcentual de atenciones por dia de la semana y sucursal '
            'durante la ultima gestion completa. El apilado al 100% permite comparar la '
            'participacion relativa de cada grupo sin que las diferencias de volumen distorsionen '
            'la visualizacion. Permite optimizar la asignacion de turnos medicos diferenciada '
            'por sede y dia de la semana.'
        ),
    },
    {
        'num': '4.6',
        'titulo': 'Estado de Atenciones por Sucursal — Ultima Gestion',
        'tipo': 'Barras apiladas 100%',
        'dimension': 'fact_atenciones (estado) + dim_sucursal',
        'sql': (
            "SELECT ds.nombre AS sucursal,\n"
            "       fa.estado,\n"
            "       COUNT(*) AS total\n"
            "FROM fact_atenciones fa\n"
            "JOIN dim_paciente dp ON dp.paciente_key = fa.paciente_key\n"
            "JOIN dim_sucursal ds ON ds.sucursal_key = dp.sucursal_key\n"
            "WHERE fa.anio = (SELECT MAX(anio)\n"
            "                 FROM fact_atenciones)\n"
            "GROUP BY ds.nombre, fa.estado\n"
            "ORDER BY ds.nombre, total DESC;"
        ),
        'decision': (
            'Los 10 estados representan el ciclo de vida de una atencion medica: '
            'Pendiente > Confirmada > En consulta > Atendida > Finalizado > Alta/Control/Completada, '
            'con Cancelada y No asistio como estados de abandono. '
            'Comparar la distribucion porcentual entre sucursales permite identificar: '
            'sucursales con cuello de botella (muchas Pendientes), perdida de productividad '
            '(muchas Canceladas o No asistio), y diferencias en el uso de estados entre sedes. '
            'El hallazgo principal es la falta de estandarizacion: cada sucursal usa los estados '
            'de forma diferente, lo que lleva a la decision de unificar criterios operativos.'
        ),
    },
]

for c in consultas:
    doc.add_heading(f'{c["num"]}. {c["titulo"]}', level=2)
    add_table(
        ['Tipo de grafico', 'Dimension principal'],
        [[c['tipo'], c['dimension']]]
    )
    doc.add_paragraph()
    doc.add_paragraph('Consulta SQL:')
    add_code(c['sql'])
    doc.add_paragraph()
    p = doc.add_paragraph()
    rb = p.add_run('Toma de decision: ')
    rb.bold = True
    rb.font.size = Pt(11)
    p.add_run(c['decision']).font.size = Pt(11)
    doc.add_paragraph()
    add_screenshot_placeholder(f'Grafico {c["num"][-1]}: {c["titulo"]} en Metabase')
    doc.add_paragraph()

doc.add_page_break()

# ============================================================
# 5. CONCLUSIONES
# ============================================================
doc.add_heading('5. CONCLUSIONES', level=1)

conclusiones = [
    'Metabase demostro ser una alternativa viable y eficiente a Power BI en entornos Linux, '
    'permitiendo conectarse directamente al DW PostgreSQL mediante Docker sin configuraciones '
    'adicionales ni costos de licencia.',
    'Las 6 consultas comparativas cubren las dimensiones clave del DW (carga laboral, evolucion '
    'historica, especialidades saturadas, perfil diagnostico, distribucion semanal y estado '
    'operativo), todas segmentadas por sucursal para facilitar la comparacion entre los 4 grupos.',
    'El uso de filtros temporales dinamicos (ultima gestion completa, gestiones completas, '
    'ultima gestion) garantiza que las visualizaciones se actualicen automaticamente conforme '
    'se incorporan nuevos datos al DW, sin necesidad de modificar las consultas.',
    'La normalizacion mediante barras apiladas al 100% permite comparar proporciones entre '
    'sucursales con volumenes de datos muy diferentes (ej. Grupo 4 con datos desde 2015 vs. '
    'Grupo 3 con datos desde 2024), eliminando el sesgo por escala.',
    'El dashboard centraliza los indicadores comparativos mas relevantes para la toma de '
    'decisiones, permitiendo a la gerencia identificar sucursales con sobrecarga, '
    'especialidades saturadas, diferencias en el perfil diagnostico y estados operativos '
    'que requieren atencion inmediata.',
]

for texto in conclusiones:
    p = doc.add_paragraph(texto, style='List Bullet')
    p.runs[0].font.size = Pt(11)

# ============================================================
# 6. REFERENCIAS
# ============================================================
doc.add_heading('6. REFERENCIAS', level=1)
refs = [
    'Metabase. (2024). Metabase Open Source Documentation. https://www.metabase.com/docs/latest/',
    'PostgreSQL Global Development Group. (2024). PostgreSQL 17 Documentation. https://www.postgresql.org/docs/17/',
    'Kimball, R., & Ross, M. (2013). The Data Warehouse Toolkit (3rd ed.). Wiley.',
    'Cepymenews. (2024). Importancia de la inteligencia de negocios, el market research y el CX. https://cepymenews.es/',
]
for ref in refs:
    p = doc.add_paragraph(ref, style='List Bullet')
    p.runs[0].font.size = Pt(10)

# ============================================================
# GUARDAR
# ============================================================
out = '/home/tunek/Universidad/MATERIAS/bd-clinica/Informe_Actividad6_Metabase.docx'
doc.save(out)
print(f'Documento generado: {out}')
