"""
Genera el informe de la Actividad 7:
Informe del Panel de control para la toma de decisiones
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
import os

CAPTURAS_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'capturas_metabase')
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ============================================================
# ESTILOS
# ============================================================
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    h = doc.styles[f'Heading {level}']
    h.font.color.rgb = RGBColor(0, 0, 0)


def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            if cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].font.size = Pt(10)
    return table


def add_image(filename, width=Inches(6.0)):
    img_path = os.path.join(CAPTURAS_DIR, filename)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(img_path, width=width)
    doc.add_paragraph()


# ============================================================
# PORTADA
# ============================================================
for _ in range(3):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('UNIVERSIDAD PRIVADA DOMINGO SAVIO')
run.bold = True
run.font.size = Pt(14)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('FACULTAD DE INGENIERIA\nCARRERA DE INGENIERIA DE SISTEMAS')
run.font.size = Pt(12)

for _ in range(3):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(
    'INFORME DEL PANEL DE CONTROL\nPARA LA TOMA DE DECISIONES')
run.bold = True
run.font.size = Pt(16)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Aplicado al Data Warehouse de la Clinica')
run.font.size = Pt(12)

for _ in range(2):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Lujan Arispe Enrique')
run.font.size = Pt(12)
run.bold = True

for _ in range(3):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Cochabamba - Bolivia\n2026')
run.bold = True
run.font.size = Pt(12)

doc.add_page_break()

# ============================================================
# 1. OBJETIVOS
# ============================================================
doc.add_heading('1. OBJETIVOS', level=1)

doc.add_heading('1.1. Objetivo general', level=2)
doc.add_paragraph(
    'Disenar e implementar un panel de control (dashboard) basado en inteligencia de '
    'negocios que permita la toma de decisiones informadas sobre la gestion clinica '
    'y administrativa de las sucursales de la Clinica, utilizando los datos consolidados '
    'en el Data Warehouse.'
)

doc.add_heading('1.2. Objetivos especificos', level=2)
objetivos = [
    'Integrar datos de cuatro fuentes heterogeneas (Grupo 1, Grupo 3, Grupo 4 y Grupo 6) '
    'en un modelo de Data Warehouse bajo un esquema hibrido estrella con snowflake parcial.',
    'Identificar indicadores clave de desempeno (KPIs) relevantes para la gestion clinica: '
    'carga laboral, evolucion historica, especialidades saturadas, perfil diagnostico, '
    'distribucion semanal y estado operativo.',
    'Construir consultas analiticas SQL con filtros temporales dinamicos que permitan '
    'comparar el rendimiento entre las 4 sucursales de forma objetiva.',
    'Generar visualizaciones interactivas mediante Metabase que faciliten la interpretacion '
    'de los datos por parte de la gerencia y el personal administrativo.',
    'Elaborar un reporte de decisiones basado en los hallazgos del panel de control, '
    'con recomendaciones accionables para la mejora operativa.',
]
for obj in objetivos:
    p = doc.add_paragraph(obj, style='List Bullet')
    p.runs[0].font.size = Pt(11)

# ============================================================
# 2. ALCANCE
# ============================================================
doc.add_heading('2. ALCANCE', level=1)

doc.add_heading('2.1. Alcance del sistema', level=2)
doc.add_paragraph(
    'El panel de control abarca la totalidad de los datos clinicos integrados en el '
    'Data Warehouse, provenientes de cuatro sucursales con diferentes origenes:'
)
add_table(
    ['Sucursal', 'Fuente', 'Periodo', 'Registros integrados'],
    [
        ['Grupo 1', 'Supabase (AWS)', '2021 - 2025', '68,123'],
        ['Grupo 3', 'PostgreSQL local', '2024 - 2026', '15,000'],
        ['Grupo 4', 'Neon (Azure)', '2015 - 2026', '301,112'],
        ['Grupo 6', 'Dump PostgreSQL', '2024 - 2026', '40,134'],
        ['Total', '', '', '424,369'],
    ]
)
doc.add_paragraph()
doc.add_paragraph(
    'En total, el DW consolida 424,369 atenciones, 426,487 pacientes y 314,668 medicos '
    'bajo un esquema hibrido estrella con snowflake parcial, donde dim_sucursal actua '
    'como subdimension de dim_paciente, permitiendo segmentar todos los indicadores '
    'por grupo de origen.'
)

doc.add_heading('2.2. Alcance del analisis', level=2)
doc.add_paragraph(
    'El analisis se centra en seis dimensiones clave para la toma de decisiones:'
)
dimensiones = [
    'Carga laboral: ratio de atenciones por medico en cada sucursal.',
    'Evolucion historica: volumen de atenciones anuales por sucursal a lo largo del tiempo.',
    'Especialidades saturadas: combinaciones sucursal-especialidad con mayor presion asistencial.',
    'Perfil diagnostico: distribucion de tipos de diagnostico por sucursal.',
    'Distribucion semanal: patrones de demanda por dia de la semana.',
    'Estado operativo: proporcion de atenciones completadas, pendientes, canceladas, etc.',
]
for d in dimensiones:
    p = doc.add_paragraph(d, style='List Bullet')
    p.runs[0].font.size = Pt(11)

doc.add_paragraph()
doc.add_paragraph(
    'Las consultas utilizan filtros temporales dinamicos (ultima gestion completa, '
    'gestiones completas) que se actualizan automaticamente conforme se incorporan '
    'nuevos datos al DW, garantizando la vigencia del panel de control.'
)

doc.add_heading('2.3. Limitaciones', level=2)
limitaciones = [
    'Los datos de cada grupo provienen de periodos diferentes, lo que limita la '
    'comparacion directa en terminos absolutos. Se mitiga este sesgo mediante '
    'visualizaciones apiladas al 100% que comparan proporciones.',
    'La calidad y completitud de los datos varia entre grupos, ya que cada uno '
    'registro la informacion de forma independiente con diferentes criterios.',
    'El panel de control no incluye datos financieros ni de costos, por lo que '
    'las recomendaciones se limitan al ambito operativo y asistencial.',
]
for lim in limitaciones:
    p = doc.add_paragraph(lim, style='List Bullet')
    p.runs[0].font.size = Pt(11)

doc.add_page_break()

# ============================================================
# 3. DESARROLLO
# ============================================================
doc.add_heading('3. DESARROLLO', level=1)

# --- 3.1 Fuentes ---
doc.add_heading('3.1. Fuentes de extraccion de la informacion', level=2)
doc.add_paragraph(
    'Los datos del Data Warehouse provienen de cuatro fuentes heterogeneas, '
    'cada una correspondiente a un grupo que implemento de forma '
    'independiente el sistema operacional de la clinica. La integracion '
    'al DW fue realizada individualmente como parte del Grupo 3:'
)
add_table(
    ['Fuente', 'Tecnologia', 'Conexion', 'Tipo de extraccion'],
    [
        ['Grupo 1', 'Supabase (PostgreSQL en AWS)', 'Remota via pooler', 'ETL directo con psql COPY'],
        ['Grupo 3', 'PostgreSQL 17 local', 'localhost:5432', 'ETL directo con psql COPY'],
        ['Grupo 4', 'Neon (PostgreSQL en Azure)', 'Remota via pooler', 'ETL directo con psql COPY'],
        ['Grupo 6', 'Dump SQL (PostgreSQL 17)', 'Importacion local', 'Restauracion de dump + ETL'],
    ]
)
doc.add_paragraph()
doc.add_paragraph(
    'Todas las fuentes comparten el mismo esquema relacional operacional '
    '(PERSONA, CITA_MEDICA, DIAGNOSTICO, TIPO_DIAGNOSTICO, ESPECIALIDAD, ZONA), '
    'lo que permite aplicar un proceso ETL unificado.'
)

# --- 3.2 Transformacion ---
doc.add_heading('3.2. Transformacion de la informacion (ETL)', level=2)
doc.add_paragraph(
    'El proceso ETL (Extract, Transform, Load) transforma los datos operacionales '
    'de cuatro fuentes heterogeneas en un modelo dimensional optimizado para el analisis. '
    'A continuacion se detalla el proceso de extraccion, analisis comparativo y '
    'transformacion de cada grupo.'
)

# --- 3.2.1 Modelo destino ---
doc.add_heading('3.2.1. Modelo Data Warehouse (destino)', level=3)
doc.add_paragraph(
    'El Data Warehouse utiliza un modelo estrella con snowflake parcial '
    '(dim_sucursal como subdimension de dim_paciente). Cada fila de '
    'fact_atenciones representa un evento de diagnostico dentro de una '
    'cita medica, vinculado a un paciente y un medico.'
)
doc.add_paragraph(
    'dim_sucursal -- dim_paciente -- fact_atenciones -- dim_medico'
)
add_table(
    ['Tabla', 'Tipo', 'Columnas', 'Descripcion'],
    [
        ['dim_sucursal', 'Subdimension', '3', 'Fuente/sucursal de los datos (G3, G1, G4, G6)'],
        ['dim_paciente', 'Dimension', '11', 'Datos del paciente + FK a dim_sucursal'],
        ['dim_medico', 'Dimension', '9', 'Datos del medico con especialidad'],
        ['fact_atenciones', 'Hechos', '17', 'Diagnostico + cita + dimensiones temporales'],
    ]
)
doc.add_paragraph()

# dim_sucursal
doc.add_paragraph('Contenido de dim_sucursal:')
add_table(
    ['sucursal_key', 'nombre', 'host'],
    [
        ['1', 'Grupo 3', 'localhost:5433 (PostgreSQL - clinica_db)'],
        ['2', 'Grupo 1', 'aws-0-us-west-2.pooler.supabase.com (Supabase)'],
        ['3', 'Grupo 6', 'PostgreSQL 17 (dump hospital_db)'],
        ['4', 'Grupo 4', 'ep-xxx.us-east-2.aws.neon.tech (Neon)'],
    ]
)
doc.add_paragraph()

# Tipo de conexion
doc.add_paragraph('Tipo de conexion a las fuentes de extraccion:')
add_table(
    ['Fuente', 'Tipo de conexion', 'Metodo'],
    [
        ['Grupo 3 (mi grupo)', 'Directa (localhost)', 'Consulta SQL local al contenedor Docker'],
        ['Grupo 1', 'Descarga CSV', 'Conexion psql a Supabase -> \\copy TO STDOUT -> CSV'],
        ['Grupo 4', 'Descarga CSV', 'Conexion psql a Neon -> \\copy TO STDOUT -> CSV'],
        ['Grupo 6', 'Backup (dump SQL)', 'Archivo dump PostgreSQL -> conversion encoding -> staging'],
    ]
)
doc.add_paragraph()

# Datos base G3
doc.add_paragraph(
    'El punto de partida es el DW cargado unicamente con datos del Grupo 3 '
    '(mi modelo original):'
)
add_table(
    ['Tabla DW', 'Registros (solo G3)'],
    [
        ['dim_sucursal', '4 (pre-cargadas)'],
        ['dim_paciente', '4,500'],
        ['dim_medico', '500'],
        ['fact_atenciones', '15,000'],
    ]
)
doc.add_paragraph()

# --- 3.2.2 ETL Grupo 3 ---
doc.add_heading('3.2.2. Extraccion: Grupo 3 (Directa)', level=3)
doc.add_paragraph(
    'Los datos del Grupo 3 residen en la misma base de datos destino (clinica_db). '
    'La carga al DW se realiza mediante consultas COPY directas:'
)
p = doc.add_paragraph()
run = p.add_run(
    '-- Extraccion directa desde el contenedor \'db\'\n'
    'docker compose exec -T db psql -U clinica_user -d clinica_db -c "\n'
    '    COPY (\n'
    '        SELECT p.CI, p.Nombre, p.Fecha_Nacimiento, ...\n'
    '        FROM PERSONA p JOIN ZONA z ON p.ID_Zona = z.ID_Zona\n'
    '        WHERE p.Matricula IS NULL\n'
    '    ) TO STDOUT WITH CSV"'
)
run.font.size = Pt(9)
run.font.name = 'Courier New'

# --- 3.2.3 ETL Grupo 1 ---
doc.add_heading('3.2.3. Analisis comparativo: Grupo 1 -> Grupo 3', level=3)
doc.add_paragraph(
    'Base de datos alojada en Supabase (AWS us-west-2). Esquema plano de 4 tablas '
    'sin normalizacion de catalogos. Se extraen todas las columnas disponibles:'
)
add_table(
    ['Tabla', 'Registros', 'Columnas (todas)'],
    [
        ['pacientes', '50,000', 'paciente_id, nombre, fecha_nacimiento, genero'],
        ['personal', '50,000', 'personal_id, nombre, cargo, especialidad'],
        ['atenciones', '50,000', 'atencion_id, paciente_id, personal_id, fecha_atencion, estado, motivo_consulta'],
        ['diagnosticos', '68,123', 'diagnostico_id, atencion_id, codigo_cie10, descripcion, severidad'],
    ]
)
doc.add_paragraph()

doc.add_paragraph(
    'El Grupo 1 tiene menos columnas que el esquema destino (Grupo 3). '
    'El ETL compensa los campos faltantes generando o derivando valores:'
)
add_table(
    ['Campo destino (G3)', 'Disponible en G1', 'Transformacion aplicada'],
    [
        ['PERSONA.CI', 'No', 'Se genera marcador sintetico: G1-PAC-{id}, G1-PER-{id}'],
        ['PERSONA.Direccion, Telefono', 'No', 'Valor por defecto: "Sin dato (Grupo 1)"'],
        ['PERSONA.ID_Zona', 'No', 'Zona especial ID=99 ("Sin zona asignada")'],
        ['PERSONA.Sexo (personal)', 'No', 'Marcador "X" (no disponible)'],
        ['PERSONA.Fecha_Nac (personal)', 'No', 'Marcador 1900-01-01'],
        ['PERSONA.Matricula', 'No', 'Solo doctores: G1-MAT-{id}'],
        ['PERSONA.ID_Especialidad', 'Texto libre', 'JOIN por nombre con tabla ESPECIALIDAD (8 de 15 coinciden)'],
        ['CITA.Fecha_Cita + Hora', 'timestamp unico', 'Se extrae DATE y TIME del timestamp'],
        ['CITA.Numero_Turno', 'No', 'ROW_NUMBER() OVER (PARTITION BY dia, medico)'],
        ['DIAGNOSTICO.ID_Tipo_Diag', 'Codigo CIE-10', 'Mapeo CIE-10 -> categoria (tabla auxiliar)'],
        ['DIAGNOSTICO.Observaciones', 'severidad', 'Concatenado: "CIE-10: {cod} | Severidad: {sev}"'],
        ['RECETA', 'No existe', 'No se migran recetas (tabla no existe en G1)'],
    ]
)
doc.add_paragraph()

doc.add_paragraph('Estrategia de IDs (offsets para evitar colision):')
add_table(
    ['Entidad', 'Rango original', 'Offset', 'Rango destino'],
    [
        ['Pacientes -> PERSONA', '1 - 50,000', '+100,000', '100,001 - 150,000'],
        ['Personal -> PERSONA', '1 - 50,000', '+200,000', '200,001 - 250,000'],
        ['Atenciones -> CITA_MEDICA', '1 - 50,000', '+100,000', '100,001 - 150,000'],
        ['Diagnosticos -> DIAGNOSTICO', '1 - 68,123', '+100,000', '100,001 - 168,123'],
    ]
)
doc.add_paragraph()

doc.add_paragraph('Mapeo CIE-10 -> Tipo de Diagnostico:')
add_table(
    ['Codigo CIE-10', 'Diagnostico', 'Categoria asignada'],
    [
        ['I10', 'Hipertension esencial', 'Clinico'],
        ['R10.9 / R51', 'Dolor abdominal / Cefalea', 'Clinico'],
        ['J02.9 / J06.9 / J20.9', 'Infecciones respiratorias', 'Clinico'],
        ['M54.5', 'Dolor lumbar bajo', 'Por Imagen'],
        ['E78.5 / E11.9', 'Hiperlipidemia / Diabetes tipo 2', 'Laboratorio'],
        ['N39.0', 'Infeccion urinaria', 'Laboratorio'],
        ['E66.9', 'Obesidad', 'Nutricional'],
        ['K29.5', 'Gastritis cronica', 'Endoscopico'],
        ['Z00.0 / Z71.1', 'Examen general / Consulta orientacion', 'Ambulatorio'],
    ]
)
doc.add_paragraph()

doc.add_paragraph('Resultado de la carga G3 + G1:')
add_table(
    ['Tabla DW', 'Antes (G3)', '+ Grupo 1', 'Despues (Total)'],
    [
        ['dim_paciente', '4,500', '+87,487', '91,987'],
        ['dim_medico', '500', '+13,337', '13,837'],
        ['fact_atenciones', '15,000', '+68,123', '83,123'],
    ]
)
doc.add_paragraph()

# --- 3.2.4 ETL Grupo 6 ---
doc.add_heading('3.2.4. Analisis comparativo: Grupo 6 -> Grupo 3', level=3)
doc.add_paragraph(
    'Esquema PostgreSQL 17 normalizado en 3NF. 25 tablas totales (incluye '
    'catalogos y extensiones), de las cuales 9 son relevantes para la migracion. '
    'A diferencia del G1, el G6 tiene un esquema mas completo y normalizado:'
)
add_table(
    ['Tabla', 'Registros', 'Descripcion'],
    [
        ['persona', '35,331', 'Datos basicos: ci, nombre, fecha_nacimiento, sexo, direccion, telefono'],
        ['paciente', '34,500', 'Tabla de rol: id_paciente -> id_persona'],
        ['personal', '831', 'Tabla de rol: id_personal -> id_persona, id_especialidad'],
        ['especialidad', '24', 'Catalogo: id_especialidad, nombre (9 mas que G3)'],
        ['zona', '30', 'Catalogo: id_zona, nombre, ciudad (10 mas que G3)'],
        ['tipo_diagnostico', '25', 'Catalogo: IDs 1-25, coinciden con G3'],
        ['cita_medica', '70,000', 'id_cita, fecha_registro, fecha_cita, hora, estado, numero_turno'],
        ['diagnostico', '40,134', 'id_diagnostico, descripcion, observaciones, id_tipo_diagnostico'],
        ['receta', '40,134', 'id_receta, medicamentos, indicaciones'],
    ]
)
doc.add_paragraph()

doc.add_paragraph(
    'El G6 comparte la mayoria de campos con G3, pero las diferencias estructurales '
    'requieren transformaciones especificas:'
)
add_table(
    ['Aspecto', 'G3', 'G6', 'Transformacion'],
    [
        ['Personas', 'PERSONA unificada', 'persona + paciente + personal', 'JOIN de tablas de rol para reconstruir PERSONA'],
        ['Especialidades', '15 registros', '24 registros', 'Remapeo de 9 adicionales por afinidad clinica'],
        ['Zonas', '20 zonas', '30 zonas', 'Zona generica ID=98 para las 10 sin equivalente'],
        ['Encoding', 'UTF-8', 'WIN1252', 'Conversion sed antes de importar el dump'],
        ['Catalogos extra', 'No existen', '9 tablas cat_*', 'Se descartan (no relevantes para el DW)'],
    ]
)
doc.add_paragraph()

doc.add_paragraph('Estrategia de IDs:')
add_table(
    ['Entidad', 'Rango original', 'Offset', 'Rango destino'],
    [
        ['Pacientes (persona)', '1 - 34,500', '+300,000', '300,001 - 334,500'],
        ['Personal (persona)', '1 - 831', '+600,000', '600,001 - 600,831'],
        ['Citas', '1 - 70,000', '+300,000', '300,001 - 370,000'],
        ['Diagnosticos', '1 - 40,134', '+300,000', '300,001 - 340,134'],
        ['Recetas', '1 - 40,134', '+300,000', '300,001 - 340,134'],
    ]
)
doc.add_paragraph()

doc.add_paragraph(
    'Extraccion del dump SQL: El G6 entrega sus datos como dump PostgreSQL '
    '(encoding WIN1252). Se convierte el encoding a UTF-8 y se renombran las '
    'tablas a staging mediante sed.'
)

doc.add_paragraph('Mapeo de especialidades (24 -> 15):')
doc.add_paragraph(
    'Las especialidades 1-15 coinciden con G3 (mapeo directo). Las 9 adicionales '
    '(16-24) se remapearon por afinidad clinica:'
)
add_table(
    ['ID G6', 'Especialidad G6', 'ID G3', 'Especialidad G3', 'Justificacion'],
    [
        ['16', 'Medicina Interna', '1', 'Medicina General', 'Subespecialidad de medicina general'],
        ['17', 'Cirugia General', '7', 'Traumatologia', 'Area quirurgica mas cercana'],
        ['18', 'Radiologia', '1', 'Medicina General', 'Servicio de apoyo diagnostico'],
        ['19', 'Anestesiologia', '1', 'Medicina General', 'Servicio transversal'],
        ['20', 'Nefrologia', '10', 'Urologia', 'Misma area anatomica (sistema renal)'],
        ['21', 'Hematologia', '1', 'Medicina General', 'Subespecialidad de medicina interna'],
        ['22', 'Reumatologia', '7', 'Traumatologia', 'Sistema musculoesqueletico'],
        ['23', 'Infectologia', '1', 'Medicina General', 'Subespecialidad de medicina interna'],
        ['24', 'Medicina Familiar', '1', 'Medicina General', 'Equivalente funcional'],
    ]
)
doc.add_paragraph()

doc.add_paragraph('Resultado de la carga G3 + G6:')
add_table(
    ['Tabla DW', 'Antes (G3)', '+ Grupo 6', 'Despues (Total)'],
    [
        ['dim_paciente', '4,500', '+34,500', '39,000'],
        ['dim_medico', '500', '+831', '1,331'],
        ['fact_atenciones', '15,000', '+40,134', '55,134'],
    ]
)
doc.add_paragraph()

# --- 3.2.5 ETL Grupo 4 ---
doc.add_heading('3.2.5. Analisis comparativo: Grupo 4 -> Grupo 3', level=3)
doc.add_paragraph(
    'Base de datos alojada en Neon (PostgreSQL serverless en AWS us-east-2). '
    'El Grupo 4 es el que aporta el mayor volumen de datos al DW, con registros '
    'desde 2015 hasta 2026.'
)

doc.add_paragraph(
    'El Grupo 4 comparte el mismo esquema relacional que el Grupo 3 '
    '(PERSONA, CITA_MEDICA, DIAGNOSTICO, TIPO_DIAGNOSTICO, ESPECIALIDAD, ZONA), '
    'por lo que la extraccion es directa sin transformaciones de estructura:'
)
add_table(
    ['Tabla', 'Registros', 'Compatibilidad con G3'],
    [
        ['PERSONA', '~330,000', 'Mismas columnas, mismo esquema'],
        ['CITA_MEDICA', '~500,000', 'Mismas columnas, mismo esquema'],
        ['DIAGNOSTICO', '~301,112', 'Mismas columnas, mismo esquema'],
        ['TIPO_DIAGNOSTICO', '25', 'IDs 1-25 coinciden con G3'],
        ['ESPECIALIDAD', '15', 'IDs 1-15 coinciden con G3'],
        ['ZONA', '20+', 'Estructura compatible'],
    ]
)
doc.add_paragraph()

doc.add_paragraph(
    'Al ser esquemas identicos, la unica diferencia es el volumen y periodo de datos. '
    'No se requieren transformaciones de estructura, solo offsets de IDs:'
)
add_table(
    ['Aspecto', 'G3', 'G4', 'Implicacion'],
    [
        ['Volumen', '~15,000 atenciones', '~301,112 atenciones', '20x mayor, es la fuente principal del DW'],
        ['Periodo', '2024 - 2026', '2015 - 2026', 'Historico mas extenso (10 anos)'],
        ['Conexion', 'localhost', 'Neon pooler remoto', 'Requiere SSL (sslmode=require)'],
        ['Transformacion', '-', 'Solo offsets de IDs', 'Sin mapeo de especialidades ni diagnosticos'],
    ]
)
doc.add_paragraph()

doc.add_paragraph(
    'El proceso de extraccion del Grupo 4 es similar al del Grupo 1: se conecta '
    'al servidor Neon mediante psql con cadena de conexion PostgreSQL (sslmode=require) '
    'y se descargan las tablas relevantes a CSV mediante \\copy TO STDOUT.'
)

doc.add_paragraph('Estrategia de IDs:')
add_table(
    ['Entidad', 'Rango original', 'Offset', 'Rango destino'],
    [
        ['Pacientes -> PERSONA', '1 - ~300,000', '+400,000', '400,001 - ~700,000'],
        ['Personal -> PERSONA', '1 - ~30,000', '+700,000', '700,001 - ~730,000'],
        ['Citas -> CITA_MEDICA', '1 - ~500,000', '+400,000', '400,001 - ~900,000'],
        ['Diagnosticos -> DIAGNOSTICO', '1 - ~301,112', '+400,000', '400,001 - ~701,112'],
    ]
)
doc.add_paragraph()

doc.add_paragraph(
    'Dado que el Grupo 4 comparte el mismo esquema relacional que el Grupo 3 '
    '(PERSONA, CITA_MEDICA, DIAGNOSTICO con las mismas tablas de catalogo), '
    'la transformacion es directa: solo requiere la aplicacion de offsets en los IDs '
    'y la asignacion del marcador \'G4\' en el campo CI para identificar el origen. '
    'No fue necesario mapeo de especialidades ni conversion de codigos diagnosticos.'
)

doc.add_paragraph('Resultado de la carga con Grupo 4:')
add_table(
    ['Tabla DW', 'Antes (G3+G1+G6)', '+ Grupo 4', 'Despues (Total)'],
    [
        ['dim_paciente', '126,487', '+300,000', '426,487'],
        ['dim_medico', '14,668', '+300,000', '314,668'],
        ['fact_atenciones', '123,257', '+301,112', '424,369'],
    ]
)
doc.add_paragraph()

# --- 3.2.6 Carga al DW ---
doc.add_heading('3.2.6. Carga final al Data Warehouse', level=3)
doc.add_paragraph(
    'La carga al modelo dimensional (estrella con snowflake parcial) se ejecuta en 6 pasos para cada grupo:'
)

doc.add_paragraph('Paso 1 - Extraccion de dim_paciente (clasificacion por grupo_origen):')
p = doc.add_paragraph()
run = p.add_run(
    'SELECT p.CI, p.Nombre, p.Fecha_Nacimiento, p.Sexo,\n'
    '       p.Direccion, p.Telefono, z.Nombre, z.Ciudad,\n'
    '       CASE\n'
    '         WHEN p.CI LIKE \'G1-%\' THEN \'G1\'\n'
    '         WHEN p.CI LIKE \'G6-%\' THEN \'G6\'\n'
    '         WHEN p.CI LIKE \'G4-%\' THEN \'G4\'\n'
    '         ELSE \'G3\'\n'
    '       END AS grupo_origen\n'
    'FROM PERSONA p\n'
    'JOIN ZONA z ON z.ID_Zona = p.ID_Zona\n'
    'WHERE p.Matricula IS NULL;'
)
run.font.size = Pt(9)
run.font.name = 'Courier New'

doc.add_paragraph('Paso 2 - Extraccion de dim_medico: SELECT de medicos con especialidad y ubicacion.')

doc.add_paragraph('Paso 3 - Asignacion de sucursal_key post-carga:')
p = doc.add_paragraph()
run = p.add_run(
    'UPDATE dim_paciente SET sucursal_key = CASE grupo_origen\n'
    '  WHEN \'G3\' THEN 1\n'
    '  WHEN \'G1\' THEN 2\n'
    '  WHEN \'G6\' THEN 3\n'
    '  WHEN \'G4\' THEN 4\n'
    'END;'
)
run.font.size = Pt(9)
run.font.name = 'Courier New'

doc.add_paragraph('Paso 4 - Carga de staging de hechos (combina DIAGNOSTICO + CITA_MEDICA + TIPO_DIAGNOSTICO):')
p = doc.add_paragraph()
run = p.add_run(
    'SELECT pac.CI, med.CI, c.Fecha_Cita,\n'
    '       EXTRACT(YEAR FROM c.Fecha_Cita)::INT AS anio,\n'
    '       EXTRACT(MONTH FROM c.Fecha_Cita)::INT AS mes,\n'
    '       EXTRACT(QUARTER FROM c.Fecha_Cita)::INT AS trimestre,\n'
    '       TRIM(TO_CHAR(c.Fecha_Cita, \'Day\')) AS dia_semana,\n'
    '       c.Hora, c.Estado, c.Numero_Turno,\n'
    '       d.Descripcion, d.Observaciones,\n'
    '       d.Tipo_Procedimiento, td.Nombre, td.Categoria,\n'
    '       grupo_origen\n'
    'FROM DIAGNOSTICO d\n'
    'JOIN CITA_MEDICA c ON c.ID_Cita = d.ID_Cita\n'
    'JOIN PERSONA pac ON pac.ID_Persona = c.ID_Paciente\n'
    'JOIN PERSONA med ON med.ID_Persona = c.ID_Medico\n'
    'JOIN TIPO_DIAGNOSTICO td ON td.ID_Tipo_Diagnostico = d.ID_Tipo_Diagnostico;'
)
run.font.size = Pt(9)
run.font.name = 'Courier New'

doc.add_paragraph('Paso 5 - Resolucion de FK en el DW mediante JOIN por CI:')
p = doc.add_paragraph()
run = p.add_run(
    'INSERT INTO fact_atenciones (...)\n'
    'SELECT dp.paciente_key, dm.medico_key,\n'
    '       sf.fecha_cita, sf.anio, sf.mes, ...\n'
    'FROM stg_fact sf\n'
    'JOIN dim_paciente dp ON dp.ci = sf.paciente_ci\n'
    'JOIN dim_medico dm ON dm.ci = sf.medico_ci;'
)
run.font.size = Pt(9)
run.font.name = 'Courier New'

doc.add_paragraph(
    'Paso 6 - Verificacion de integridad: se valida que no existan registros '
    'huerfanos ni FK rotas en el modelo final.'
)

doc.add_paragraph()
doc.add_paragraph('Verificacion de integridad del DW final:')
add_table(
    ['Verificacion', 'Resultado'],
    [
        ['Pacientes sin sucursal asignada', '0'],
        ['Atenciones sin paciente valido (FK)', '0'],
        ['Atenciones sin medico valido (FK)', '0'],
        ['Registros huerfanos en fact_atenciones', '0'],
    ]
)
doc.add_paragraph()

# Resultado final consolidado
doc.add_paragraph('Resultado final del DW con los 4 grupos:')
add_table(
    ['Tabla DW', 'G3', 'G1', 'G6', 'G4', 'Total'],
    [
        ['dim_sucursal', '-', '-', '-', '-', '4'],
        ['dim_paciente', '4,500', '87,487', '34,500', '300,000', '426,487'],
        ['dim_medico', '500', '13,337', '831', '300,000', '314,668'],
        ['fact_atenciones', '15,000', '68,123', '40,134', '301,112', '424,369'],
    ]
)
doc.add_paragraph()

doc.add_paragraph('Distribucion porcentual de atenciones (hechos) por grupo:')
add_table(
    ['Grupo', 'Atenciones', '%'],
    [
        ['G3', '15,000', '3.5%'],
        ['G1', '68,123', '16.1%'],
        ['G6', '40,134', '9.5%'],
        ['G4', '301,112', '70.9%'],
        ['Total', '424,369', '100%'],
    ]
)
doc.add_paragraph()

doc.add_paragraph(
    'Cada fuente presento desafios distintos de integracion: el Grupo 1 requirio '
    'mapeo de codigos CIE-10 a categorias diagnosticas y generacion de identificadores '
    'sinteticos; el Grupo 6 necesito conversion de encoding (WIN1252 -> UTF-8) y '
    'remapeo de 24 especialidades a 15; el Grupo 4, al compartir el mismo esquema, '
    'solo requirio offsets de IDs. Los tres tipos de conexion utilizados (directa local, '
    'descarga CSV desde servidor remoto y carga de backup SQL) demuestran la capacidad '
    'del proceso ETL para integrar fuentes con distintos mecanismos de acceso.'
)

# --- 3.3 Metodos y modelos ---
doc.add_heading('3.3. Metodos y modelos usados', level=2)

doc.add_heading('Modelo dimensional: Esquema hibrido estrella-snowflake', level=3)
doc.add_paragraph(
    'Se utiliza el modelo estrella (Star Schema) propuesto por Kimball, donde la tabla '
    'de hechos fact_atenciones se conecta a tres dimensiones (dim_paciente, dim_medico, '
    'dim_sucursal). La variante snowflake parcial surge porque dim_paciente contiene '
    'una referencia a dim_sucursal, permitiendo segmentar los datos por grupo de origen '
    'sin duplicar informacion.'
)

doc.add_heading('Metodo de analisis: ROLAP', level=3)
doc.add_paragraph(
    'Las consultas del panel de control operan bajo el paradigma ROLAP (Relational OLAP), '
    'donde el analisis multidimensional se realiza directamente mediante consultas SQL '
    'sobre las tablas relacionales del DW. Este enfoque ofrece flexibilidad para crear '
    'nuevas consultas sin necesidad de precalcular cubos, y aprovecha la capacidad de '
    'PostgreSQL para ejecutar agregaciones, JOINs y subconsultas de forma eficiente.'
)

doc.add_heading('Herramienta de visualizacion: Metabase', level=3)
doc.add_paragraph(
    'Metabase es una plataforma de inteligencia de negocios (BI) de codigo abierto que '
    'permite conectarse a bases de datos PostgreSQL, ejecutar consultas nativas SQL y '
    'generar visualizaciones interactivas. Se despliega como contenedor Docker dentro '
    'de la misma red del proyecto, conectandose directamente al DW sin configuraciones '
    'adicionales.'
)

doc.add_heading('Tecnicas de visualizacion aplicadas', level=3)
tecnicas = [
    'Barras agrupadas: para comparar valores absolutos entre sucursales '
    '(carga laboral, evolucion historica).',
    'Barras apiladas al 100%: para comparar distribuciones porcentuales entre '
    'sucursales con volumenes muy diferentes, eliminando el sesgo por escala '
    '(diagnosticos, dias de semana, estados).',
    'Barras apiladas 100% con Top N: para mostrar las especialidades mas demandadas '
    'por sucursal, limitando a las principales para evitar ruido visual.',
    'Filtros temporales dinamicos: subconsultas que seleccionan automaticamente '
    'la ultima gestion completa o las gestiones sin datos parciales.',
]
for t in tecnicas:
    p = doc.add_paragraph(t, style='List Bullet')
    p.runs[0].font.size = Pt(11)

doc.add_page_break()

# --- 3.4 Reporte de toma de decisiones ---
doc.add_heading('3.4. Reporte de la toma de decisiones', level=2)
doc.add_paragraph(
    'A continuacion se presentan los hallazgos obtenidos del panel de control '
    'y las decisiones que cada uno sustenta. Los datos corresponden a la gestion '
    '2025 (ultima gestion completa disponible en el DW).'
)

# Hallazgo 1: Carga laboral
doc.add_heading('Hallazgo 1: Desbalance critico en la carga laboral', level=3)
add_table(
    ['Sucursal', 'Atenciones', 'Medicos activos', 'Carga/medico'],
    [
        ['Grupo 6', '19,370', '831', '23.3'],
        ['Grupo 3', '6,845', '500', '13.7'],
        ['Grupo 1', '5,085', '1,032', '4.9'],
        ['Grupo 4', '24,773', '23,747', '1.0'],
    ]
)
doc.add_paragraph()
p = doc.add_paragraph()
rb = p.add_run('Analisis: ')
rb.bold = True
rb.font.size = Pt(11)
p.add_run(
    'Grupo 6 presenta la mayor sobrecarga con 23.3 atenciones por medico, '
    'casi 5 veces mas que Grupo 1 (4.9) y 23 veces mas que Grupo 4 (1.0). '
    'Esto indica que Grupo 6 opera con un numero insuficiente de medicos '
    'para su volumen de demanda, mientras que Grupo 4 tiene una dotacion '
    'de personal significativamente mayor.'
).font.size = Pt(11)

p = doc.add_paragraph()
rb = p.add_run('Decision: ')
rb.bold = True
rb.font.size = Pt(11)
p.add_run(
    'Reasignar medicos desde Grupo 4 (con excedente) hacia Grupo 6 y Grupo 3. '
    'Alternativamente, contratar personal adicional para Grupo 6 y evaluar la '
    'redistribucion de pacientes entre sedes para equilibrar la carga.'
).font.size = Pt(11)

doc.add_paragraph()
add_image('image6.png')

# Hallazgo 2: Evolucion historica
doc.add_heading('Hallazgo 2: Incorporacion progresiva de sucursales', level=3)
doc.add_paragraph(
    'El analisis historico revela que Grupo 4 opera desde 2015 con un volumen '
    'constante de ~25,000 atenciones anuales. Grupo 1 se incorporo en 2021 con '
    '~16,000 atenciones/ano. Grupo 3 y Grupo 6 son los mas recientes, iniciando '
    'en 2024. Esta diferencia en antiguedad explica parcialmente las diferencias '
    'de volumen total entre sucursales.'
)
p = doc.add_paragraph()
rb = p.add_run('Decision: ')
rb.bold = True
rb.font.size = Pt(11)
p.add_run(
    'Las comparaciones absolutas entre sucursales deben considerar el periodo de '
    'operacion. Para decisiones de inversion, priorizar las sucursales mas nuevas '
    '(Grupo 3 y Grupo 6) que estan en fase de crecimiento y necesitan recursos '
    'para consolidarse.'
).font.size = Pt(11)

doc.add_paragraph()
add_image('image7.png')

# Hallazgo 3: Especialidades mas demandadas
doc.add_heading('Hallazgo 3: Concentracion de demanda por especialidad', level=3)
doc.add_paragraph(
    'El analisis porcentual de las especialidades mas demandadas durante la ultima '
    'gestion completa revela que las 4 sucursales comparten las mismas especialidades '
    'principales (Medicina General, Cardiologia, Pediatria, Ginecologia, Neurologia), '
    'pero con proporciones diferentes. Grupo 4 presenta la distribucion mas equilibrada, '
    'mientras que Grupo 6 concentra mayor demanda en pocas especialidades.'
)
p = doc.add_paragraph()
rb = p.add_run('Decision: ')
rb.bold = True
rb.font.size = Pt(11)
p.add_run(
    'Reforzar las especialidades con mayor demanda en cada sucursal. Evaluar la '
    'posibilidad de derivar pacientes hacia sedes con menor carga en especialidades '
    'saturadas. Planificar contrataciones enfocadas en las especialidades que cada '
    'sede necesita segun su perfil de demanda.'
).font.size = Pt(11)

doc.add_paragraph()
add_image('image8.png')

# Hallazgo 4: Perfil diagnostico
doc.add_heading('Hallazgo 4: Diferencias en el perfil diagnostico', level=3)
doc.add_paragraph(
    'El analisis porcentual de tipos de diagnostico revela que los tres primeros '
    'tipos (Clinico, Laboratorio e Imagen) concentran las atenciones de todos los '
    'grupos, pero con proporciones diferentes. Grupo 4 domina en diagnosticos '
    'especializados (Diferencial, Presuntivo, Definitivo, Prenatal, Patologico, '
    'Molecular, Genetico, Funcional) que los demas grupos practicamente no registran.'
)
p = doc.add_paragraph()
rb = p.add_run('Decision: ')
rb.bold = True
rb.font.size = Pt(11)
p.add_run(
    'Estandarizar la clasificacion de diagnosticos entre sucursales para permitir '
    'una comparacion mas precisa. Evaluar si Grupo 1, 3 y 6 necesitan capacitacion '
    'o equipamiento para ampliar su oferta diagnostica a los tipos especializados '
    'que actualmente solo ofrece Grupo 4.'
).font.size = Pt(11)

doc.add_paragraph()
add_image('image9.png')

# Hallazgo 5: Distribucion semanal
doc.add_heading('Hallazgo 5: Distribucion uniforme de la demanda semanal', level=3)
doc.add_paragraph(
    'El analisis por dia de la semana muestra una distribucion practicamente uniforme '
    'en todas las sucursales: cada dia concentra entre 13% y 15% del total semanal. '
    'No se observan picos significativos en ningun dia especifico, lo que sugiere '
    'que la demanda es constante a lo largo de la semana.'
)
p = doc.add_paragraph()
rb = p.add_run('Decision: ')
rb.bold = True
rb.font.size = Pt(11)
p.add_run(
    'Mantener una dotacion de personal uniforme durante todos los dias de la semana. '
    'No se justifica reducir turnos en fines de semana ni reforzar dias especificos, '
    'dado que la demanda es homogenea.'
).font.size = Pt(11)

doc.add_paragraph()
add_image('image10.png')

# Hallazgo 6: Estado operativo
doc.add_heading('Hallazgo 6: Diferencias en la gestion operativa', level=3)
doc.add_paragraph(
    'La distribucion de estados de atencion varia significativamente entre sucursales '
    'en la ultima gestion:'
)
estados = [
    'Grupo 4 presenta una distribucion equilibrada entre Finalizado (25%), '
    'Pendiente (25%), Alta (25%) y Control (25%), indicando un flujo operativo '
    'completo y sistematico.',
    'Grupo 6 concentra el 99% de sus atenciones en estado "Atendida", '
    'sugiriendo que utiliza un sistema de estados simplificado.',
    'Grupo 3 opera con estados "Confirmada" y "Completada" principalmente, '
    'con un volumen significativamente menor en la gestion actual.',
]
for est in estados:
    p = doc.add_paragraph(est, style='List Bullet')
    p.runs[0].font.size = Pt(11)

p = doc.add_paragraph()
rb = p.add_run('Decision: ')
rb.bold = True
rb.font.size = Pt(11)
p.add_run(
    'Estandarizar los estados de atencion entre todas las sucursales, adoptando '
    'el modelo de Grupo 4 (Pendiente -> Atendida -> Finalizado -> Alta/Control) '
    'como referencia. Capacitar al personal de Grupo 6 y Grupo 3 en el uso '
    'correcto de los estados para mejorar la trazabilidad del ciclo de atencion.'
).font.size = Pt(11)

doc.add_paragraph()
add_image('image11.png')

doc.add_paragraph()
add_image('image5.png')

doc.add_page_break()

# ============================================================
# 4. CONCLUSIONES
# ============================================================
doc.add_heading('4. CONCLUSIONES', level=1)

conclusiones = [
    'La inteligencia de negocios aplicada al Data Warehouse de la Clinica permitio '
    'identificar hallazgos concretos y accionables que no serian visibles con reportes '
    'operacionales tradicionales. La consolidacion de datos de 4 fuentes heterogeneas '
    'en un modelo dimensional unificado habilita la comparacion objetiva entre sucursales.',

    'El hallazgo mas critico es el desbalance en la carga laboral: Grupo 6 tiene una '
    'carga 23 veces mayor que Grupo 4 por medico, lo que representa un riesgo operativo '
    'inmediato que requiere intervencion. Las especialidades mas afectadas son Oncologia '
    '(71.4 atenciones/medico) y Urologia (63.0 atenciones/medico).',

    'Las diferencias en el perfil diagnostico y los estados de atencion revelan una falta '
    'de estandarizacion entre sucursales. Grupo 4 ofrece diagnosticos especializados que '
    'los demas grupos no registran, y cada sucursal utiliza los estados de atencion de '
    'forma diferente, dificultando la comparacion directa.',

    'El uso de filtros temporales dinamicos y visualizaciones apiladas al 100% resuelve '
    'el desafio de comparar sucursales con volumenes y periodos de datos muy diferentes, '
    'permitiendo un analisis proporcional justo.',

    'El panel de control implementado en Metabase cumple con los principios de la '
    'inteligencia de negocios: transforma datos crudos en informacion estructurada, '
    'y esta informacion en conocimiento accionable para la toma de decisiones '
    'gerenciales y administrativas.',
]

for texto in conclusiones:
    p = doc.add_paragraph(texto, style='List Bullet')
    p.runs[0].font.size = Pt(11)

# ============================================================
# 5. RECOMENDACIONES
# ============================================================
doc.add_heading('5. RECOMENDACIONES', level=1)

recomendaciones = [
    ('Redistribucion de personal medico',
     'Implementar un plan de redistribucion de medicos entre sucursales, priorizando '
     'Grupo 6 (carga: 23.3) y Grupo 3 (carga: 13.7). Evaluar la posibilidad de '
     'reasignar medicos desde Grupo 4 (carga: 1.0), que cuenta con excedente de personal.'),

    ('Refuerzo de especialidades criticas',
     'Contratar oncologos, urologos y neurologos adicionales para las sucursales con '
     'mayor saturacion. Oncologia en Grupo 6 (71.4 atenciones/medico) requiere atencion '
     'inmediata.'),

    ('Estandarizacion de procesos',
     'Unificar los criterios de clasificacion de diagnosticos y estados de atencion '
     'entre todas las sucursales, utilizando el modelo de Grupo 4 como referencia. '
     'Esto mejorara la comparabilidad de los datos y la trazabilidad del ciclo de atencion.'),

    ('Actualizacion periodica del DW',
     'Establecer un proceso de actualizacion periodica (mensual o trimestral) del '
     'Data Warehouse para mantener vigente el panel de control. Los filtros temporales '
     'dinamicos garantizan que las visualizaciones se adapten automaticamente.'),

    ('Ampliacion del panel de control',
     'Incorporar indicadores financieros (costo por atencion, facturacion por sucursal) '
     'y de satisfaccion del paciente en futuras versiones del dashboard, para complementar '
     'el analisis operativo con la perspectiva economica y de calidad de servicio.'),

    ('Capacitacion del personal',
     'Capacitar al personal administrativo en el uso de Metabase para que puedan consultar '
     'el panel de control de forma autonoma, reduciendo la dependencia del equipo tecnico '
     'para la generacion de reportes.'),
]

for titulo, texto in recomendaciones:
    p = doc.add_paragraph()
    rb = p.add_run(f'{titulo}: ')
    rb.bold = True
    rb.font.size = Pt(11)
    p.add_run(texto).font.size = Pt(11)

# ============================================================
# 6. REFERENCIAS
# ============================================================
doc.add_heading('6. REFERENCIAS', level=1)
refs = [
    'Cepymenews. (2024). Importancia de la inteligencia de negocios, el market '
    'research y el CX. https://cepymenews.es/',
    'Kimball, R., & Ross, M. (2013). The Data Warehouse Toolkit: The Definitive '
    'Guide to Dimensional Modeling (3rd ed.). Wiley.',
    'Metabase. (2024). Metabase Open Source Documentation. '
    'https://www.metabase.com/docs/latest/',
    'PostgreSQL Global Development Group. (2024). PostgreSQL 17 Documentation. '
    'https://www.postgresql.org/docs/17/',
    'Laudon, K. C., & Laudon, J. P. (2020). Management Information Systems: '
    'Managing the Digital Firm (16th ed.). Pearson.',
]
for ref in refs:
    p = doc.add_paragraph(ref, style='List Bullet')
    p.runs[0].font.size = Pt(10)

# ============================================================
# GUARDAR
# ============================================================
out = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'Informe_Actividad7_PanelControl.docx')
doc.save(out)
print(f'Documento generado: {out}')
