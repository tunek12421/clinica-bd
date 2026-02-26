"""
Genera el informe técnico de migración ETL al Data Warehouse (modelo estrella).
Flujo: G3 solo → G3+G1 → rollback → G3+G6 → rollback → G3+G1+G6 (final)

Uso: python3 generar-informe-dw.py
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# ============================================================
# ESTILOS
# ============================================================
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    h = doc.styles[f'Heading {level}']
    h.font.color.rgb = RGBColor(0, 0, 0)

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            if cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].font.size = Pt(10)
    return table

def add_code(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Pt(18)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(8.5)

# ============================================================
# PORTADA
# ============================================================
for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('UNIVERSIDAD PRIVADA DOMINGO SAVIO')
run.bold = True; run.font.size = Pt(14)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('FACULTAD DE INGENIERÍA\nCARRERA DE INGENIERÍA DE SISTEMAS')
run.font.size = Pt(12)

for _ in range(3):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('INFORME DE RECONOCIMIENTO Y\nMIGRACIÓN ETL AL DATA WAREHOUSE')
run.bold = True; run.font.size = Pt(16)

for _ in range(2):
    doc.add_paragraph()

for nombre in [
    'Argote Gonzales Marco Ariel',
    'Benavides Arancibia Jorge Enrique',
    'Lujan Arispe Enrique',
    'Lujan Arispe William',
    'Mejillones Iraizos Rebeca',
    'Valencia Vidaurre Marcos Daniel',
    'Villarroel Espinoza Gustavo Ernesto',
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(nombre)
    run.bold = True; run.font.size = Pt(11)

for _ in range(3):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Cochabamba - Bolivia\n2026')
run.bold = True; run.font.size = Pt(12)

doc.add_page_break()

# ============================================================
# 1. MODELO DW (DESTINO)
# ============================================================
doc.add_heading('1. MODELO DATA WAREHOUSE (DESTINO)', level=1)

doc.add_heading('1.1. Esquema Estrella', level=2)
doc.add_paragraph(
    'El Data Warehouse utiliza un modelo estrella con snowflake parcial '
    '(dim_sucursal como subdimensión de dim_paciente). '
    'Cada fila de fact_atenciones representa un evento de diagnóstico dentro de una cita médica.'
)
add_code('dim_sucursal ── dim_paciente ── fact_atenciones ── dim_medico')
doc.add_paragraph()
add_table(
    ['Tabla', 'Tipo', 'Columnas', 'Descripción'],
    [
        ['dim_sucursal', 'Subdimensión', '3', 'Fuente/sucursal de los datos (G3, G1, G6)'],
        ['dim_paciente', 'Dimensión', '11', 'Datos del paciente + FK a dim_sucursal'],
        ['dim_medico', 'Dimensión', '9', 'Datos del médico con especialidad'],
        ['fact_atenciones', 'Hechos', '17', 'Diagnóstico + cita + dimensiones temporales'],
    ]
)

doc.add_heading('1.2. Definición del esquema (SQL)', level=2)
doc.add_paragraph('Creación de las tablas del DW:')
add_code(
    'CREATE TABLE dim_sucursal (\n'
    '    sucursal_key  SERIAL PRIMARY KEY,\n'
    '    nombre        VARCHAR(100) NOT NULL,\n'
    '    host          VARCHAR(200) NOT NULL\n'
    ');\n\n'
    'CREATE TABLE dim_paciente (\n'
    '    paciente_key      SERIAL PRIMARY KEY,\n'
    '    sucursal_key      INT REFERENCES dim_sucursal(sucursal_key),\n'
    '    ci                VARCHAR(20) NOT NULL,\n'
    '    nombre            VARCHAR(150) NOT NULL,\n'
    '    fecha_nacimiento  DATE,\n'
    '    sexo              CHAR(1),\n'
    '    direccion         VARCHAR(255),\n'
    '    telefono          VARCHAR(20),\n'
    '    zona              VARCHAR(100),\n'
    '    ciudad            VARCHAR(100),\n'
    '    grupo_origen      VARCHAR(10) NOT NULL\n'
    ');\n\n'
    'CREATE TABLE dim_medico (\n'
    '    medico_key    SERIAL PRIMARY KEY,\n'
    '    ci            VARCHAR(20) NOT NULL,\n'
    '    nombre        VARCHAR(150) NOT NULL,\n'
    '    matricula     VARCHAR(50),\n'
    '    sexo          CHAR(1),\n'
    '    especialidad  VARCHAR(100),\n'
    '    zona          VARCHAR(100),\n'
    '    ciudad        VARCHAR(100),\n'
    '    grupo_origen  VARCHAR(10) NOT NULL\n'
    ');\n\n'
    'CREATE TABLE fact_atenciones (\n'
    '    atencion_key        SERIAL PRIMARY KEY,\n'
    '    paciente_key        INT NOT NULL REFERENCES dim_paciente(paciente_key),\n'
    '    medico_key          INT NOT NULL REFERENCES dim_medico(medico_key),\n'
    '    fecha_cita          DATE NOT NULL,\n'
    '    anio                INT NOT NULL,\n'
    '    mes                 INT NOT NULL,\n'
    '    trimestre           INT NOT NULL,\n'
    '    dia_semana          VARCHAR(15) NOT NULL,\n'
    '    hora                TIME,\n'
    '    estado              VARCHAR(50),\n'
    '    numero_turno        INT,\n'
    '    descripcion         TEXT,\n'
    '    observaciones       TEXT,\n'
    '    tipo_procedimiento  VARCHAR(100),\n'
    '    tipo_diagnostico    VARCHAR(100),\n'
    '    categoria           VARCHAR(100),\n'
    '    grupo_origen        VARCHAR(10) NOT NULL\n'
    ');'
)

doc.add_heading('1.3. dim_sucursal', level=2)
add_table(
    ['sucursal_key', 'nombre', 'host'],
    [
        ['1', 'Grupo 3', 'localhost:5433 (PostgreSQL - clinica_db)'],
        ['2', 'Grupo 1', 'aws-0-us-west-2.pooler.supabase.com (Supabase)'],
        ['3', 'Grupo 6', 'PostgreSQL 17 (dump hospital_db)'],
    ]
)

doc.add_heading('1.4. Datos base: Solo Grupo 3', level=2)
doc.add_paragraph(
    'El punto de partida es el DW cargado únicamente con datos del Grupo 3 '
    '(nuestro modelo original).'
)
add_table(
    ['Tabla DW', 'Registros (solo G3)'],
    [
        ['dim_sucursal', '3 (pre-cargadas)'],
        ['dim_paciente', '4,500'],
        ['dim_medico', '500'],
        ['fact_atenciones', '15,000'],
    ]
)

doc.add_page_break()

# ============================================================
# 2. GRUPO 1 → G3
# ============================================================
doc.add_heading('2. ANÁLISIS COMPARATIVO GRUPO 1 → GRUPO 3', level=1)

doc.add_heading('2.1. Esquema del Grupo 1 (Fuente)', level=2)
doc.add_paragraph(
    'Base de datos alojada en Supabase (AWS us-west-2). '
    '4 tablas planas sin normalización de catálogos.'
)
add_table(
    ['Tabla', 'Registros', 'Columnas principales'],
    [
        ['pacientes', '50,000', 'paciente_id, nombre, fecha_nacimiento, genero'],
        ['personal', '50,000', 'personal_id, nombre, cargo, especialidad'],
        ['atenciones', '50,000', 'atencion_id, paciente_id, personal_id, fecha_atencion, estado, motivo'],
        ['diagnosticos', '68,123', 'diagnostico_id, atencion_id, codigo_cie10, descripcion, severidad'],
    ]
)

doc.add_heading('2.2. Diferencias identificadas', level=2)
add_table(
    ['Aspecto', 'Grupo 3', 'Grupo 1', 'Observación'],
    [
        ['Identificación', 'CI obligatorio', 'Sin CI', 'Se genera marcador G1-PAC-{id}'],
        ['Ubicación', 'zona + ciudad', 'No existe', 'Se asigna zona especial ID=99'],
        ['Especialidad', 'ID normalizado (15)', 'Texto libre', 'JOIN por nombre con ESPECIALIDAD'],
        ['Diagnóstico', 'FK → TIPO_DIAGNOSTICO', 'Código CIE-10', 'Mapeo CIE-10 → categoría'],
        ['Fecha/hora', 'fecha_cita + hora', 'timestamp único', 'Se extrae DATE y TIME'],
        ['Turno', 'numero_turno', 'No existe', 'ROW_NUMBER() por día+médico'],
        ['Recetas', 'Tabla RECETA', 'No tiene', 'No se migran recetas'],
    ]
)

doc.add_heading('2.3. Información a transformar', level=2)

doc.add_heading('2.3.1. Estrategia de IDs', level=3)
add_table(
    ['Entidad', 'Rango original', 'Offset', 'Rango destino'],
    [
        ['Pacientes → PERSONA', '1 - 50,000', '+100,000', '100,001 - 150,000'],
        ['Personal → PERSONA', '1 - 50,000', '+200,000', '200,001 - 250,000'],
        ['Atenciones → CITA_MEDICA', '1 - 50,000', '+100,000', '100,001 - 150,000'],
        ['Diagnósticos → DIAGNOSTICO', '1 - 68,123', '+100,000', '100,001 - 168,123'],
    ]
)

doc.add_heading('2.3.2. Transformación de pacientes', level=3)
doc.add_paragraph('Extracción y transformación de pacientes del G1:')
add_code(
    "SELECT\n"
    "    paciente_id + 100000           AS id_persona,\n"
    "    'G1-PAC-' || paciente_id       AS ci,\n"
    "    nombre                          AS nombre,\n"
    "    fecha_nacimiento                AS fecha_nacimiento,\n"
    "    genero                          AS sexo,\n"
    "    'Sin dato (Grupo 1)'           AS direccion,\n"
    "    'Sin dato (Grupo 1)'           AS telefono,\n"
    "    NULL                            AS matricula,\n"
    "    99                              AS id_zona,\n"
    "    NULL                            AS id_especialidad\n"
    "INTO tfm_g1_persona_pacientes\n"
    "FROM stg_g1_pacientes;"
)

doc.add_heading('2.3.3. Transformación de personal', level=3)
add_code(
    "SELECT\n"
    "    personal_id + 200000           AS id_persona,\n"
    "    'G1-PER-' || personal_id       AS ci,\n"
    "    nombre, '1900-01-01' AS fecha_nacimiento,\n"
    "    'X' AS sexo,\n"
    "    'G1-MAT-' || personal_id       AS matricula,\n"
    "    99                              AS id_zona,\n"
    "    e.ID_Especialidad              AS id_especialidad\n"
    "INTO tfm_g1_persona_personal\n"
    "FROM stg_g1_personal p\n"
    "LEFT JOIN ESPECIALIDAD e ON e.Nombre = p.especialidad;"
)

doc.add_heading('2.3.4. Transformación de citas', level=3)
add_code(
    "SELECT\n"
    "    atencion_id + 100000                 AS id_cita,\n"
    "    fecha_atencion::DATE                 AS fecha_registro,\n"
    "    fecha_atencion::DATE                 AS fecha_cita,\n"
    "    fecha_atencion::TIME                 AS hora,\n"
    "    ROW_NUMBER() OVER(\n"
    "        PARTITION BY fecha_atencion::DATE, personal_id\n"
    "        ORDER BY fecha_atencion\n"
    "    )                                    AS numero_turno,\n"
    "    estado,\n"
    "    paciente_id + 100000                 AS id_paciente,\n"
    "    personal_id + 200000                 AS id_medico\n"
    "INTO tfm_g1_cita_medica\n"
    "FROM stg_g1_atenciones;"
)

doc.add_heading('2.3.5. Mapeo CIE-10 → Tipo de Diagnóstico', level=3)
doc.add_paragraph(
    'Los códigos CIE-10 del G1 se mapearon a categorías diagnósticas del G3:'
)
add_table(
    ['Código CIE-10', 'Diagnóstico', 'Categoría asignada'],
    [
        ['I10', 'Hipertensión esencial', 'Clínico'],
        ['R10.9', 'Dolor abdominal', 'Clínico'],
        ['R51', 'Cefalea', 'Clínico'],
        ['J02.9/J06.9/J20.9', 'Infecciones respiratorias', 'Clínico'],
        ['L30.9', 'Dermatitis', 'Clínico'],
        ['M54.5', 'Dolor lumbar bajo', 'Por Imagen'],
        ['E78.5', 'Hiperlipidemia', 'Laboratorio'],
        ['E11.9', 'Diabetes mellitus tipo 2', 'Laboratorio'],
        ['N39.0', 'Infección urinaria', 'Laboratorio'],
        ['E66.9', 'Obesidad', 'Nutricional'],
        ['K29.5', 'Gastritis crónica', 'Endoscópico'],
        ['Z00.0', 'Examen general', 'Ambulatorio'],
        ['Z71.1', 'Consulta orientación', 'Ambulatorio'],
    ]
)

doc.add_heading('2.4. Resultado de la carga: G3 + G1', level=2)
add_table(
    ['Tabla DW', 'Antes (G3)', '+ Grupo 1', 'Después (Total)'],
    [
        ['dim_paciente', '4,500', '+87,487', '91,987'],
        ['dim_medico', '500', '+13,337', '13,837'],
        ['fact_atenciones', '15,000', '+68,123', '83,123'],
    ]
)

doc.add_heading('2.5. Reversibilidad', level=2)
doc.add_paragraph(
    'El rollback elimina datos por rango de IDs y restaura secuencias:'
)
add_code(
    "-- Eliminar en orden de dependencias FK\n"
    "DELETE FROM DIAGNOSTICO WHERE ID_Diagnostico BETWEEN 100001 AND 168123;\n"
    "DELETE FROM CITA_MEDICA WHERE ID_Cita BETWEEN 100001 AND 150000;\n"
    "DELETE FROM PERSONA WHERE ID_Persona BETWEEN 100001 AND 250000;\n\n"
    "-- Restaurar secuencias\n"
    "SELECT setval('persona_id_persona_seq', (SELECT MAX(ID_Persona) FROM PERSONA));"
)

doc.add_page_break()

# ============================================================
# 3. GRUPO 6 → G3
# ============================================================
doc.add_heading('3. ANÁLISIS COMPARATIVO GRUPO 6 → GRUPO 3', level=1)

doc.add_heading('3.1. Esquema del Grupo 6 (Fuente)', level=2)
doc.add_paragraph(
    'Esquema PostgreSQL 17 normalizado en 3NF. '
    '25 tablas totales (incluye catálogos y extensiones del Grupo 2), '
    'de las cuales 9 son relevantes para la migración.'
)
add_table(
    ['Tabla', 'Registros', 'Columnas principales'],
    [
        ['persona', '35,331', 'id_persona, ci, nombre, fecha_nacimiento, sexo, direccion, telefono, matricula, id_zona'],
        ['paciente', '34,500', 'id_paciente, id_persona (tabla de rol)'],
        ['personal', '831', 'id_personal, id_persona, id_especialidad, id_cargo, id_turno'],
        ['especialidad', '24', 'id_especialidad, nombre'],
        ['zona', '30', 'id_zona, nombre, ciudad'],
        ['tipo_diagnostico', '25', 'id_tipo_diagnostico, nombre, categoria'],
        ['cita_medica', '70,000', 'id_cita, fecha_registro, fecha_cita, hora, numero_turno, estado'],
        ['diagnostico', '40,134', 'id_diagnostico, descripcion, observaciones, tipo_procedimiento'],
        ['receta', '40,134', 'id_receta, medicamentos, indicaciones, id_diagnostico'],
    ]
)

doc.add_heading('3.2. Diferencias identificadas', level=2)
add_table(
    ['Aspecto', 'Grupo 3', 'Grupo 6', 'Observación'],
    [
        ['Personas', 'PERSONA unificada', 'persona + paciente + personal', 'Roles en tablas separadas'],
        ['Especialidades', '15 registros', '24 registros', '9 adicionales sin equivalente en G3'],
        ['Zonas', '20 zonas', '30 zonas', 'Se asigna zona genérica ID=98'],
        ['Catálogos extra', 'No existen', '9 tablas cat_*', 'No relevantes para el DW'],
        ['Extensiones', 'No existen', 'pago, medicamento, procedimiento', 'Fuera del alcance del DW'],
        ['Tipo diagnóstico', 'IDs 1-25', 'IDs 1-25 (coinciden)', 'Mapeo directo'],
    ]
)

doc.add_heading('3.3. Información a transformar', level=2)

doc.add_heading('3.3.1. Estrategia de IDs', level=3)
add_table(
    ['Entidad', 'Rango original', 'Offset', 'Rango destino'],
    [
        ['Pacientes (persona)', '1 - 34,500', '+300,000', '300,001 - 334,500'],
        ['Personal (persona)', '1 - 831', '+600,000', '600,001 - 600,831'],
        ['Citas', '1 - 70,000', '+300,000', '300,001 - 370,000'],
        ['Diagnósticos', '1 - 40,134', '+300,000', '300,001 - 340,134'],
        ['Recetas', '1 - 40,134', '+300,000', '300,001 - 340,134'],
    ]
)

doc.add_heading('3.3.2. Extracción del dump SQL', level=3)
doc.add_paragraph(
    'El G6 entrega sus datos como dump PostgreSQL (encoding WIN1252). '
    'Se convierte encoding y se renombran tablas a staging:'
)
add_code(
    "iconv -f WINDOWS-1252 -t UTF-8 hospital_db_inserts.sql | \\\n"
    "  sed 's/INSERT INTO public\\.\\([a-z_]*\\)/INSERT INTO stg_g6_\\1/g' | \\\n"
    "  grep '^INSERT INTO stg_g6_\\(persona\\|paciente\\|personal\\|...\\) ' \\\n"
    "  > g6v2_inserts_staging.sql"
)

doc.add_heading('3.3.3. Transformación de pacientes', level=3)
doc.add_paragraph(
    'Se identifican pacientes mediante JOIN con la tabla de rol paciente:'
)
add_code(
    "SELECT\n"
    "    p.id_persona + 300000           AS id_persona,\n"
    "    'G6-PAC-' || p.id_persona       AS ci,\n"
    "    p.nombre, p.fecha_nacimiento, p.sexo,\n"
    "    p.direccion, p.telefono,\n"
    "    NULL                             AS matricula,\n"
    "    98                               AS id_zona\n"
    "INTO tfm_g6_persona_pacientes\n"
    "FROM stg_g6_persona p\n"
    "INNER JOIN stg_g6_paciente pac ON pac.id_persona = p.id_persona;"
)

doc.add_heading('3.3.4. Transformación de personal', level=3)
add_code(
    "SELECT\n"
    "    p.id_persona + 600000           AS id_persona,\n"
    "    'G6-MED-' || p.id_persona       AS ci,\n"
    "    p.nombre, p.fecha_nacimiento, p.sexo,\n"
    "    p.direccion, p.telefono,\n"
    "    'G6-MAT-' || per.id_personal    AS matricula,\n"
    "    98                               AS id_zona,\n"
    "    COALESCE(me.g3_id, 1)           AS id_especialidad\n"
    "INTO tfm_g6_persona_personal\n"
    "FROM stg_g6_persona p\n"
    "INNER JOIN stg_g6_personal per ON per.id_persona = p.id_persona\n"
    "LEFT JOIN tfm_g6_mapeo_especialidad me ON me.g6_id = per.id_especialidad;"
)

doc.add_heading('3.3.5. Transformación de citas y diagnósticos', level=3)
add_code(
    "-- CITA_MEDICA: offsets en paciente y médico\n"
    "SELECT\n"
    "    c.id_cita + 300000              AS id_cita,\n"
    "    c.fecha_registro, c.fecha_cita, c.hora,\n"
    "    c.numero_turno, c.estado,\n"
    "    c.id_paciente + 300000          AS id_paciente,\n"
    "    c.id_medico + 600000            AS id_medico\n"
    "INTO tfm_g6_cita_medica\n"
    "FROM stg_g6_cita_medica c;\n\n"
    "-- DIAGNOSTICO: tipo_diagnostico IDs 1-25 coinciden con G3\n"
    "SELECT\n"
    "    d.id_diagnostico + 300000       AS id_diagnostico,\n"
    "    d.descripcion, d.observaciones, d.tipo_procedimiento,\n"
    "    d.id_cita + 300000              AS id_cita,\n"
    "    d.id_tipo_diagnostico           AS id_tipo_diagnostico\n"
    "INTO tfm_g6_diagnostico\n"
    "FROM stg_g6_diagnostico d;"
)

doc.add_heading('3.3.6. Mapeo de Especialidades (24 → 15)', level=3)
doc.add_paragraph(
    'Las especialidades 1-15 coinciden con G3 (mapeo directo). '
    'Las 9 adicionales (16-24) se remapearon por afinidad clínica:'
)
add_table(
    ['ID G6', 'Especialidad G6', 'ID G3', 'Especialidad G3', 'Justificación'],
    [
        ['16', 'Medicina Interna', '1', 'Medicina General', 'Subespecialidad de medicina general'],
        ['17', 'Cirugía General', '7', 'Traumatología', 'Área quirúrgica más cercana'],
        ['18', 'Radiología', '1', 'Medicina General', 'Servicio de apoyo diagnóstico'],
        ['19', 'Anestesiología', '1', 'Medicina General', 'Servicio transversal'],
        ['20', 'Nefrología', '10', 'Urología', 'Misma área anatómica (sistema renal)'],
        ['21', 'Hematología', '1', 'Medicina General', 'Subespecialidad de medicina interna'],
        ['22', 'Reumatología', '7', 'Traumatología', 'Sistema musculoesquelético'],
        ['23', 'Infectología', '1', 'Medicina General', 'Subespecialidad de medicina interna'],
        ['24', 'Medicina Familiar', '1', 'Medicina General', 'Equivalente funcional'],
    ]
)

doc.add_heading('3.4. Resultado de la carga: G3 + G6', level=2)
add_table(
    ['Tabla DW', 'Antes (G3)', '+ Grupo 6', 'Después (Total)'],
    [
        ['dim_paciente', '4,500', '+34,500', '39,000'],
        ['dim_medico', '500', '+831', '1,331'],
        ['fact_atenciones', '15,000', '+40,134', '55,134'],
    ]
)

doc.add_heading('3.5. Reversibilidad', level=2)
add_code(
    "-- Eliminar datos G6 por rango de IDs\n"
    "DELETE FROM RECETA WHERE ID_Receta >= 300001;\n"
    "DELETE FROM DIAGNOSTICO WHERE ID_Diagnostico >= 300001;\n"
    "DELETE FROM CITA_MEDICA WHERE ID_Cita >= 300001;\n"
    "DELETE FROM PERSONA WHERE ID_Persona BETWEEN 300001 AND 400000;\n"
    "DELETE FROM PERSONA WHERE ID_Persona BETWEEN 600001 AND 700000;\n\n"
    "-- Limpiar staging y transformación\n"
    "DROP TABLE IF EXISTS stg_g6_* CASCADE;\n"
    "DROP TABLE IF EXISTS tfm_g6_* CASCADE;\n\n"
    "-- Restaurar secuencias\n"
    "SELECT setval('persona_id_persona_seq', (SELECT MAX(ID_Persona) FROM PERSONA));"
)

doc.add_page_break()

# ============================================================
# 4. RESULTADO FINAL: G3 + G1 + G6
# ============================================================
doc.add_heading('4. ANÁLISIS COMPARATIVO GRUPO 1 Y GRUPO 6 → GRUPO 3', level=1)

doc.add_heading('4.1. Resultado final del DW', level=2)
doc.add_paragraph(
    'Con ambos grupos cargados, el DW contiene datos de las 3 fuentes:'
)
add_table(
    ['Tabla DW', 'G3', 'G1', 'G6', 'Total'],
    [
        ['dim_sucursal', '—', '—', '—', '3'],
        ['dim_paciente', '4,500', '87,487', '34,500', '126,487'],
        ['dim_medico', '500', '13,337', '831', '14,668'],
        ['fact_atenciones', '15,000', '68,123', '40,134', '123,257'],
    ]
)

doc.add_heading('4.2. Distribución porcentual', level=2)

doc.add_paragraph('Pacientes por sucursal:')
add_table(
    ['Sucursal', 'Pacientes', '%'],
    [
        ['Grupo 3', '4,500', '3.6%'],
        ['Grupo 1', '87,487', '69.2%'],
        ['Grupo 6', '34,500', '27.3%'],
    ]
)
doc.add_paragraph()
doc.add_paragraph('Atenciones (hechos) por grupo:')
add_table(
    ['Grupo', 'Atenciones', '%'],
    [
        ['G3', '15,000', '12.2%'],
        ['G1', '68,123', '55.3%'],
        ['G6', '40,134', '32.6%'],
    ]
)
doc.add_paragraph()
doc.add_paragraph('Médicos por grupo:')
add_table(
    ['Grupo', 'Médicos', '%'],
    [
        ['G3', '500', '3.4%'],
        ['G1', '13,337', '90.9%'],
        ['G6', '831', '5.7%'],
    ]
)

doc.add_heading('4.3. Verificación de integridad', level=2)
add_table(
    ['Verificación', 'Resultado'],
    [
        ['Pacientes sin sucursal asignada', '0'],
        ['Atenciones sin paciente válido (FK)', '0'],
        ['Atenciones sin médico válido (FK)', '0'],
        ['Registros huérfanos en fact_atenciones', '0'],
    ]
)

doc.add_heading('4.4. ETL al DW', level=2)
doc.add_paragraph(
    'La carga al modelo estrella se ejecuta en 6 pasos. '
    'Extracción de dim_paciente:'
)
add_code(
    "SELECT p.CI, p.Nombre, p.Fecha_Nacimiento, p.Sexo,\n"
    "       p.Direccion, p.Telefono, z.Nombre, z.Ciudad,\n"
    "       CASE\n"
    "           WHEN p.CI LIKE 'G1-%' THEN 'G1'\n"
    "           WHEN p.CI LIKE 'G6-%' THEN 'G6'\n"
    "           ELSE 'G3'\n"
    "       END AS grupo_origen\n"
    "FROM PERSONA p\n"
    "JOIN ZONA z ON z.ID_Zona = p.ID_Zona\n"
    "WHERE p.Matricula IS NULL;"
)
doc.add_paragraph('Asignación de sucursal_key post-carga:')
add_code(
    "UPDATE dim_paciente SET sucursal_key = CASE grupo_origen\n"
    "    WHEN 'G3' THEN 1\n"
    "    WHEN 'G1' THEN 2\n"
    "    WHEN 'G6' THEN 3\n"
    "END;"
)
doc.add_paragraph('Carga de fact_atenciones (combina DIAGNOSTICO + CITA_MEDICA + TIPO_DIAGNOSTICO):')
add_code(
    "SELECT pac.CI, med.CI, c.Fecha_Cita,\n"
    "       EXTRACT(YEAR FROM c.Fecha_Cita)::INT  AS anio,\n"
    "       EXTRACT(MONTH FROM c.Fecha_Cita)::INT AS mes,\n"
    "       EXTRACT(QUARTER FROM c.Fecha_Cita)::INT AS trimestre,\n"
    "       TRIM(TO_CHAR(c.Fecha_Cita, 'Day'))    AS dia_semana,\n"
    "       c.Hora, c.Estado, c.Numero_Turno,\n"
    "       d.Descripcion, d.Observaciones,\n"
    "       d.Tipo_Procedimiento, td.Nombre, td.Categoria,\n"
    "       grupo_origen\n"
    "FROM DIAGNOSTICO d\n"
    "JOIN CITA_MEDICA c ON c.ID_Cita = d.ID_Cita\n"
    "JOIN PERSONA pac ON pac.ID_Persona = c.ID_Paciente\n"
    "JOIN PERSONA med ON med.ID_Persona = c.ID_Medico\n"
    "JOIN TIPO_DIAGNOSTICO td ON td.ID_Tipo_Diagnostico = d.ID_Tipo_Diagnostico;"
)
doc.add_paragraph('Resolución de FK en el DW mediante JOIN por CI:')
add_code(
    "INSERT INTO fact_atenciones (...)\n"
    "SELECT dp.paciente_key, dm.medico_key,\n"
    "       sf.fecha_cita, sf.anio, sf.mes, ...\n"
    "FROM stg_fact sf\n"
    "JOIN dim_paciente dp ON dp.ci = sf.paciente_ci\n"
    "JOIN dim_medico dm ON dm.ci = sf.medico_ci;"
)

# ============================================================
# GUARDAR
# ============================================================
output_path = '/home/tunek/Universidad/MATERIAS/bd-clinica/Informe_ETL_DataWarehouse.docx'
doc.save(output_path)
print(f'Documento generado: {output_path}')
