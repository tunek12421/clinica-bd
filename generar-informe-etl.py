#!/usr/bin/env python3
"""
Genera informe Word completo del ETL: Grupo 1 y Grupo 6
1. Análisis comparativo de esquemas
2. Información a transformar
3. Gráficos estadísticos antes y después
"""

import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

OUTPUT = '/home/tunek/Universidad/MATERIAS/bd-clinica/Informe_ETL_Grupo1.docx'
IMG_DIR = '/tmp/etl_charts'
os.makedirs(IMG_DIR, exist_ok=True)

# =============================================================================
# DATOS
# =============================================================================

# --- Citas por Especialidad ---
esp_antes = [
    ('Medicina General', 3542), ('Cardiología', 3097), ('Pediatría', 2170),
    ('Dermatología', 1905), ('Neurología', 1764), ('Ginecología', 1660),
    ('Oftalmología', 1220), ('Traumatología', 1001), ('Otorrinolaringología', 760),
    ('Urología', 741), ('Gastroenterología', 706), ('Neumología', 483),
    ('Psiquiatría', 426), ('Endocrinología', 331), ('Oncología', 194),
]

esp_despues = [
    ('Medicina General', 71813), ('Traumatología', 30307), ('Urología', 20429),
    ('Cardiología', 18921), ('Pediatría', 17088), ('Neurología', 16830),
    ('Dermatología', 16157), ('Ginecología', 14473), ('Oftalmología', 10706),
    ('Oncología', 10229), ('Otorrinolaringología', 8756), ('Endocrinología', 6514),
    ('Psiquiatría', 6236), ('Gastroenterología', 706), ('Neumología', 483),
]

# --- Citas por Mes ---
meses_nombre = ['Ene','Feb','Mar','Abr','May','Jun','Jul','Ago','Sep','Oct','Nov','Dic']
mes_antes = [1432, 1432, 1707, 1682, 1780, 2406, 2573, 2660, 1230, 1178, 1198, 722]
mes_despues = [23268, 21466, 23532, 22924, 22545, 22796, 23362, 23357, 21417, 22013, 21623, 21697]

# --- Diagnósticos por Tipo (top 10) ---
diag_antes = [
    ('Clínico', 2142), ('Por Imagen', 1735), ('Laboratorio', 1389),
    ('Diferencial', 1228), ('Presuntivo', 1084), ('Definitivo', 1027),
    ('Prenatal', 890), ('Molecular', 750), ('Patológico', 737),
    ('Funcional', 620),
]

diag_despues = [
    ('Definitivo', 43780), ('Presuntivo', 43721), ('Clínico', 33940),
    ('Laboratorio', 15025), ('Ambulatorio', 8954), ('Por Imagen', 6307),
    ('Nutricional', 4802), ('Endoscópico', 4628), ('Diferencial', 1228),
    ('Prenatal', 890),
]

# =============================================================================
# GENERAR GRÁFICOS
# =============================================================================
plt.rcParams['font.size'] = 10
plt.rcParams['figure.dpi'] = 150

def color_list(n, cmap='tab20'):
    cm = plt.get_cmap(cmap)
    return [cm(i/n) for i in range(n)]

# --- Especialidades ANTES ---
fig, ax = plt.subplots(figsize=(10, 5))
names = [e[0] for e in esp_antes]; vals = [e[1] for e in esp_antes]
ax.barh(names[::-1], vals[::-1], color=color_list(len(names))[::-1])
ax.set_xlabel('Total de Citas'); ax.set_title('Citas por Especialidad — ANTES (Solo Grupo 3)', fontweight='bold', fontsize=12)
for i, v in enumerate(vals[::-1]): ax.text(v + 50, i, f'{v:,}', va='center', fontsize=8)
plt.tight_layout(); plt.savefig(f'{IMG_DIR}/esp_antes.png'); plt.close()

# --- Especialidades DESPUÉS ---
fig, ax = plt.subplots(figsize=(10, 5))
names = [e[0] for e in esp_despues]; vals = [e[1] for e in esp_despues]
ax.barh(names[::-1], vals[::-1], color=color_list(len(names))[::-1])
ax.set_xlabel('Total de Citas'); ax.set_title('Citas por Especialidad — DESPUÉS (G3 + G1 + G6)', fontweight='bold', fontsize=12)
for i, v in enumerate(vals[::-1]): ax.text(v + 500, i, f'{v:,}', va='center', fontsize=8)
plt.tight_layout(); plt.savefig(f'{IMG_DIR}/esp_despues.png'); plt.close()

# --- Meses ANTES ---
fig, ax = plt.subplots(figsize=(10, 4.5))
x = np.arange(12)
ax.bar(x, mes_antes, color='#4472C4', width=0.6); ax.set_xticks(x); ax.set_xticklabels(meses_nombre)
ax.set_ylabel('Total'); ax.set_title('Citas por Mes — ANTES (Solo Grupo 3)', fontweight='bold', fontsize=12)
for i, v in enumerate(mes_antes): ax.text(i, v+40, f'{v:,}', ha='center', fontsize=8)
plt.tight_layout(); plt.savefig(f'{IMG_DIR}/mes_antes.png'); plt.close()

# --- Meses DESPUÉS ---
fig, ax = plt.subplots(figsize=(10, 4.5))
ax.bar(x, mes_despues, color='#ED7D31', width=0.6); ax.set_xticks(x); ax.set_xticklabels(meses_nombre)
ax.set_ylabel('Total'); ax.set_title('Citas por Mes — DESPUÉS (G3 + G1 + G6)', fontweight='bold', fontsize=12)
for i, v in enumerate(mes_despues): ax.text(i, v+200, f'{v:,}', ha='center', fontsize=8)
plt.tight_layout(); plt.savefig(f'{IMG_DIR}/mes_despues.png'); plt.close()

# --- Diagnósticos ANTES ---
fig, ax = plt.subplots(figsize=(10, 5))
names = [d[0] for d in diag_antes]; vals = [d[1] for d in diag_antes]
ax.barh(names[::-1], vals[::-1], color=color_list(len(names), 'Set3')[::-1])
ax.set_xlabel('Total'); ax.set_title('Diagnósticos por Tipo (Top 10) — ANTES (Solo Grupo 3)', fontweight='bold', fontsize=12)
for i, v in enumerate(vals[::-1]): ax.text(v+20, i, f'{v:,}', va='center', fontsize=8)
plt.tight_layout(); plt.savefig(f'{IMG_DIR}/diag_antes.png'); plt.close()

# --- Diagnósticos DESPUÉS ---
fig, ax = plt.subplots(figsize=(10, 5))
names = [d[0] for d in diag_despues]; vals = [d[1] for d in diag_despues]
ax.barh(names[::-1], vals[::-1], color=color_list(len(names), 'Set3')[::-1])
ax.set_xlabel('Total'); ax.set_title('Diagnósticos por Tipo (Top 10) — DESPUÉS (G3 + G1 + G6)', fontweight='bold', fontsize=12)
for i, v in enumerate(vals[::-1]): ax.text(v+300, i, f'{v:,}', va='center', fontsize=8)
plt.tight_layout(); plt.savefig(f'{IMG_DIR}/diag_despues.png'); plt.close()

print("Gráficos generados.")

# =============================================================================
# DOCUMENTO WORD
# =============================================================================
doc = Document()
style = doc.styles['Normal']; style.font.name = 'Calibri'; style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs: run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    return h

def add_table(headers, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(headers), style='Light Grid Accent 1')
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]; cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs: r.bold = True; r.font.size = Pt(9)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri+1].cells[ci]; cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs: r.font.size = Pt(9)
    return t

# --- PORTADA ---
doc.add_paragraph(); doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Informe de Reconocimiento y Migración ETL'); r.bold = True; r.font.size = Pt(24); r.font.color.rgb = RGBColor(0x1F,0x3A,0x5F)

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Grupo 1 y Grupo 6 → Grupo 3'); r.font.size = Pt(16); r.font.color.rgb = RGBColor(0x4F,0x81,0xBD)

doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('Materia: Base de Datos\n').font.size = Pt(12)
p.add_run('Universidad Privada Domingo Savio\n').font.size = Pt(12)
p.add_run('Febrero 2026').font.size = Pt(12)
doc.add_page_break()

# =============================================================================
# SECCIÓN 1: ANÁLISIS
# =============================================================================
add_heading('1. Análisis Comparativo', level=1)
doc.add_paragraph(
    'Se comparan las bases de datos del Grupo 1 y Grupo 6 (fuentes) con la base de datos '
    'del Grupo 3 (destino), identificando diferencias estructurales, volumétricas y de calidad.'
)

# --- 1.1 Esquema Grupo 3 ---
add_heading('1.1 Esquema del Grupo 3 (Destino)', level=2)
doc.add_paragraph('Nuestra base de datos cuenta con 8 tablas normalizadas en PostgreSQL, '
    '5 vistas de documentación y comentarios internos (COMMENT ON).')
add_table(['Tabla', 'Columnas', 'Registros', 'Estructura'], [
    ('ESPECIALIDAD', 'ID, Nombre', '15', '2 cols'),
    ('TIPO_DIAGNOSTICO', 'ID, Nombre, Categoria', '25', '3 cols'),
    ('ZONA', 'ID, Nombre, Ciudad', '20', '3 cols'),
    ('PERSONA', 'ID, CI, Nombre, FechaNac, Sexo, Dir, Tel, Matric, Zona, Esp', '5,000', '10 cols'),
    ('HORARIO_MEDICO', 'ID, Dia, HoraInicio, HoraFin, Cupo, ID_Persona', '2,500', '6 cols'),
    ('CITA_MEDICA', 'ID, FechaReg, FechaCita, Hora, Turno, Estado, Paciente, Medico', '20,000', '8 cols'),
    ('DIAGNOSTICO', 'ID, Descripcion, Observaciones, TipoProc, ID_Cita, ID_Tipo', '15,000', '6 cols'),
    ('RECETA', 'ID, Medicamentos, Indicaciones, ID_Diagnostico', '7,440', '4 cols'),
])

# --- 1.2 Esquema Grupo 1 ---
add_heading('1.2 Esquema del Grupo 1 (Fuente)', level=2)
doc.add_paragraph('Base de datos alojada en Supabase (PostgreSQL gestionado en AWS us-west-2). '
    '4 tablas sin normalización de catálogos, sin documentación interna.')
add_table(['Tabla', 'Columnas', 'Registros', 'Observación'], [
    ('pacientes', 'paciente_id, nombre, fecha_nac, genero', '50,000', 'Sin CI, dirección, teléfono, zona'),
    ('personal', 'personal_id, nombre, cargo, especialidad', '50,000', 'Especialidad como texto, no normalizada'),
    ('atenciones', 'atencion_id, paciente_id, personal_id, fecha, estado, motivo', '50,000', 'Sin fecha_registro, sin turno'),
    ('diagnosticos', 'diagnostico_id, atencion_id, codigo_cie10, descripcion, severidad', '68,123', 'Usa CIE-10 (estándar internacional)'),
])

# --- 1.3 Esquema Grupo 6 ---
add_heading('1.3 Esquema del Grupo 6 (Fuente)', level=2)
doc.add_paragraph('Base de datos en SQLite (archivo hospital.db de 121 MB). '
    '9 tablas con estructura más completa incluyendo historia clínica, signos vitales y recetas.')
add_table(['Tabla', 'Columnas', 'Registros', 'Observación'], [
    ('Paciente', 'paciente_id, nombres, apellidos, tel, correo, dir, fecha_nac', '300,000', 'Tiene dirección y teléfono'),
    ('Especialidad', 'especialidad_id, nombre, descripcion', '24', '24 especialidades (vs 15 nuestras)'),
    ('Personal', 'personal_id, nombres, apellidos, rol, especialidad_id, fecha_contrat', '3,000', 'Roles: Médico, Enfermero, Administrativo'),
    ('Cita', 'cita_id, paciente_id, medico_id, fecha, hora, estado, fecha_creacion', '200,000', 'Tiene fecha_creación'),
    ('Historia_Clinica', 'historia_id, paciente_id, fecha_apertura, estado', '300,000', 'Sin equivalente en nuestro esquema'),
    ('Atencion_Medica', 'atencion_id, historia_id, cita_id, fecha_hora, motivo, notas', '55,390', 'Vincula historia con cita'),
    ('Signos_Vitales', 'signo_id, atencion_id, presion, FC, FR, temp, SpO2', '85,390', 'Sin equivalente en nuestro esquema'),
    ('Diagnostico', 'diagnostico_id, atencion_id, codigo_cie10, descripcion, tipo', '85,390', 'Tipos: Presuntivo y Confirmado'),
    ('Receta', 'receta_id, diagnostico_id, medicamento, dosis, frec, duración, indic', '230,696', 'Detalle por medicamento individual'),
])

# --- 1.4 Comparación ---
add_heading('1.4 Tabla Comparativa General', level=2)
add_table(['Aspecto', 'Grupo 3 (Destino)', 'Grupo 1', 'Grupo 6'], [
    ('Motor', 'PostgreSQL 17 (Docker)', 'PostgreSQL (Supabase)', 'SQLite'),
    ('Tablas', '8', '4', '9'),
    ('Registros', '50,000', '218,123', '959,866'),
    ('Normalización', 'Alta (3FN)', 'Baja', 'Media'),
    ('Catálogos', 'Especialidad, Zona, TipoDiag', 'Ninguno', 'Solo Especialidad'),
    ('Documentación', 'COMMENT ON + 5 vistas', 'Ninguna', 'Comentarios en SQL'),
    ('CI / Documento', 'Sí (NOT NULL)', 'No', 'No'),
    ('Dirección', 'Sí (NOT NULL)', 'No', 'Sí'),
    ('Teléfono', 'Sí (NOT NULL)', 'No', 'Sí'),
    ('Zonas', 'Sí (20 zonas)', 'No', 'No'),
    ('Recetas', 'Sí', 'No', 'Sí (detallado)'),
    ('Hist. Clínica', 'No', 'No', 'Sí'),
    ('Signos Vitales', 'No', 'No', 'Sí'),
    ('CIE-10', 'No', 'Sí', 'Sí'),
])

add_heading('1.5 Resumen del Análisis', level=2)
doc.add_paragraph(
    'El Grupo 6 aporta el mayor volumen de datos (959,866 registros) con un esquema más '
    'completo que incluye historia clínica, signos vitales y recetas detalladas. El Grupo 1 '
    'aporta 218,123 registros con un esquema más simple pero con códigos CIE-10. '
    'Ambos carecen de información de CI (cédula) y zonas geográficas. '
    'El Grupo 6 sí tiene dirección y teléfono de pacientes, datos que se preservan en la migración.'
)

doc.add_page_break()

# =============================================================================
# SECCIÓN 2: INFORMACIÓN A TRANSFORMAR
# =============================================================================
add_heading('2. Información a Transformar', level=1)
doc.add_paragraph(
    'El proceso ETL se implementó en archivos SQL separados por grupo y fase, '
    'ubicados en etl/grupo1/ y etl/grupo6/. Cada proceso es completamente reversible.'
)

# --- 2.1 ETL Grupo 1 ---
add_heading('2.1 Transformaciones — Grupo 1', level=2)

add_heading('Estrategia de IDs', level=3)
add_table(['Entidad', 'Rango Original', 'Offset', 'Rango Destino'], [
    ('Pacientes → PERSONA', '1 - 50,000', '+100,000', '100,001 - 150,000'),
    ('Personal → PERSONA', '1 - 50,000', '+200,000', '200,001 - 250,000'),
    ('Atenciones → CITA_MEDICA', '1 - 50,000', '+100,000', '100,001 - 150,000'),
    ('Diagnósticos → DIAGNOSTICO', '1 - 68,123', '+100,000', '100,001 - 168,123'),
])

add_heading('Mapeo de Campos Principales', level=3)
add_table(['Campo Destino', 'Disponible', 'Transformación'], [
    ('PERSONA.CI', 'No', "'G1-PAC-{id}' / 'G1-PER-{id}'"),
    ('PERSONA.Nombre', 'Sí', 'Directo'),
    ('PERSONA.Sexo', 'Solo pacientes', "Personal: 'X' (no disponible)"),
    ('PERSONA.Direccion', 'No', "'Sin dato (Grupo 1)'"),
    ('PERSONA.ID_Zona', 'No', 'Zona especial ID=99'),
    ('PERSONA.ID_Especialidad', 'Texto', 'JOIN con tabla ESPECIALIDAD por nombre'),
    ('CITA.Hora', 'En timestamp', 'fecha_atencion::TIME'),
    ('CITA.Numero_Turno', 'No', 'ROW_NUMBER() por día+médico'),
    ('DIAG.Observaciones', 'Parcial', "'CIE-10: {cod} | Severidad: {sev}'"),
    ('DIAG.ID_Tipo_Diagnostico', 'CIE-10', 'Mapeo código → categoría clínica'),
])

add_heading('Mapeo CIE-10 → Tipo de Diagnóstico', level=3)
add_table(['Código CIE-10', 'Diagnóstico', 'Tipo Asignado'], [
    ('I10, R10.9, R51, J02.9, J06.9, J20.9, L30.9', 'Hipertensión, dolor, infecciones resp.', 'Clínico'),
    ('M54.5', 'Dolor lumbar bajo', 'Por Imagen'),
    ('E78.5, E11.9, N39.0', 'Hiperlipidemia, diabetes, ITU', 'Laboratorio'),
    ('E66.9', 'Obesidad', 'Nutricional'),
    ('K29.5', 'Gastritis crónica', 'Endoscópico'),
    ('Z00.0, Z71.1', 'Examen general, consulta', 'Ambulatorio'),
])

# --- 2.2 ETL Grupo 6 ---
add_heading('2.2 Transformaciones — Grupo 6', level=2)

add_heading('Estrategia de IDs', level=3)
add_table(['Entidad', 'Rango Original', 'Offset', 'Rango Destino'], [
    ('Pacientes → PERSONA', '1 - 300,000', '+300,000', '300,001 - 600,000'),
    ('Personal → PERSONA', '1 - 3,000', '+600,000', '600,001 - 603,000'),
    ('Citas → CITA_MEDICA', '1 - 200,000', '+300,000', '300,001 - 500,000'),
    ('Diagnósticos → DIAGNOSTICO', '1 - 85,390', '+300,000', '300,001 - 385,390'),
    ('Recetas → RECETA', '1 - 230,696', '+300,000', '300,001 - 530,696'),
])

add_heading('Mapeo de Campos Principales', level=3)
add_table(['Campo Destino', 'Disponible', 'Transformación'], [
    ('PERSONA.CI', 'No', "'G6-PAC-{id}' / 'G6-PER-{id}'"),
    ('PERSONA.Nombre', 'nombres + apellidos', 'Concatenar con espacio'),
    ('PERSONA.Sexo', 'No', "'X' (no disponible)"),
    ('PERSONA.Direccion', 'Sí (pacientes)', 'Directo / personal: Sin dato'),
    ('PERSONA.Telefono', 'Sí (pacientes)', 'Directo / personal: Sin dato'),
    ('PERSONA.ID_Zona', 'No', 'Zona especial ID=98'),
    ('PERSONA.ID_Especialidad', 'FK normalizada', 'Mapeo 24 esp. G6 → 15 nuestras'),
    ('CITA.Fecha_Registro', 'fecha_creacion', 'Directo (disponible en G6)'),
    ('DIAG.Observaciones', 'Parcial', "'CIE-10: {cod} | Tipo: {tipo}'"),
    ('RECETA.Medicamentos', 'medicamento + dosis', 'Concatenar'),
    ('RECETA.Indicaciones', 'frecuencia + duración + indic.', 'Concatenar campos'),
])

add_heading('Mapeo de Especialidades (24 → 15)', level=3)
doc.add_paragraph('De las 24 especialidades del Grupo 6, 13 coinciden directamente con las nuestras. '
    'Las 11 restantes se mapearon a la especialidad más cercana:')
add_table(['Especialidad Grupo 6', '→ Nuestra Especialidad'], [
    ('Medicina Interna, Anestesiología, Radiología, Hematología', '→ Medicina General'),
    ('Infectología, Medicina Familiar, Med. Preventiva, Med. Emergencias, Patología', '→ Medicina General'),
    ('Cirugía General, Reumatología', '→ Traumatología'),
    ('Nefrología', '→ Urología'),
    ('(13 restantes)', '→ Coincidencia directa'),
])

# --- 2.3 Resultado ---
add_heading('2.3 Resultado de la Carga', level=2)
add_table(['Tabla', 'Grupo 3', '+ Grupo 1', '+ Grupo 6', 'Total Final'], [
    ('PERSONA', '5,000', '+100,000', '+303,000', '408,000'),
    ('CITA_MEDICA', '20,000', '+50,000', '+200,000', '270,000'),
    ('DIAGNOSTICO', '15,000', '+68,123', '+85,390', '168,513'),
    ('RECETA', '7,440', '—', '+230,696', '238,136'),
    ('TOTAL', '47,440', '+218,123', '+819,086', '1,084,649'),
])

doc.add_paragraph('Integridad referencial verificada post-carga: 0 errores en todas las relaciones FK.')

add_heading('2.4 Reversibilidad', level=2)
doc.add_paragraph(
    'Cada ETL es completamente reversible mediante sus respectivos scripts de rollback '
    '(04-rollback.sql), que eliminan los datos por rango de IDs y restauran las secuencias.'
)

doc.add_page_break()

# =============================================================================
# SECCIÓN 3: GRÁFICOS
# =============================================================================
add_heading('3. Gráficos Estadísticos: Antes y Después', level=1)
doc.add_paragraph(
    'Los siguientes gráficos comparan los datos originales del Grupo 3 (antes) con los datos '
    'consolidados después de integrar los Grupos 1 y 6 (después). '
    'Se analizan tres dimensiones: especialidad, temporalidad y tipo de diagnóstico.'
)

# --- 3.1 ---
add_heading('3.1 Citas por Especialidad', level=2)
add_heading('Antes (Solo Grupo 3 — 20,000 citas)', level=3)
doc.add_picture(f'{IMG_DIR}/esp_antes.png', width=Inches(6))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
add_heading('Después (Grupo 3 + Grupo 1 + Grupo 6 — 270,000 citas)', level=3)
doc.add_picture(f'{IMG_DIR}/esp_despues.png', width=Inches(6))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph(
    'Observación: Medicina General se consolida como la especialidad dominante (71,813 citas) '
    'debido al mapeo de múltiples especialidades del Grupo 6 (Medicina Interna, Familiar, '
    'Preventiva, Emergencias, etc.). Traumatología sube al segundo lugar (30,307) por la '
    'inclusión de Cirugía General y Reumatología del Grupo 6. El volumen total pasa de '
    '20,000 a 270,000 citas (incremento de 13.5x).'
)

# --- 3.2 ---
add_heading('3.2 Citas por Mes', level=2)
add_heading('Antes (Solo Grupo 3 — 20,000 citas)', level=3)
doc.add_picture(f'{IMG_DIR}/mes_antes.png', width=Inches(6))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
add_heading('Después (Grupo 3 + Grupo 1 + Grupo 6 — 270,000 citas)', level=3)
doc.add_picture(f'{IMG_DIR}/mes_despues.png', width=Inches(6))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph(
    'Observación: El patrón estacional pronunciado del Grupo 3 (pico invernal junio-agosto, 35%) '
    'se suaviza drásticamente al incorporar los datos uniformes de los Grupos 1 y 6. '
    'La distribución final es casi uniforme (~22,000/mes), con variación máxima de solo 1.1x '
    'entre el mes más alto (marzo, 23,532) y el más bajo (septiembre, 21,417). '
    'Esto indica que los patrones estacionales son específicos de nuestro dataset.'
)

# --- 3.3 ---
add_heading('3.3 Diagnósticos por Tipo', level=2)
add_heading('Antes (Solo Grupo 3 — 15,000 diagnósticos)', level=3)
doc.add_picture(f'{IMG_DIR}/diag_antes.png', width=Inches(6))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
add_heading('Después (Grupo 3 + Grupo 1 + Grupo 6 — 168,513 diagnósticos)', level=3)
doc.add_picture(f'{IMG_DIR}/diag_despues.png', width=Inches(6))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph(
    'Observación: El Grupo 6 aporta masivamente diagnósticos de tipo Definitivo (43,780) y '
    'Presuntivo (43,721), que eran minoritarios en nuestros datos. Esto refleja que el Grupo 6 '
    'clasifica sus diagnósticos por certeza (Confirmado/Presuntivo) mientras que nosotros '
    'clasificamos por método (Clínico, Laboratorio, Imagen, etc.). El Diagnóstico Clínico '
    'se mantiene en tercer lugar (33,940) por la contribución del Grupo 1 vía mapeo CIE-10.'
)

doc.save(OUTPUT)
print(f"Informe generado: {OUTPUT}")
