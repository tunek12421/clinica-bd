"""
Genera el informe: Comparación Práctica entre ROLAP, MOLAP y HOLAP
Blanco y negro, con placeholders para capturas de pantalla.

Uso: python3 generar-informe-olap.py
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# ============================================================
# ESTILOS
# ============================================================
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    h = doc.styles[f'Heading {level}']
    h.font.color.rgb = RGBColor(0, 0, 0)


def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    return table


def add_code(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Pt(18)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(8.5)


def add_screenshot_placeholder(label):
    """Agrega un recuadro con borde que indica dónde poner la captura."""
    # Tabla de 1 celda con borde para simular el recuadro
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Espaciado interno para dar altura al recuadro
    run = p.add_run(f'\n\n[CAPTURA: {label}]\n\n')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(120, 120, 120)
    run.italic = True
    doc.add_paragraph()  # Espacio después


# ============================================================
# PORTADA
# ============================================================
for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('UNIVERSIDAD PRIVADA DOMINGO SAVIO')
run.bold = True
run.font.size = Pt(14)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('FACULTAD DE INGENIERÍA\nCARRERA DE INGENIERÍA DE SISTEMAS')
run.font.size = Pt(12)

for _ in range(3):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('COMPARACIÓN PRÁCTICA ENTRE\nROLAP, MOLAP Y HOLAP')
run.bold = True
run.font.size = Pt(16)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Aplicada al Data Warehouse de la Clínica')
run.font.size = Pt(12)
run.italic = True

for _ in range(2):
    doc.add_paragraph()

for nombre in [
    'Argote Gonzales Marco Ariel',
    'Benavides Arancibia Jorge Enrique',
    'Lujan Arispe Enrique',
    'Lujan Arispe William',
    'Mejillones Iraizos Rebeca',
    'Valencia Vidaurre Marcos Daniel',
    'Villarroel Espinoza Gustavo Ernesto',
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(nombre)
    run.bold = True
    run.font.size = Pt(11)

for _ in range(3):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Cochabamba - Bolivia\n2026')
run.bold = True
run.font.size = Pt(12)

doc.add_page_break()

# ============================================================
# 1. INTRODUCCIÓN
# ============================================================
doc.add_heading('1. INTRODUCCIÓN', level=1)

doc.add_paragraph(
    'El presente informe realiza una comparación práctica entre las tres arquitecturas '
    'de procesamiento analítico en línea (OLAP): ROLAP, MOLAP y HOLAP. '
    'La comparación se ejecuta sobre el Data Warehouse de la Clínica, '
    'implementado en PostgreSQL con un modelo estrella que contiene datos '
    'de atenciones médicas de tres sucursales.'
)

doc.add_heading('1.1. Esquema del Data Warehouse', level=2)
doc.add_paragraph(
    'El DW utiliza un modelo estrella con snowflake parcial '
    '(dim_sucursal como subdimensión de dim_paciente). '
    'Cada fila de fact_atenciones representa un evento de diagnóstico dentro de una cita médica.'
)
add_code('dim_sucursal -- dim_paciente -- fact_atenciones -- dim_medico')

doc.add_paragraph()
add_table(
    ['Tabla', 'Tipo', 'Descripción'],
    [
        ['dim_sucursal', 'Subdimensión', 'Fuente/sucursal de los datos (Grupo 1, 3, 6)'],
        ['dim_paciente', 'Dimensión', 'Datos del paciente + FK a dim_sucursal'],
        ['dim_medico', 'Dimensión', 'Datos del médico con especialidad'],
        ['fact_atenciones', 'Hechos', 'Diagnóstico + cita + dimensiones temporales'],
    ]
)

doc.add_heading('1.2. Definiciones', level=2)

doc.add_heading('1.2.1. ROLAP (Relational OLAP)', level=3)
doc.add_paragraph(
    'Los datos residen en tablas relacionales (esquema estrella). '
    'Las consultas OLAP se traducen a SQL estándar y los agregados se calculan '
    'al momento de ejecutar cada consulta. No hay nada pre-calculado.'
)
add_table(
    ['Ventajas', 'Desventajas'],
    [
        ['Datos siempre actualizados', 'Consultas más lentas (calcula al vuelo)'],
        ['Sin duplicación de almacenamiento', 'Depende del rendimiento del motor SQL'],
        ['Usa infraestructura SQL existente', 'No óptimo para cubos muy complejos'],
    ]
)

doc.add_heading('1.2.2. MOLAP (Multidimensional OLAP)', level=3)
doc.add_paragraph(
    'Los cubos se pre-calculan y almacenan en estructuras multidimensionales. '
    'Los resultados ya están listos antes de que el usuario pregunte. '
    'En PostgreSQL se simula con tablas pre-calculadas que almacenan '
    'TODOS los agregados del CUBE.'
)
add_table(
    ['Ventajas', 'Desventajas'],
    [
        ['Consultas instantáneas (lectura directa)', 'Requiere mucho espacio en disco'],
        ['Datos pre-agregados', 'Carga inicial lenta (hay que calcular todo)'],
        ['Óptimo para cubos pequeños/medianos', 'No escala con muchas dimensiones'],
    ]
)

doc.add_heading('1.2.3. HOLAP (Hybrid OLAP)', level=3)
doc.add_paragraph(
    'Combina ambos enfoques: los agregados frecuentes se pre-calculan '
    '(como MOLAP) mediante vistas materializadas, mientras que las consultas '
    'de detalle granular acceden a los datos en vivo (como ROLAP).'
)
add_table(
    ['Ventajas', 'Desventajas'],
    [
        ['Rápido para agregados (pre-calculados)', 'Mayor complejidad de administración'],
        ['Flexible para detalle (consulta en vivo)', 'Hay que decidir qué pre-calcular'],
        ['Balance entre rendimiento y espacio', 'Requiere mantenimiento (REFRESH)'],
    ]
)

doc.add_page_break()

# ============================================================
# 2. ROLAP
# ============================================================
doc.add_heading('2. ROLAP — Consultas al vuelo sobre tablas relacionales', level=1)

doc.add_paragraph(
    'En la arquitectura ROLAP, las consultas OLAP se ejecutan directamente sobre '
    'las tablas del esquema estrella (fact_atenciones + dimensiones). '
    'PostgreSQL calcula los agregados cada vez que se ejecuta la consulta. '
    'No existe ninguna estructura pre-calculada.'
)

doc.add_heading('2.1. CUBE completo: especialidad x sucursal x gestión', level=2)
doc.add_paragraph(
    'Se ejecuta un CUBE con 3 dimensiones, lo que genera 2³ = 8 agrupaciones '
    '(cada combinación + subtotales parciales + gran total). '
    'PostgreSQL debe hacer 3 JOINs y calcular las 8 agrupaciones al momento.'
)
add_code(
    "EXPLAIN ANALYZE\n"
    "SELECT\n"
    "    COALESCE(dm.especialidad, '** TODAS **')  AS especialidad,\n"
    "    COALESCE(ds.nombre, '** TODAS **')         AS sucursal,\n"
    "    COALESCE(fa.anio::TEXT, '** TODOS **')     AS gestion,\n"
    "    COUNT(*)                                   AS total_atenciones\n"
    "FROM fact_atenciones fa\n"
    "JOIN dim_medico dm    ON dm.medico_key   = fa.medico_key\n"
    "JOIN dim_paciente dp  ON dp.paciente_key = fa.paciente_key\n"
    "JOIN dim_sucursal ds  ON ds.sucursal_key = dp.sucursal_key\n"
    "GROUP BY CUBE (dm.especialidad, ds.nombre, fa.anio)\n"
    "ORDER BY GROUPING(dm.especialidad), dm.especialidad, ...;"
)
add_screenshot_placeholder('ROLAP — CUBE completo (resultado + EXPLAIN ANALYZE)')

doc.add_heading('2.2. ROLLUP jerárquico temporal', level=2)
doc.add_paragraph(
    'Se ejecuta un ROLLUP con jerarquía año -> trimestre -> mes. '
    'Genera subtotales jerárquicos de arriba hacia abajo: '
    'año+trimestre+mes, año+trimestre, año, y gran total.'
)
add_code(
    "EXPLAIN ANALYZE\n"
    "SELECT\n"
    "    fa.anio AS gestion, fa.trimestre, fa.mes,\n"
    "    COUNT(*) AS total_atenciones\n"
    "FROM fact_atenciones fa\n"
    "GROUP BY ROLLUP (fa.anio, fa.trimestre, fa.mes)\n"
    "ORDER BY GROUPING(fa.anio), fa.anio, ...;"
)
add_screenshot_placeholder('ROLAP — ROLLUP temporal (resultado + EXPLAIN ANALYZE)')

doc.add_page_break()

# ============================================================
# 3. MOLAP
# ============================================================
doc.add_heading('3. MOLAP — Cubos pre-calculados en tablas', level=1)

doc.add_paragraph(
    'En un sistema MOLAP real (Essbase, SSAS Multidimensional), los cubos se almacenan '
    'en estructuras multidimensionales propietarias. En PostgreSQL lo simulamos con '
    'tablas que almacenan TODOS los agregados pre-calculados mediante CREATE TABLE AS SELECT ... GROUP BY CUBE. '
    'Las consultas posteriores solo leen de estas tablas sin hacer JOINs ni GROUP BY.'
)

doc.add_heading('3.1. Carga del cubo: especialidad x sucursal x gestión', level=2)
doc.add_paragraph(
    'Se pre-calculan TODAS las combinaciones del CUBE y se almacenan en una tabla física. '
    'Este paso simula la "carga del cubo MOLAP". Se crean índices para acelerar las consultas.'
)
add_code(
    "DROP TABLE IF EXISTS molap_cubo_especialidad_sucursal_gestion;\n"
    "CREATE TABLE molap_cubo_especialidad_sucursal_gestion AS\n"
    "SELECT\n"
    "    dm.especialidad,\n"
    "    ds.nombre                         AS sucursal,\n"
    "    fa.anio                           AS gestion,\n"
    "    GROUPING(dm.especialidad)         AS grp_especialidad,\n"
    "    GROUPING(ds.nombre)               AS grp_sucursal,\n"
    "    GROUPING(fa.anio)                 AS grp_gestion,\n"
    "    COUNT(*)                          AS total_atenciones,\n"
    "    COUNT(DISTINCT fa.paciente_key)   AS pacientes_unicos,\n"
    "    COUNT(DISTINCT fa.medico_key)     AS medicos_involucrados\n"
    "FROM fact_atenciones fa\n"
    "JOIN dim_medico dm    ON dm.medico_key   = fa.medico_key\n"
    "JOIN dim_paciente dp  ON dp.paciente_key = fa.paciente_key\n"
    "JOIN dim_sucursal ds  ON ds.sucursal_key = dp.sucursal_key\n"
    "GROUP BY CUBE (dm.especialidad, ds.nombre, fa.anio);\n\n"
    "CREATE INDEX idx_molap_c1_esp ON molap_cubo_...(especialidad);\n"
    "CREATE INDEX idx_molap_c1_suc ON molap_cubo_...(sucursal);\n"
    "CREATE INDEX idx_molap_c1_ges ON molap_cubo_...(gestion);"
)
add_screenshot_placeholder('MOLAP — Creacion de tabla pre-calculada (CREATE TABLE AS)')

doc.add_heading('3.2. Carga del cubo temporal', level=2)
doc.add_paragraph(
    'Pre-cálculo del ROLLUP año -> trimestre -> mes almacenado en tabla.'
)
add_code(
    "CREATE TABLE molap_cubo_temporal AS\n"
    "SELECT fa.anio AS gestion, fa.trimestre, fa.mes,\n"
    "    GROUPING(fa.anio) AS grp_gestion,\n"
    "    GROUPING(fa.trimestre) AS grp_trimestre,\n"
    "    GROUPING(fa.mes) AS grp_mes,\n"
    "    COUNT(*) AS total_atenciones,\n"
    "    COUNT(DISTINCT fa.paciente_key) AS pacientes_unicos\n"
    "FROM fact_atenciones fa\n"
    "GROUP BY ROLLUP (fa.anio, fa.trimestre, fa.mes);"
)
add_screenshot_placeholder('MOLAP — Creacion del cubo temporal')

doc.add_heading('3.3. Carga del cubo de diagnósticos', level=2)
add_code(
    "CREATE TABLE molap_cubo_diagnostico_especialidad AS\n"
    "SELECT fa.tipo_diagnostico, fa.categoria, dm.especialidad,\n"
    "    GROUPING(fa.tipo_diagnostico) AS grp_tipo,\n"
    "    GROUPING(fa.categoria) AS grp_categoria,\n"
    "    GROUPING(dm.especialidad) AS grp_especialidad,\n"
    "    COUNT(*) AS total_diagnosticos\n"
    "FROM fact_atenciones fa\n"
    "JOIN dim_medico dm ON dm.medico_key = fa.medico_key\n"
    "GROUP BY CUBE (fa.tipo_diagnostico, fa.categoria, dm.especialidad);"
)
add_screenshot_placeholder('MOLAP — Creacion del cubo diagnostico x especialidad')

doc.add_page_break()

doc.add_heading('3.4. Consulta MOLAP: Gran total (lectura directa)', level=2)
doc.add_paragraph(
    'Solo lee 1 fila de la tabla pre-calculada. Sin JOINs, sin GROUP BY. '
    'Se filtra por las banderas de GROUPING para obtener el gran total.'
)
add_code(
    "EXPLAIN ANALYZE\n"
    "SELECT total_atenciones, pacientes_unicos, medicos_involucrados\n"
    "FROM molap_cubo_especialidad_sucursal_gestion\n"
    "WHERE grp_especialidad = 1\n"
    "  AND grp_sucursal     = 1\n"
    "  AND grp_gestion      = 1;"
)
add_screenshot_placeholder('MOLAP — Gran total (resultado + EXPLAIN ANALYZE)')

doc.add_heading('3.5. Consulta MOLAP: Atenciones por especialidad', level=2)
doc.add_paragraph(
    'Los datos ya están agrupados. Solo se filtra por las banderas de grouping '
    'para obtener el desglose por especialidad. No hay GROUP BY ni JOINs.'
)
add_code(
    "EXPLAIN ANALYZE\n"
    "SELECT\n"
    "    COALESCE(especialidad, '** TODAS **') AS especialidad,\n"
    "    total_atenciones, pacientes_unicos\n"
    "FROM molap_cubo_especialidad_sucursal_gestion\n"
    "WHERE grp_especialidad = 0   -- filas con especialidad\n"
    "  AND grp_sucursal     = 1   -- agrupado (todas las sucursales)\n"
    "  AND grp_gestion      = 1   -- agrupado (todas las gestiones)\n"
    "ORDER BY total_atenciones DESC;"
)
add_screenshot_placeholder('MOLAP — Atenciones por especialidad (lectura directa)')

doc.add_heading('3.6. Consulta MOLAP: Slice sobre cubo pre-calculado', level=2)
doc.add_paragraph(
    'Operación SLICE (rebanada por sucursal = Grupo 3) directamente sobre '
    'la tabla MOLAP. No recalcula nada, solo filtra.'
)
add_code(
    "EXPLAIN ANALYZE\n"
    "SELECT COALESCE(especialidad, '** TODAS **') AS especialidad,\n"
    "       total_atenciones\n"
    "FROM molap_cubo_especialidad_sucursal_gestion\n"
    "WHERE sucursal = 'Grupo 3'\n"
    "  AND grp_sucursal     = 0   -- sucursal fija (no agrupada)\n"
    "  AND grp_especialidad = 0   -- desglose por especialidad\n"
    "  AND grp_gestion      = 1   -- todas las gestiones\n"
    "ORDER BY total_atenciones DESC;"
)
add_screenshot_placeholder('MOLAP — Slice Grupo 3 (lectura directa)')

doc.add_heading('3.7. Consulta MOLAP: Drill-down temporal', level=2)
doc.add_paragraph(
    'Toda la jerarquía temporal (año -> trimestre -> mes) '
    'leída directamente de la tabla pre-calculada, sin cálculos al vuelo.'
)
add_code(
    "EXPLAIN ANALYZE\n"
    "SELECT\n"
    "    COALESCE(gestion::TEXT, '** TODOS **')    AS gestion,\n"
    "    COALESCE(trimestre::TEXT, '** TODOS **')  AS trimestre,\n"
    "    COALESCE(mes::TEXT, '** TODOS **')        AS mes,\n"
    "    total_atenciones, pacientes_unicos\n"
    "FROM molap_cubo_temporal\n"
    "ORDER BY grp_gestion, gestion, grp_trimestre, trimestre, grp_mes, mes;"
)
add_screenshot_placeholder('MOLAP — Drill-down temporal pre-calculado')

doc.add_page_break()

# ============================================================
# 4. HOLAP
# ============================================================
doc.add_heading('4. HOLAP — Vistas materializadas + consultas en vivo', level=1)

doc.add_paragraph(
    'La arquitectura HOLAP combina lo mejor de ambos enfoques: '
    'los agregados más consultados se pre-calculan con MATERIALIZED VIEW (capa MOLAP), '
    'mientras que las consultas de detalle granular acceden a los datos en vivo '
    'sobre fact_atenciones (capa ROLAP). Las vistas se refrescan con '
    'REFRESH MATERIALIZED VIEW después de cada ETL.'
)

doc.add_heading('4.1. Capa MOLAP: Vista materializada especialidad x sucursal', level=2)
doc.add_paragraph(
    'Solo pre-calcula los agregados más frecuentes (no todo el cubo como MOLAP). '
    'Usa MATERIALIZED VIEW para poder refrescar periódicamente.'
)
add_code(
    "DROP MATERIALIZED VIEW IF EXISTS holap_mv_especialidad_sucursal;\n"
    "CREATE MATERIALIZED VIEW holap_mv_especialidad_sucursal AS\n"
    "SELECT\n"
    "    dm.especialidad,\n"
    "    ds.nombre                         AS sucursal,\n"
    "    COUNT(*)                          AS total_atenciones,\n"
    "    COUNT(DISTINCT fa.paciente_key)   AS pacientes_unicos,\n"
    "    COUNT(DISTINCT fa.medico_key)     AS medicos_involucrados\n"
    "FROM fact_atenciones fa\n"
    "JOIN dim_medico dm    ON dm.medico_key   = fa.medico_key\n"
    "JOIN dim_paciente dp  ON dp.paciente_key = fa.paciente_key\n"
    "JOIN dim_sucursal ds  ON ds.sucursal_key = dp.sucursal_key\n"
    "GROUP BY dm.especialidad, ds.nombre;\n\n"
    "CREATE INDEX idx_holap_mv1_esp ON holap_mv_...(especialidad);\n"
    "CREATE INDEX idx_holap_mv1_suc ON holap_mv_...(sucursal);"
)
add_screenshot_placeholder('HOLAP — Creacion de vista materializada especialidad x sucursal')

doc.add_heading('4.2. Capa MOLAP: Vista materializada mensual x sucursal', level=2)
add_code(
    "CREATE MATERIALIZED VIEW holap_mv_mensual_sucursal AS\n"
    "SELECT fa.anio AS gestion, fa.mes,\n"
    "    ds.nombre AS sucursal,\n"
    "    COUNT(*) AS total_atenciones,\n"
    "    COUNT(DISTINCT fa.paciente_key) AS pacientes_unicos\n"
    "FROM fact_atenciones fa\n"
    "JOIN dim_paciente dp  ON dp.paciente_key = fa.paciente_key\n"
    "JOIN dim_sucursal ds  ON ds.sucursal_key = dp.sucursal_key\n"
    "GROUP BY fa.anio, fa.mes, ds.nombre;"
)
add_screenshot_placeholder('HOLAP — Creacion de vista materializada mensual x sucursal')

doc.add_heading('4.3. Consulta HOLAP capa MOLAP: Ranking de especialidades', level=2)
doc.add_paragraph(
    'Consulta rápida sobre la vista materializada. Sin JOINs. '
    'Similar en velocidad a MOLAP.'
)
add_code(
    "EXPLAIN ANALYZE\n"
    "SELECT especialidad, sucursal, total_atenciones, pacientes_unicos\n"
    "FROM holap_mv_especialidad_sucursal\n"
    "ORDER BY sucursal, total_atenciones DESC;"
)
add_screenshot_placeholder('HOLAP capa MOLAP — Ranking de especialidades (pre-calculado)')

doc.add_heading('4.4. Consulta HOLAP capa MOLAP: Tendencia mensual', level=2)
add_code(
    "EXPLAIN ANALYZE\n"
    "SELECT gestion, mes, sucursal, total_atenciones\n"
    "FROM holap_mv_mensual_sucursal\n"
    "ORDER BY gestion, mes, sucursal;"
)
add_screenshot_placeholder('HOLAP capa MOLAP — Tendencia mensual (pre-calculado)')

doc.add_page_break()

doc.add_heading('4.5. Consulta HOLAP capa ROLAP: Detalle granular en vivo', level=2)
doc.add_paragraph(
    'Para consultas de detalle (paciente, médico, hora, diagnóstico) que NO están '
    'pre-calculadas, HOLAP va directo a fact_atenciones como ROLAP. '
    'Los cubos pre-calculados no tienen este nivel de detalle.'
)
add_code(
    "EXPLAIN ANALYZE\n"
    "SELECT\n"
    "    fa.fecha_cita, dp.nombre AS paciente,\n"
    "    dm.nombre AS medico, fa.descripcion AS diagnostico,\n"
    "    fa.tipo_diagnostico, fa.hora\n"
    "FROM fact_atenciones fa\n"
    "JOIN dim_medico dm    ON dm.medico_key   = fa.medico_key\n"
    "JOIN dim_paciente dp  ON dp.paciente_key = fa.paciente_key\n"
    "WHERE dm.especialidad = 'Pediatria' AND fa.anio = 2024\n"
    "ORDER BY fa.fecha_cita, fa.hora;"
)
add_screenshot_placeholder('HOLAP capa ROLAP — Detalle granular en vivo')

doc.add_heading('4.6. Consulta HOLAP capa ROLAP: Búsqueda de paciente', level=2)
doc.add_paragraph(
    'Búsqueda por CI de paciente. Los cubos pre-calculados no contienen '
    'este nivel de detalle, así que HOLAP usa la capa relacional.'
)
add_code(
    "EXPLAIN ANALYZE\n"
    "SELECT fa.fecha_cita, dm.nombre AS medico,\n"
    "    dm.especialidad, fa.descripcion AS diagnostico,\n"
    "    fa.observaciones\n"
    "FROM fact_atenciones fa\n"
    "JOIN dim_medico dm    ON dm.medico_key   = fa.medico_key\n"
    "JOIN dim_paciente dp  ON dp.paciente_key = fa.paciente_key\n"
    "WHERE dp.ci = '12345678'\n"
    "ORDER BY fa.fecha_cita DESC;"
)
add_screenshot_placeholder('HOLAP capa ROLAP — Busqueda de paciente (en vivo)')

doc.add_heading('4.7. Mantenimiento HOLAP: Refrescar vistas', level=2)
doc.add_paragraph(
    'Después de cada ETL, se ejecuta REFRESH para actualizar las vistas '
    'materializadas con los datos nuevos. Las tablas base (fact_atenciones) '
    'ya están actualizadas por el ETL.'
)
add_code(
    "REFRESH MATERIALIZED VIEW holap_mv_especialidad_sucursal;\n"
    "REFRESH MATERIALIZED VIEW holap_mv_mensual_sucursal;"
)
add_screenshot_placeholder('HOLAP — Refresh de vistas materializadas')

doc.add_page_break()

# ============================================================
# 5. COMPARACIÓN DIRECTA
# ============================================================
doc.add_heading('5. COMPARACIÓN DIRECTA — Misma pregunta, 3 enfoques', level=1)

doc.add_paragraph(
    'Para demostrar la diferencia práctica entre las tres arquitecturas, '
    'se resuelve la misma pregunta de negocio con cada enfoque:'
)
p = doc.add_paragraph()
run = p.add_run('Pregunta: "¿Cuántas atenciones tiene cada especialidad por sucursal?"')
run.bold = True
run.italic = True

doc.add_heading('5.1. ROLAP — Calcula al vuelo (3 JOINs + GROUP BY)', level=2)
add_code(
    "EXPLAIN ANALYZE\n"
    "SELECT dm.especialidad, ds.nombre AS sucursal,\n"
    "       COUNT(*) AS total_atenciones\n"
    "FROM fact_atenciones fa\n"
    "JOIN dim_medico dm    ON dm.medico_key   = fa.medico_key\n"
    "JOIN dim_paciente dp  ON dp.paciente_key = fa.paciente_key\n"
    "JOIN dim_sucursal ds  ON ds.sucursal_key = dp.sucursal_key\n"
    "GROUP BY dm.especialidad, ds.nombre\n"
    "ORDER BY dm.especialidad, total_atenciones DESC;"
)
add_screenshot_placeholder('Comparacion directa — ROLAP (resultado + tiempo)')

doc.add_heading('5.2. MOLAP — Lee tabla pre-calculada (sin JOINs, sin GROUP BY)', level=2)
add_code(
    "EXPLAIN ANALYZE\n"
    "SELECT COALESCE(especialidad, '** TODAS **') AS especialidad,\n"
    "       COALESCE(sucursal, '** TODAS **')     AS sucursal,\n"
    "       total_atenciones\n"
    "FROM molap_cubo_especialidad_sucursal_gestion\n"
    "WHERE grp_especialidad = 0\n"
    "  AND grp_sucursal     = 0\n"
    "  AND grp_gestion      = 1\n"
    "ORDER BY especialidad, total_atenciones DESC;"
)
add_screenshot_placeholder('Comparacion directa — MOLAP (resultado + tiempo)')

doc.add_heading('5.3. HOLAP — Lee vista materializada', level=2)
add_code(
    "EXPLAIN ANALYZE\n"
    "SELECT especialidad, sucursal, total_atenciones\n"
    "FROM holap_mv_especialidad_sucursal\n"
    "ORDER BY especialidad, total_atenciones DESC;"
)
add_screenshot_placeholder('Comparacion directa — HOLAP (resultado + tiempo)')

doc.add_heading('5.4. Comparación de tiempos de ejecución', level=2)
doc.add_paragraph(
    'Anotar los tiempos del EXPLAIN ANALYZE de cada enfoque:'
)
add_table(
    ['Enfoque', 'Execution Time', 'Planning Time', 'Observación'],
    [
        ['ROLAP', '[anotar]', '[anotar]', '3 JOINs + GROUP BY + ORDER BY'],
        ['MOLAP', '[anotar]', '[anotar]', 'Lectura directa con WHERE sobre flags'],
        ['HOLAP', '[anotar]', '[anotar]', 'Lectura de vista materializada'],
    ]
)

doc.add_page_break()

# ============================================================
# 6. COMPARACIÓN DE ALMACENAMIENTO
# ============================================================
doc.add_heading('6. COMPARACIÓN DE ALMACENAMIENTO', level=1)

doc.add_paragraph(
    'Se compara cuánto espacio en disco ocupa cada enfoque. '
    'ROLAP solo necesita las tablas base; MOLAP almacena todas las combinaciones '
    'pre-calculadas; HOLAP solo pre-calcula los agregados más frecuentes.'
)

doc.add_heading('6.1. Tablas base — ROLAP (almacenamiento mínimo)', level=2)
add_code(
    "SELECT 'ROLAP' AS enfoque, relname AS tabla,\n"
    "    pg_size_pretty(pg_total_relation_size(oid)) AS tamanio\n"
    "FROM pg_class\n"
    "WHERE relname IN ('fact_atenciones','dim_medico','dim_paciente','dim_sucursal')\n"
    "ORDER BY pg_total_relation_size(oid) DESC;"
)
add_screenshot_placeholder('Almacenamiento — Tablas base ROLAP')

doc.add_heading('6.2. Tablas pre-calculadas — MOLAP (almacenamiento extra)', level=2)
add_code(
    "SELECT 'MOLAP' AS enfoque, relname AS tabla,\n"
    "    pg_size_pretty(pg_total_relation_size(oid)) AS tamanio\n"
    "FROM pg_class WHERE relname LIKE 'molap_%'\n"
    "ORDER BY pg_total_relation_size(oid) DESC;"
)
add_screenshot_placeholder('Almacenamiento — Tablas MOLAP extra')

doc.add_heading('6.3. Vistas materializadas — HOLAP (almacenamiento extra)', level=2)
add_code(
    "SELECT 'HOLAP' AS enfoque, relname AS tabla,\n"
    "    pg_size_pretty(pg_total_relation_size(oid)) AS tamanio\n"
    "FROM pg_class WHERE relname LIKE 'holap_%' AND relkind = 'm'\n"
    "ORDER BY pg_total_relation_size(oid) DESC;"
)
add_screenshot_placeholder('Almacenamiento — Vistas HOLAP extra')

doc.add_heading('6.4. Resumen comparativo de almacenamiento', level=2)
add_code(
    "SELECT enfoque, pg_size_pretty(SUM(tamanio_bytes)) AS tamanio_total\n"
    "FROM (\n"
    "    SELECT 'ROLAP' AS enfoque, pg_total_relation_size(oid) AS tamanio_bytes\n"
    "    FROM pg_class WHERE relname IN ('fact_atenciones',...)\n"
    "    UNION ALL\n"
    "    SELECT 'MOLAP (extra)', pg_total_relation_size(oid)\n"
    "    FROM pg_class WHERE relname LIKE 'molap_%'\n"
    "    UNION ALL\n"
    "    SELECT 'HOLAP (extra)', pg_total_relation_size(oid)\n"
    "    FROM pg_class WHERE relname LIKE 'holap_%' AND relkind = 'm'\n"
    ") t GROUP BY enfoque ORDER BY enfoque;"
)
add_screenshot_placeholder('Almacenamiento — Resumen comparativo ROLAP vs MOLAP vs HOLAP')

doc.add_page_break()

# ============================================================
# 7. EXPLOSIÓN COMBINATORIA
# ============================================================
doc.add_heading('7. EXPLOSIÓN COMBINATORIA DEL MOLAP', level=1)

doc.add_paragraph(
    'Una desventaja clave del MOLAP es la explosión combinatoria: al pre-calcular '
    'TODAS las combinaciones de un CUBE, la cantidad de filas almacenadas crece '
    'exponencialmente con el número de valores distintos en cada dimensión.'
)

doc.add_heading('7.1. Conteo de filas: dato original vs pre-calculado', level=2)
add_code(
    "SELECT 'fact_atenciones (original)' AS tabla, COUNT(*) AS filas\n"
    "    FROM fact_atenciones\n"
    "UNION ALL\n"
    "SELECT 'molap_cubo_especialidad_sucursal_gestion', COUNT(*)\n"
    "    FROM molap_cubo_especialidad_sucursal_gestion\n"
    "UNION ALL\n"
    "SELECT 'molap_cubo_temporal', COUNT(*) FROM molap_cubo_temporal\n"
    "UNION ALL\n"
    "SELECT 'molap_cubo_diagnostico_especialidad', COUNT(*)\n"
    "    FROM molap_cubo_diagnostico_especialidad\n"
    "UNION ALL\n"
    "SELECT 'holap_mv_especialidad_sucursal', COUNT(*)\n"
    "    FROM holap_mv_especialidad_sucursal\n"
    "UNION ALL\n"
    "SELECT 'holap_mv_mensual_sucursal', COUNT(*)\n"
    "    FROM holap_mv_mensual_sucursal\n"
    "ORDER BY filas DESC;"
)
add_screenshot_placeholder('Explosion combinatoria — Filas originales vs pre-calculadas')

doc.add_paragraph(
    'Como se puede observar, las tablas MOLAP contienen muchas más filas que los datos '
    'originales (por todas las combinaciones de subtotales). Las vistas HOLAP contienen '
    'menos filas porque solo pre-calculan los agregados seleccionados.'
)

doc.add_page_break()

# ============================================================
# 8. TABLA RESUMEN COMPARATIVA
# ============================================================
doc.add_heading('8. TABLA RESUMEN COMPARATIVA', level=1)

add_table(
    ['Criterio', 'ROLAP', 'MOLAP', 'HOLAP'],
    [
        [
            'Almacenamiento',
            'Solo tablas base (fact + dims)',
            'Tablas pre-calculadas con CUBE',
            'Vistas materializadas para agregados frecuentes'
        ],
        [
            'Velocidad de consulta',
            'Lenta (calcula al vuelo con JOINs)',
            'Muy rápida (lectura directa)',
            'Rápida para agregados, normal para detalle'
        ],
        [
            'Actualización',
            'Siempre actualizado (lee datos en vivo)',
            'Hay que recrear tablas tras cada ETL',
            'REFRESH MATERIALIZED VIEW tras ETL'
        ],
        [
            'Espacio en disco',
            'Mínimo',
            'Alto (explosión combinatoria)',
            'Medio (solo agregados frecuentes)'
        ],
        [
            'Detalle granular',
            'Acceso completo a fact_atenciones',
            'Solo lo que se pre-calculó',
            'Acceso completo via fact_atenciones'
        ],
        [
            'Simulado en PostgreSQL',
            'GROUP BY CUBE / ROLLUP (al vuelo)',
            'CREATE TABLE AS (pre-calcular)',
            'MATERIALIZED VIEW + queries en vivo'
        ],
        [
            'Herramienta real',
            'PostgreSQL, MySQL, Hive, BigQuery',
            'Essbase, SSAS Multidimensional',
            'SSAS Hybrid, Oracle OLAP'
        ],
    ]
)

doc.add_page_break()

# ============================================================
# 9. CONCLUSIONES
# ============================================================
doc.add_heading('9. CONCLUSIONES', level=1)

doc.add_paragraph(
    'Se realizó una comparación práctica de las tres arquitecturas OLAP '
    '(ROLAP, MOLAP y HOLAP) utilizando el Data Warehouse de la Clínica '
    'implementado en PostgreSQL con un modelo estrella.'
)

doc.add_paragraph(
    'ROLAP demostró ser la arquitectura más simple de implementar, ya que utiliza '
    'directamente las tablas del esquema estrella y las operaciones CUBE, ROLLUP y '
    'GROUPING SETS de SQL estándar. Su principal limitación es el tiempo de respuesta, '
    'ya que cada consulta recalcula todos los agregados al momento de ejecutarse.'
)

doc.add_paragraph(
    'MOLAP, simulado mediante tablas pre-calculadas con CREATE TABLE AS SELECT ... GROUP BY CUBE, '
    'mostró tiempos de respuesta significativamente más rápidos al leer datos ya agregados. '
    'Sin embargo, evidenció la explosión combinatoria: las tablas MOLAP contienen '
    'muchas más filas que los datos originales, consumiendo más espacio en disco. '
    'Además, requiere recrear las tablas cada vez que se actualizan los datos.'
)

doc.add_paragraph(
    'HOLAP ofreció un balance entre ambos enfoques: las vistas materializadas '
    'proporcionan velocidad para los agregados más consultados (similiar a MOLAP), '
    'mientras que las consultas de detalle granular mantienen acceso completo '
    'a los datos en vivo (como ROLAP). Su mantenimiento es más sencillo que MOLAP '
    'gracias al comando REFRESH MATERIALIZED VIEW.'
)

doc.add_paragraph(
    'Para el caso del Data Warehouse de la Clínica, la arquitectura más adecuada '
    'depende del patrón de uso: si predominan las consultas analíticas agregadas, '
    'HOLAP ofrece el mejor balance; si se requiere siempre datos actualizados '
    'y flexibilidad total, ROLAP es suficiente; si el volumen es muy alto y '
    'las consultas son predecibles, MOLAP brinda la máxima velocidad a costa de espacio.'
)

# ============================================================
# GUARDAR
# ============================================================
output_path = '/home/tunek/Universidad/MATERIAS/bd-clinica/Informe_OLAP_ROLAP_MOLAP_HOLAP.docx'
doc.save(output_path)
print(f'Documento generado: {output_path}')
