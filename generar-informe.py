from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

# ============================================================
# PORTADA
# ============================================================
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('INFORME DE RECONOCIMIENTO Y MIGRACION\nDE BASE DE DATOS')
run.font.size = Pt(22)
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Sistema de Gestion de Clinica Medica')
run.font.size = Pt(14)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Base de Datos - 2026')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(100, 100, 100)

doc.add_page_break()

# ============================================================
# INDICE
# ============================================================
doc.add_heading('Indice', level=1)
indice = [
    'Parte 1: Reconocimiento de la BD Remota (Grupo 4)',
    '   1.1 Datos de conexion',
    '   1.2 Estructura de tablas',
    '   1.3 Mapa de relaciones',
    '   1.4 Observaciones',
    'Parte 2: Migracion de Datos',
    '   2.1 Objetivo',
    '   2.2 Herramienta usada',
    '   2.3 Diferencias entre estructuras',
    '   2.4 Mapeo de tablas',
    '   2.5 Resultados',
    '   2.6 Dificultades y conclusiones',
]
for item in indice:
    doc.add_paragraph(item, style='List Number' if not item.startswith(' ') else 'Normal')

doc.add_page_break()

# ============================================================
# PARTE 1: RECONOCIMIENTO
# ============================================================
doc.add_heading('Parte 1: Reconocimiento de la BD Remota', level=1)

doc.add_heading('1.1 Datos de conexion', level=2)
doc.add_paragraph('Se recibieron las credenciales de solo lectura del Grupo 4 (Benitez Paco Joel), cuya base de datos esta alojada en Neon (PostgreSQL serverless en Azure).')

t = doc.add_table(rows=6, cols=2, style='Table Grid')
t.alignment = WD_TABLE_ALIGNMENT.CENTER
data = [
    ('Campo', 'Valor'),
    ('Host', 'ep-curly-snow-a8psiq7k-pooler.eastus2.azure.neon.tech'),
    ('Puerto', '5432'),
    ('Base de datos', 'neondb'),
    ('Usuario', 'usuario_lectura'),
    ('Version', 'PostgreSQL 17.7 (Neon)'),
]
for i, (k, v) in enumerate(data):
    t.rows[i].cells[0].text = k
    t.rows[i].cells[1].text = v
    if i == 0:
        for cell in t.rows[i].cells:
            cell.paragraphs[0].runs[0].bold = True

doc.add_heading('1.2 Estructura de tablas', level=2)
doc.add_paragraph('La BD remota cuenta con 12 tablas, 24 foreign keys, y usa UUID como primary key en todas las tablas.')

tablas_remotas = [
    ('personas', '45,000', 'Tabla central de todas las personas del sistema'),
    ('pacientes', '31,499', 'Extiende personas con datos clinicos (alergias, enfermedades cronicas)'),
    ('personal_medico', '9,000', 'Extiende personas con datos laborales (especialidad, colegiatura, salario)'),
    ('servicios', '45,000', 'Catalogo de servicios (consulta, cirugia, laboratorio, imagen)'),
    ('diagnosticos', '45,000', 'Catalogo de diagnosticos con codigo CIE-10'),
    ('citas', '45,000', 'Citas medicas programadas'),
    ('atenciones', '45,000', 'Atenciones realizadas (vinculadas a citas)'),
    ('diagnosticos_atencion', '112,606', 'Tabla pivote N:M entre atenciones y diagnosticos'),
    ('prescripciones', '45,000', 'Recetas con medicamento, dosis, frecuencia'),
    ('procedimientos', '13,549', 'Procedimientos quirurgicos o medicos'),
    ('facturacion', '45,000', 'Facturacion por servicios prestados'),
    ('ausencias', '45,000', 'Control de ausencias del personal medico'),
]

t = doc.add_table(rows=len(tablas_remotas)+1, cols=3, style='Table Grid')
t.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ('Tabla', 'Registros', 'Descripcion')
for j, h in enumerate(headers):
    t.rows[0].cells[j].text = h
    t.rows[0].cells[j].paragraphs[0].runs[0].bold = True
for i, (tabla, reg, desc) in enumerate(tablas_remotas):
    t.rows[i+1].cells[0].text = tabla
    t.rows[i+1].cells[1].text = reg
    t.rows[i+1].cells[2].text = desc

doc.add_heading('1.3 Mapa de relaciones', level=2)
doc.add_paragraph('Patron de herencia: personas es la tabla padre con pacientes y personal_medico como tablas hijas (relacion 1:1 via UNIQUE en id_persona).')

relaciones = [
    ('personas -> pacientes', 'id_persona', 'ON DELETE CASCADE'),
    ('personas -> personal_medico', 'id_persona', 'ON DELETE CASCADE'),
    ('pacientes -> citas', 'id_paciente', ''),
    ('pacientes -> atenciones', 'id_paciente', ''),
    ('pacientes -> prescripciones', 'id_paciente', ''),
    ('pacientes -> procedimientos', 'id_paciente', ''),
    ('pacientes -> facturacion', 'id_paciente', ''),
    ('personal_medico -> citas', 'id_personal', ''),
    ('personal_medico -> atenciones', 'id_medico', ''),
    ('personal_medico -> prescripciones', 'id_medico', ''),
    ('personal_medico -> procedimientos', 'id_cirujano', ''),
    ('personal_medico -> ausencias', 'id_personal', ''),
    ('servicios -> citas', 'id_servicio', ''),
    ('servicios -> procedimientos', 'id_servicio', ''),
    ('servicios -> facturacion', 'id_servicio', ''),
    ('diagnosticos -> atenciones', 'diagnostico_principal', ''),
    ('diagnosticos -> diagnosticos_atencion', 'id_diagnostico', ''),
    ('diagnosticos -> procedimientos', 'diagnostico_asociado', ''),
    ('diagnosticos -> ausencias', 'diagnostico_asociado', ''),
    ('citas -> atenciones', 'id_cita', ''),
    ('atenciones -> diagnosticos_atencion', 'id_atencion', 'ON DELETE CASCADE'),
    ('atenciones -> prescripciones', 'id_atencion', ''),
    ('atenciones -> procedimientos', 'id_atencion', ''),
    ('atenciones -> facturacion', 'id_atencion', ''),
]

t = doc.add_table(rows=len(relaciones)+1, cols=3, style='Table Grid')
t.alignment = WD_TABLE_ALIGNMENT.CENTER
for j, h in enumerate(('Relacion', 'Columna FK', 'Nota')):
    t.rows[0].cells[j].text = h
    t.rows[0].cells[j].paragraphs[0].runs[0].bold = True
for i, (rel, col, nota) in enumerate(relaciones):
    t.rows[i+1].cells[0].text = rel
    t.rows[i+1].cells[1].text = col
    t.rows[i+1].cells[2].text = nota

doc.add_heading('1.4 Observaciones', level=2)
obs = [
    'Todas las PKs son UUID con gen_random_uuid(), no usan SERIAL.',
    'Usan CHECK constraints para validar estados (ej: PROGRAMADA, CONFIRMADA, COMPLETADA, CANCELADA, NO_ASISTIO).',
    'No tienen documentacion interna (COMMENT ON). Las tablas no tienen descripcion.',
    'Tienen indices en campos de fecha (idx_citas_fecha, idx_atenciones_fecha, idx_facturacion_fecha).',
    'Solo 2 tablas tienen datos masivos reales (diagnosticos y servicios con 45,000 c/u). El resto fue generado.',
    'El usuario de lectura permite ver todos los roles del sistema via pg_roles (cloud_admin es SUPERUSER).',
    'La version PostgreSQL 17.7 esta parcheada contra los CVEs conocidos (CVE-2025-12818, CVE-2025-12817, CVE-2025-4207).',
]
for o in obs:
    doc.add_paragraph(o, style='List Bullet')

doc.add_page_break()

# ============================================================
# PARTE 2: MIGRACION
# ============================================================
doc.add_heading('Parte 2: Migracion de Datos', level=1)

doc.add_heading('2.1 Objetivo', level=2)
doc.add_paragraph('Extraer datos de la BD del Grupo 4 (Neon) e insertarlos en nuestra BD local (Docker/PostgreSQL), adaptando la estructura de sus tablas a la nuestra.')

doc.add_heading('2.2 Herramienta usada: dblink', level=2)
doc.add_paragraph('Se uso la extension dblink de PostgreSQL, que permite hacer consultas SQL a una base de datos remota como si fuera una tabla local:')
doc.add_paragraph('CREATE EXTENSION dblink;\nSELECT * FROM dblink(conexion_remota, consulta_sql) AS t(columnas...);', style='Normal').runs[0].font.size = Pt(9)
doc.add_paragraph('Esto permite extraer e insertar en un solo paso, sin archivos intermedios.')

doc.add_heading('2.3 Diferencias entre estructuras', level=2)

diffs = [
    ('Primary Keys', 'UUID (gen_random_uuid)', 'INT (SERIAL autoincremental)'),
    ('Personas', '3 tablas: personas + pacientes + personal_medico', '1 tabla: PERSONA (Matricula distingue rol)'),
    ('Especialidad', 'Columna VARCHAR en personal_medico', 'Tabla separada ESPECIALIDAD'),
    ('Zonas', 'No tienen', 'Tabla ZONA'),
    ('Horarios', 'No tienen', 'Tabla HORARIO_MEDICO'),
    ('Diagnosticos', 'Tabla catalogo (nombre, CIE-10, flags)', 'Tabla transaccional (por cita)'),
    ('Recetas', 'prescripciones (medicamento, dosis, frecuencia separados)', 'RECETA (texto libre)'),
    ('Facturacion', 'Tabla FACTURACION', 'No tenemos'),
    ('Ausencias', 'Tabla AUSENCIAS', 'No tenemos'),
]

t = doc.add_table(rows=len(diffs)+1, cols=3, style='Table Grid')
t.alignment = WD_TABLE_ALIGNMENT.CENTER
for j, h in enumerate(('Aspecto', 'Ellos (Neon)', 'Nosotros (Docker)')):
    t.rows[0].cells[j].text = h
    t.rows[0].cells[j].paragraphs[0].runs[0].bold = True
for i, (asp, ellos, nosotros) in enumerate(diffs):
    t.rows[i+1].cells[0].text = asp
    t.rows[i+1].cells[1].text = ellos
    t.rows[i+1].cells[2].text = nosotros

doc.add_heading('2.4 Mapeo de tablas', level=2)
doc.add_paragraph('La clave para migrar fue identificar campos equivalentes entre ambos modelos. El numero de documento (CI) fue el campo natural comun para vincular personas.')

doc.add_heading('ESPECIALIDAD', level=3)
doc.add_paragraph('Ellos no tienen tabla de especialidades. Es un VARCHAR suelto en personal_medico. Se extrajo con SELECT DISTINCT y se inserto como catalogo.')

doc.add_heading('ZONA', level=3)
doc.add_paragraph('Ellos no tienen zonas. Se creo una zona generica "Zona Migrada" para no violar la FK obligatoria en PERSONA.')

doc.add_heading('PERSONA (medicos)', level=3)
mapeo_med = [
    ('personas.numero_documento', 'PERSONA.CI'),
    ('personas.nombre_completo', 'PERSONA.Nombre'),
    ('personas.fecha_nacimiento', 'PERSONA.Fecha_Nacimiento'),
    ('personas.genero (M/F/OTRO)', 'PERSONA.Sexo (M/F)'),
    ('personas.direccion', 'PERSONA.Direccion'),
    ('personas.telefono', 'PERSONA.Telefono'),
    ('personal_medico.numero_colegiatura', 'PERSONA.Matricula (NOT NULL = medico)'),
    ('personal_medico.especialidad', 'PERSONA.ID_Especialidad (via lookup)'),
]
t = doc.add_table(rows=len(mapeo_med)+1, cols=2, style='Table Grid')
t.alignment = WD_TABLE_ALIGNMENT.CENTER
for j, h in enumerate(('Origen (Neon)', 'Destino (nuestra BD)')):
    t.rows[0].cells[j].text = h
    t.rows[0].cells[j].paragraphs[0].runs[0].bold = True
for i, (orig, dest) in enumerate(mapeo_med):
    t.rows[i+1].cells[0].text = orig
    t.rows[i+1].cells[1].text = dest

doc.add_heading('PERSONA (pacientes)', level=3)
doc.add_paragraph('Mismo mapeo que medicos pero con Matricula = NULL e ID_Especialidad = NULL.')

doc.add_heading('CITA_MEDICA', level=3)
mapeo_cita = [
    ('citas.fecha_solicitud::DATE', 'CITA_MEDICA.Fecha_Registro'),
    ('citas.fecha_cita', 'CITA_MEDICA.Fecha_Cita'),
    ('citas.hora_cita', 'CITA_MEDICA.Hora'),
    ('ROW_NUMBER() (generado)', 'CITA_MEDICA.Numero_Turno'),
    ('citas.estado', 'CITA_MEDICA.Estado'),
    ('lookup por CI del paciente', 'CITA_MEDICA.ID_Paciente'),
    ('lookup por CI del medico', 'CITA_MEDICA.ID_Medico'),
]
t = doc.add_table(rows=len(mapeo_cita)+1, cols=2, style='Table Grid')
t.alignment = WD_TABLE_ALIGNMENT.CENTER
for j, h in enumerate(('Origen (Neon)', 'Destino (nuestra BD)')):
    t.rows[0].cells[j].text = h
    t.rows[0].cells[j].paragraphs[0].runs[0].bold = True
for i, (orig, dest) in enumerate(mapeo_cita):
    t.rows[i+1].cells[0].text = orig
    t.rows[i+1].cells[1].text = dest

doc.add_heading('DIAGNOSTICO', level=3)
doc.add_paragraph('Se mapeo atenciones.plan_tratamiento a Descripcion y atenciones.pronostico a Observaciones. La categoria del diagnostico remoto se uso para buscar el ID_Tipo_Diagnostico local.')

doc.add_heading('RECETA', level=3)
doc.add_paragraph('Se concateno prescripciones.medicamento + dosis en Medicamentos, y frecuencia + duracion_dias en Indicaciones.')

doc.add_heading('2.5 Resultados', level=2)

resultados = [
    ('ESPECIALIDAD', '15', '4', '19'),
    ('TIPO_DIAGNOSTICO', '25', '5', '30'),
    ('ZONA', '20', '1', '21'),
    ('PERSONA', '5,000', '6,000', '11,000'),
    ('CITA_MEDICA', '20,000', '87', '20,087'),
    ('DIAGNOSTICO', '15,000', '0', '15,000'),
    ('RECETA', '7,440', '0', '7,440'),
]

t = doc.add_table(rows=len(resultados)+1, cols=4, style='Table Grid')
t.alignment = WD_TABLE_ALIGNMENT.CENTER
for j, h in enumerate(('Tabla', 'Propios', 'Migrados', 'Total')):
    t.rows[0].cells[j].text = h
    t.rows[0].cells[j].paragraphs[0].runs[0].bold = True
for i, (tabla, prop, mig, tot) in enumerate(resultados):
    t.rows[i+1].cells[0].text = tabla
    t.rows[i+1].cells[1].text = prop
    t.rows[i+1].cells[2].text = mig
    t.rows[i+1].cells[3].text = tot

doc.add_heading('2.6 Dificultades y conclusiones', level=2)

doc.add_heading('Por que migraron pocos diagnosticos y recetas', level=3)
doc.add_paragraph('Los diagnosticos y recetas dependen de una cadena de JOINs: receta -> diagnostico -> cita -> paciente + medico. Para migrar una receta se necesita que la cita ya exista con coincidencia exacta de paciente, medico, fecha y hora. Como ambas BDs tienen datos generados aleatoriamente, las coincidencias fueron minimas.')

doc.add_heading('Que hizo posible la migracion', level=3)
factores = [
    'Ambas BDs modelan el mismo dominio (clinica medica) con tablas equivalentes.',
    'El numero de documento (CI) es el campo comun que vincula personas entre sistemas.',
    'dblink permite consultar la BD remota como tabla local, sin archivos intermedios.',
    'Las FKs se resuelven por lookup: no se copian UUIDs, se busca el registro equivalente por CI y se asigna nuestro ID entero.',
]
for f in factores:
    doc.add_paragraph(f, style='List Bullet')

doc.add_heading('Que complica la migracion', level=3)
problemas = [
    'PKs diferentes (UUID vs SERIAL): no se pueden copiar directo.',
    'Estructura diferente: ellos separan personas en 3 tablas, nosotros en 1.',
    'Campos que no existen en el otro modelo: ellos no tienen zonas ni horarios, nosotros no tenemos facturacion ni ausencias.',
    'Datos dependientes: migrar diagnosticos y recetas requiere que toda la cadena previa ya exista.',
]
for p in problemas:
    doc.add_paragraph(p, style='List Bullet')

doc.save('/home/tunek/Universidad/MATERIAS/bd-clinica/informe-reconocimiento-migracion.docx')
print('Documento generado: informe-reconocimiento-migracion.docx')
