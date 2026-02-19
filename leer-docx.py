from docx import Document
from docx.oxml.ns import qn
import json

doc = Document('/home/tunek/Universidad/MATERIAS/bd-clinica/Informe_Hospital_NE.docx')

print("=== PROPIEDADES ===")
print(f"Secciones: {len(doc.sections)}")
print(f"Paragraphs: {len(doc.paragraphs)}")
print(f"Tables: {len(doc.tables)}")

print("\n=== IMAGENES ===")
for rel in doc.part.rels.values():
    if "image" in rel.reltype:
        print(f"  Imagen: {rel.target_ref} (tipo: {rel.reltype.split('/')[-1]})")

print("\n=== CONTENIDO COMPLETO ===")
for i, p in enumerate(doc.paragraphs):
    style = p.style.name if p.style else 'None'
    text = p.text.strip()
    has_image = bool(p._element.findall('.//' + qn('wp:inline'))) or bool(p._element.findall('.//' + qn('wp:anchor')))
    if text or has_image:
        img_tag = " [IMAGEN]" if has_image else ""
        print(f"[{i}] ({style}){img_tag}: {text}")

print("\n=== TABLAS ===")
for ti, table in enumerate(doc.tables):
    print(f"\nTabla {ti} ({len(table.rows)} filas x {len(table.columns)} cols):")
    for ri, row in enumerate(table.rows):
        cells = [cell.text.strip().replace('\n', ' | ') for cell in row.cells]
        print(f"  Fila {ri}: {cells}")
